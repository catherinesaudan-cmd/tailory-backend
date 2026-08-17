"""
Tailory Backend — Pipeline documentaire pédagogique
(numéro de version : UNE seule déclaration, la constante VERSION ci-dessous —
 l'en-tête n'en porte plus ; voir JOURNAL BACKEND v2.23, défaut de livraison
 « le fichier annonçait 2.18 »)
FastAPI + python-docx + pdf2docx + Anthropic proxy + pictogrammes ARASAAC

Endpoints:
  POST /pictos   → résolution de mots-clés en pictogrammes ARASAAC (base64)
                   (chantier prioritaire validé : supports visuels du mode
                    Participation — non-lecteur, non-verbal/CAA, allophone)
  POST /parse    → DOCX/PDF/ODT → structure JSON pédagogique
                   (ODT : converti en PDF via LibreOffice puis pipeline PDF —
                    indispensable car les ODT contiennent souvent des formes
                    vectorielles natives invisibles pour l'extracteur DOCX)
  POST /generate → proxy Anthropic avec retry + chunking
  POST /export   → structure JSON adaptée → DOCX
  POST /convert  → PDF → DOCX (existant, conservé)
  GET  /health   → vérification

V2.11 (chantier 6.8 — segments courts étiquetés, 31 juillet 2026) :
  · Un trait fin de 7 à 14 mm devient une figure candidate s'il porte une
    ÉTIQUETTE DE SÉRIE juste à sa gauche — une lettre ou un chiffre seul
    suivi d'une parenthèse ou d'un point : a) b) e) 1) 2. — ET qu'un autre
    trait étiqueté de longueur différente (> 5 %) existe sur la même page.
    Mesuré sur essai_10117 : le segment e) fait 8,0 mm dans la source,
    sous le plancher de 14,1 mm (40 pt) qui écarte tirets, puces et
    soulignements ; dans un exercice de MESURE, un trait court qui porte
    sa lettre est une vraie figure, pas un parasite.
  · Ce qui reste dehors, et pourquoi : les tirets et puces (jamais
    étiquetés) ; les lignes de réponse d'une liste numérotée « 1) ___ »
    (étiquetées mais TOUTES de la même longueur — les segments d'un
    exercice de mesure ont des longueurs toutes différentes, c'est le
    principe de l'exercice) ; un trait court étiqueté isolé sans compagnon
    (le plancher de 14,1 mm garde le dernier mot) ; l'étiquette au-dessus
    du trait (trop proche d'un mot souligné ou d'un numérateur — porte
    fermée pour cette version). Les promus subissent ensuite les mêmes
    filtres anti-bruit que les longs (isolement, souligné, familles) et le
    filtre v2.5 des pages de mesure.

V2.10 (grilles et formes composites, 27 juillet 2026) :
  · Les quadrillages et les tableaux d'une source ne partent plus en miettes.
    Trois symptômes réparés, tous mesurés sur un cas d'essai fabriqué :
      - un quadrillage vide disparaissait entièrement (ses traits réguliers
        étaient éliminés comme du bruit) : le modèle recevait des taches
        coloriées sans repère pour compter, d'où les aires inventées ;
      - un tableau ressortait en file de mots, sans plus dire quelle valeur
        allait avec quelle entrée (« 36 / 24 / 1 2 3 … 12 ») ;
      - ce même tableau repartait AUSSI en figure découpée : vu deux fois,
        une fois juste et une fois faux.
    Le serveur ne résout rien : il mesure et il le dit. Une grille produit sa
    forme (« quadrillage 8 colonnes x 6 lignes, case 10,0 mm »), son contenu
    si elle porte du texte, et le nombre de cases pleines de chaque figure —
    ce dernier UNIQUEMENT si les figures épousent la grille. Une forme
    oblique, un triangle, un décalage : rien n'est annoncé, et la raison du
    silence est transmise. Un comptage douteux est plus dangereux qu'absent.
    Module `grilles.py`, batterie `test_grilles_v210.py` (21 cas, à vide).
    Si le module est absent au déploiement, le pipeline retombe à l'identique
    sur le comportement 2.9.3 — aucune panne, seulement l'ancien défaut.

V2.9.3 (expérience du mode dégradé, 26 juillet 2026) :
  · PDF_B64_MAX : le plafond d'attachement du PDF converti (ODT) passe de
    150 000 à 4 000 000 caractères base64. Il doublait en silence le plafond
    du frontend : remonter celui du frontend seul ne changeait rien, le PDF
    ne quittait jamais le backend. Quand le plafond est dépassé, la réponse
    porte désormais « pdf_b64_skipped_bytes » — plus de rejet muet.

V2.9.2 (chantier ARASAAC seul — périmètre validé par Catherine) :
  · POST /pictos : {"mots": ["araignée", …], "lang": "fr"} → pour chaque mot,
    le meilleur pictogramme ARASAAC en data-URL PNG 300 px, prêt à injecter
    dans le HTML (l'export reste autonome, aucune dépendance réseau à
    l'impression). Recherche via bestsearch avec repli search ; normalisation
    des mots (minuscules, déterminants élidés « l'araignée » → « araignée ») ;
    langues fr/de/en… (code ARASAAC).
  · CACHE deux niveaux, zéro dépendance nouvelle (urllib stdlib) :
    mémoire (mot→id, id→png) + disque (ARASAAC_CACHE_DIR, défaut
    /tmp/arasaac_cache — les pictos sont immuables). Un mot introuvable est
    aussi mis en cache (négatif, TTL session) pour ne pas re-frapper l'API.
  · Limites de sécurité : 24 mots max par requête, timeout 8 s par appel
    ARASAAC, mots introuvables → null (le frontend applique ses replis
    émoji/sans-image). Meta d'attribution CC BY-NC-SA incluse dans chaque
    réponse (obligation de licence — ligne à imprimer en pied de fiche).
  · /health expose arasaac:"ok"/"unreachable" (sonde légère, 1 requête test
    mise en cache) pour vérifier l'accès réseau depuis Render au déploiement.

V2.9.1 (retour essai 44 — régression Pacôme corrigée) :
  · COMPOSITES = SCÈNES SEULEMENT : la garde par paires de la V2.9 laissait
    fusionner les grands rasters qui se CHEVAUCHENT (scans à marges
    blanches) — les dessins Pacôme/Momo ont refusionné (161×153 mm) en
    avalant le fragment « er, hululer… », le bug même de l'essai 39.
    Un composite exige désormais >= 3 membres tous petits (min-dim
    <= 120 pt) ; toute paire et tout composant contenant un grand raster
    sont dissous en rasters individuels (comportement V2.8 restauré pour
    les scans/photos/cliparts, composites conservés pour les scènes :
    pièces, billets, tas de tomates, vignettes).

V2.9 (audit essai 43, validé localement sur EvalSerie1maths5P.odt — 168 → 93 figures) :
  · CLUSTERING RASTER-RASTER : la règle V2.8 « rasters autonomes » pulvérisait
    les scènes composées de nombreux petits rasters (série 5P : 16 pièces de
    monnaie, 15 tas de tomates, ~19 billets, 17 vignettes… = 128 rasters) ;
    le plafond de miniatures du frontend saturait et le modèle déclarait
    « non disponibles » des images existantes (essai 43 : 154 figures
    inutilisées, 12 placeholders « coller ici »). Les rasters se regroupent
    désormais ENTRE EUX (_cluster_rasters, marge 16 pt calibrée sur les
    écarts mesurés : intra-scène 0-14,6 pt, figures distinctes ≥ 18 pt),
    jamais avec les vecteurs — l'acquis V2.8 est conservé.
  · GARDE PHOTOS : deux grands rasters (min-dim > 120 pt ≈ 42 mm) sans
    recouvrement ne fusionnent jamais — les photos côte à côte restent
    autonomes (essai 42 insectes : 15/15 photos, à ne pas casser).
  · SCISSION TABLEAU : un composite multi-rasters est re-scindé à chaque
    bordure horizontale de tableau qui le traverse sans couper aucun membre
    (p5 de la 5P : écarts intra/inter-collections indistinguables — seules
    les bordures séparent les 4 collections, retrouvées une à une : 1, 8,
    4 et 3 pièces).

V2.8 (retour essai 39, validé localement sur Pacomefantome.odt) :
  · RASTERS AUTONOMES : une image raster (scan, photo, dessin importé) est
    déjà une unité sémantique complète — elle n'est PLUS clusterisée, ni avec
    les autres rasters, ni avec les vecteurs. Cas prouvé (essai 39) : trois
    dessins scannés aux marges blanches se chevauchant fusionnaient en un
    bloc de 180×134 mm dont la bbox avalait les lignes basses de la banque
    de mots — le fragment de texte tronqué « er, hululer… » partait comme
    figure et était placé tel quel sur la fiche. Seuls les rasters quasi
    identiques (recouvrement ≥ 80 %) sont dédupliqués. Le clustering ne
    s'applique désormais qu'aux tracés VECTORIELS (une figure = plusieurs
    primitives), sa raison d'être.

V2.7 (retours essais 36-37, validé localement sur AireDefinitionetMesure.odt) :
  · COMPOSANTES CONNEXES : le clustering par union de bounding-box était
    structurellement vorace — chaque fusion créait une boîte plus grande qui
    absorbait tout ce qu'elle SURVOLAIT sans en être réellement proche.
    Cas prouvé (essais 36-37) : le trait séparateur pointillé passait à
    10,5 pt du tissu E → fusion → la bbox résultante (x45-185 mm)
    intersectait D, puis C, puis B → composite unique de 169×109 mm,
    impossible à placer figure par figure côté modèle (placeholders
    contradictoires, « B et E » alors que C et D y étaient aussi).
    Désormais : composantes connexes calculées sur les primitives D'ORIGINE
    (arête si écart < marge), bbox par composante calculée À LA FIN.
    Résultat validé : les 5 tissus A-E sortent individuellement.
  · MARGE 14 → 10 pt : l'écart réel entre les tissus B et C est de 12,2 pt ;
    les sous-éléments d'une même figure (cellules de grille, hachures) se
    touchent ou restent < 10 pt. Aucune figure légitime connue du corpus
    n'a d'écart interne de 10-14 pt.
  · SÉPARATEURS PARTIELS : une RANGÉE de barres fines (h < 16 pt) alignées
    couvrant ensemble > 50 % de la largeur de page, sans voisin plein,
    est un trait de section coupé par du texte — exclue avant clustering.
    Le filtre pleine largeur (0.92·pw) ne voyait pas ces morceaux courts.
  · BANDES DE TEXTE : une bande plate (h < 32 pt ≈ 11 mm) contenant du texte
    (≥ 2 mots de ≥ 3 lettres, ou 1 mot couvrant > 30 % de sa largeur) est
    une cellule d'en-tête de tableau ou un titre décoré, PAS une figure —
    exclue avant clustering. Validé : « Figures | Calcul de l'aire »,
    « Exercice N », « L'aire du triangle », « Définition N » — précisément
    les bandeaux gris qui fuyaient dans les fiches des essais 36-37.
    Les figures plates légitimes survivent : tissus sans texte (lettres
    vectorielles), lignes graduées (chiffres seulement).

V2.3 (retours essais 25-29) :
  · FONDS DE PAGE : les rectangles couvrant > 85 % de la page (LibreOffice
    exporte un fond blanc pleine page selon le modèle de document) sont exclus
    AVANT clustering — en v2.2 un tel fond absorbait toutes les régions en un
    cluster pleine page, jeté ensuite → 0 figure transmise (cas essai 27 :
    les 15 images d'animaux et les schémas existaient dans la source).
  · RE-SPLIT : un cluster qui couvre malgré tout > 85 % de la page est
    re-clusterisé plus finement (marge 5) au lieu d'être jeté ; en dernier
    recours, ses primitives pleines sont conservées individuellement.
  · LIGNES FINES : les segments à mesurer et lignes graduées (une dimension
    quasi nulle, l'autre ≥ 40 pt) sont désormais capturés comme figures —
    en v2.2 le filtre « dim < 10 » les jetait (cas essai 29 : segments a-f
    « non disponibles » alors qu'ils existaient). Ils partent avec leur
    taille physique réelle (w_mm/h_mm du clip rasterisé) pour l'impression
    à l'échelle (class="img-echelle") — mesure à la règle valide.
    Anti-bruit, une ligne fine n'est une figure QUE si :
      - elle est isolée (rien à moins de 3 pt — élimine les bordures de
        tableaux, qui se croisent aux coins),
      - elle n'est pas un souligné de texte (texte juste au-dessus couvrant
        ≥ 60 % de sa longueur),
      - elle n'appartient pas à une famille de ≥ 3 lignes parallèles de même
        longueur (lignes d'écriture, grilles) — les segments à mesurer ont
        des longueurs toutes différentes, c'est le principe de l'exercice.
"""

from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import anthropic
import base64
import io
import json
import os
import re
import time
import uuid
import zipfile
import glob
import tempfile
import subprocess
from typing import Optional

# python-docx
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# pdf2docx (conversion PDF→DOCX existante)
from pdf2docx import Converter

# V2.9.3 — plafond d'attachement du PDF converti, en caractères base64.
# Doit rester égal à PDF_B64_MAX du frontend (tailoryv10_62.html) : un plafond
# plus bas ici filtrerait en amont, invisiblement, quoi que fasse le frontend.
PDF_B64_MAX = 4_000_000  # ~3 Mo de PDF, une trentaine de pages illustrées

# Seule déclaration du numéro de version du backend. /health la LIT — jamais
# recopiée à la main ailleurs (même règle que pour grilles/formes ci-dessous).
# (voir JOURNAL BACKEND v2.23)
VERSION = "2.35"

# ═══════════════════════════════════════════════════════════════════════════
# v2.34 — C11 : UN CADRE SANS DESSIN N'EST PAS UNE FIGURE
#
# Condition : dans le cadre découpé, aucun objet dessiné.
# Effet    : le cadre n'est pas transmis au modèle.
# Le POURQUOI, les chiffres, le panel et les témoins : JOURNAL BACKEND v2.34.
#
# CE QUE CETTE FONCTION LIT, et qui existe bien à l'étape où elle tourne
# (§ 2.5 du protocole, vérifié dans le code et non de mémoire) :
#   · page.get_drawings()      — primitives de tracé de la page
#   · page.get_text("dict")    — blocs de texte et blocs d'image de la page
#   · rasters                  — images collées, déjà constituées en amont
# Aucune déclaration d'une autre couche, aucun mot, aucune police, aucun nom
# de fichier : deux proportions géométriques. Ce n'est PAS une reconnaissance
# d'apparence au sens du § 2.1 — elle lit les primitives de dessin de la
# SOURCE, pas ce que le modèle écrit.
# ═══════════════════════════════════════════════════════════════════════════
C11_AIRE_MIN_PT2 = 20.0    # sous cette aire, un tracé est un filet, pas un objet
C11_PART_TEXTE = 0.10      # palier mesuré : 5 %, 10 % et 15 % donnent le même
                           # résultat sur les 20 témoins — seuil au milieu
C11_PLAFOND_CADRE = 0.60   # un tracé qui couvre plus de 60 % du cadre EST le
                           # cadre — MAIS seulement si le cadre porte du texte


def _v234_porte_un_dessin(page, clip, rasters):
    """Vrai si le cadre porte au moins un objet dessiné (image collée ou
    tracé qui n'est pas le contenant d'un texte)."""
    dic = page.get_text("dict")
    spans = [sp for b in dic["blocks"] if b["type"] == 0
             for l in b["lines"] for sp in l["spans"]]
    dedans = [sp for sp in spans if clip.intersects(fitz.Rect(sp["bbox"]))]
    # une image collée est toujours un dessin
    for ra in rasters:
        if not (ra & clip).is_empty and (ra & clip).get_area() >= 0.5 * ra.get_area():
            return True
    aire_clip = clip.get_area()
    for d in page.get_drawings():
        r = fitz.Rect(d["rect"])
        if not clip.contains(r) or r.width < 3 or r.height < 3:
            continue
        if r.get_area() < C11_AIRE_MIN_PT2:
            continue
        # un cadre qui n'est QU'une forme, c'est la forme elle-même : le
        # plafond ne s'applique que si le cadre porte du texte.
        if dedans and r.get_area() > C11_PLAFOND_CADRE * aire_clip:
            continue
        aire_txt = sum((fitz.Rect(sp["bbox"]) & r).get_area()
                       for sp in spans if r.intersects(fitz.Rect(sp["bbox"])))
        if aire_txt / r.get_area() < C11_PART_TEXTE:
            return True
    return False

# v2.28 — LE SERVEUR SIGNE SON CONTENU, PAS SEULEMENT SON NOM (11.08.2026).
# Condition : le module se charge. Effet : l'empreinte du fichier lui-même est
# calculée UNE fois et servie par /health. Une étiquette n'est pas le code (leçon
# de la 2.18 puis du tirage 10371 : « même numéro, contenu différent ») ; la
# signature rend la divergence visible d'un regard. (voir JOURNAL BACKEND v2.28)
import hashlib as _hl
try:
    with open(__file__, "rb") as _f:
        EMPREINTE = _hl.sha256(_f.read()).hexdigest()[:12]
except Exception:
    EMPREINTE = "illisible"
# v2.29 — LE SERVEUR ANNONCE AUSSI AVEC QUOI IL LIT (11.08.2026). Condition : le
# module se charge. Effet : la version de la bibliothèque d'extraction est lue
# UNE fois et servie par /health — deux machines qui comptent différemment se
# voient d'un regard. Née de l'écart 12/11 du tirage 10371, environnement de
# Render non lisible autrement. (voir JOURNAL BACKEND v2.29)
try:
    import pymupdf as _pm
    LECTEUR_PDF = "pymupdf " + getattr(_pm, "__version__", "?")
except Exception:
    try:
        LECTEUR_PDF = (fitz.__doc__ or "pymupdf ?").split(":")[0].strip().lower()
    except Exception:
        LECTEUR_PDF = "illisible"

app = FastAPI(title="Tailory Backend V2")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://classedobsdecath.ch",
        "http://localhost",
        "http://127.0.0.1",
        "null"  # v2.6 — fichier tailory ouvert en local (file://) : Origin: null
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────────
# MODÈLES DE TYPES D'EXERCICES
# ─────────────────────────────────────────────
EXERCISE_KEYWORDS = {
    "relier":     ["reli", "associe", "associer", "relie", "relient"],
    "entourer":   ["entoure", "entourer", "encercle", "encercler", "barre", "barrer", "souligne"],
    "compléter":  ["complète", "compléter", "écris", "écrire", "inscris", "note"],
    "numéroter":  ["numérote", "numéroter", "numérotez", "numérotes"],
    "classer":    ["découpe", "découper", "colle", "coller", "classe", "classer", "range", "ranger"],
    "dessiner":   ["dessine", "dessiner", "dessinez", "colorie", "colorier"],
    "observer":   ["observe", "observer", "regardes", "regarde"],
    "lire":       ["lis", "lire", "lisez", "lecture"],
}

def detect_exercise_type(text: str) -> str:
    text_lower = text.lower()
    for ex_type, keywords in EXERCISE_KEYWORDS.items():
        for kw in keywords:
            if kw in text_lower:
                return ex_type
    return "autre"


# ─────────────────────────────────────────────
# PDF : extraction PyMuPDF — raster + vectoriel + lignes fines
# Clustering v2.7 : composantes connexes (fini l'absorption par bbox)
# ─────────────────────────────────────────────
import fitz  # PyMuPDF, déjà installé (dépendance pdf2docx)


def _connected_components(rects, margin=10.0):
    """v2.7 — Regroupe les primitives par COMPOSANTES CONNEXES : arête entre
    deux primitives d'ORIGINE si leur écart < margin ; la bbox de chaque
    composante est calculée à la fin. Remplace l'union itérative de bbox,
    qui absorbait toute primitive SURVOLÉE par une boîte déjà fusionnée
    sans qu'elle soit réellement proche (cas essais 36-37 : le séparateur
    collé au tissu E entraînait D, C et B dans un composite unique)."""
    rects = [fitz.Rect(r) for r in rects]
    n = len(rects)
    parent = list(range(n))

    def find(a):
        while parent[a] != a:
            parent[a] = parent[parent[a]]
            a = parent[a]
        return a

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[ra] = rb

    for i in range(n):
        gi = fitz.Rect(rects[i].x0 - margin, rects[i].y0 - margin,
                       rects[i].x1 + margin, rects[i].y1 + margin)
        for j in range(i + 1, n):
            if gi.intersects(rects[j]):
                union(i, j)

    comps = {}
    for i in range(n):
        comps.setdefault(find(i), []).append(rects[i])

    out = []
    for members in comps.values():
        bb = fitz.Rect(members[0])
        for m in members[1:]:
            bb |= m
        out.append(bb)
    return out


def _cluster_rasters(rasters, page, margin=16.0):
    """v2.9 — Les rasters se regroupent ENTRE EUX (jamais avec les vecteurs,
    acquis v2.8 conservé) : les scènes composées de nombreux petits rasters
    (pièces de monnaie, billets, tas de tomates, vignettes de situations —
    128 rasters sur la seule série 5P) partaient en figures individuelles,
    saturaient le plafond de miniatures du frontend et sortaient en
    « coller ici » (essai 43). Trois règles :
    1. Composantes connexes, marge 16 pt (écarts intra-scène mesurés : 0-14,6 pt ;
       figures distinctes du corpus : >= 18 pt).
    2. SCÈNES SEULEMENT (v2.9.1, retour essai 44) : un composite exige
       >= 3 membres TOUS petits (min-dim <= 120 pt ~ 42 mm). Tout composant
       contenant un grand raster (scan, photo, clipart) ou réduit à une paire
       est dissous en rasters individuels : les scans à marges blanches qui se
       chevauchent (Pacôme/Momo, essai 39 et 44) et les photos côte à côte
       (essai 42 insectes) ne fusionnent jamais.
    3. SCISSION TABLEAU : un composite multi-rasters est re-scindé à chaque
       bordure horizontale de tableau qui le traverse sans toucher aucun de
       ses membres (collections en lignes de tableau, p5 : écarts
       intra/inter-collections indistinguables géométriquement)."""
    n = len(rasters)
    if n <= 1:
        return list(rasters)

    def _gap(a, b):
        dx = max(0.0, max(a.x0, b.x0) - min(a.x1, b.x1))
        dy = max(0.0, max(a.y0, b.y0) - min(a.y1, b.y1))
        return (dx * dx + dy * dy) ** 0.5

    parent = list(range(n))

    def find(a):
        while parent[a] != a:
            parent[a] = parent[parent[a]]
            a = parent[a]
        return a

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[ra] = rb

    for i in range(n):
        for j in range(i + 1, n):
            if _gap(rasters[i], rasters[j]) < margin:
                union(i, j)

    comps = {}
    for i in range(n):
        comps.setdefault(find(i), []).append(rasters[i])

    # v2.9.1 — FILTRE AU NIVEAU DU COMPOSANT (retour essai 44 : régression sur
    # Pacôme). La garde v2.9 par paires (« deux grands rasters SANS recouvrement
    # ne fusionnent pas ») laissait passer les scans à marges blanches qui se
    # CHEVAUCHENT — exactement le bug de l'essai 39 que la V2.8 avait éliminé :
    # les dessins de Pacôme et Momo ont refusionné en un composite de
    # 161×153 mm avalant le fragment « er, hululer… » de la banque de mots.
    # Nouvelle règle : un composite n'est légitime que pour une SCÈNE —
    # au moins 3 membres, TOUS petits (min-dim <= 120 pt ~ 42 mm : pièces,
    # billets, tas, vignettes). Tout composant contenant un grand raster, ou
    # réduit à une paire, est dissous en rasters individuels (comportement
    # V2.8). Les scans, photos et cliparts ne fusionnent donc plus jamais,
    # avec ou sans recouvrement.
    def _scene(members):
        return len(members) >= 3 and all(
            min(m.width, m.height) <= 120 for m in members)

    # Bordures horizontales candidates pour la scission (triplets de traits
    # LibreOffice regroupés à 3 pt près)
    hlines = []
    try:
        for d in page.get_drawings():
            r = d.get("rect")
            if r is not None and r.height < 4 and r.width > 100:
                hlines.append(fitz.Rect(r))
    except Exception:
        pass

    out = []
    for members in comps.values():
        if len(members) == 1:
            out.append(members[0])
            continue
        if not _scene(members):
            out.extend(members)  # v2.9.1 — pas une scène : chacun reste autonome
            continue
        bb = fitz.Rect(members[0])
        for m in members[1:]:
            bb |= m
        # Coupes : bordures traversant le composite entre les membres
        cuts = []
        for t in hlines:
            cy = (t.y0 + t.y1) / 2
            if not (bb.y0 + 4 < cy < bb.y1 - 4):
                continue
            if t.x0 > bb.x0 + 0.2 * bb.width or t.x1 < bb.x1 - 0.2 * bb.width:
                continue  # ne traverse pas le composite
            if any(m.y0 < cy < m.y1 for m in members):
                continue  # couperait un raster : pas une frontière de lignes
            if not any(abs(c - cy) <= 3 for c in cuts):
                cuts.append(cy)
        if not cuts:
            out.append(bb)
            continue
        cuts.sort()
        bands = {}
        for m in members:
            k = sum(1 for c in cuts if c < (m.y0 + m.y1) / 2)
            bands.setdefault(k, []).append(m)
        for grp in bands.values():
            gb = fitz.Rect(grp[0])
            for m in grp[1:]:
                gb |= m
            out.append(gb)
    return out


def _drop_separator_rows(solids, pw):
    """v2.7 — RANGÉE de barres fines (h < 16 pt) alignées couvrant ensemble
    > 50 % de la largeur de page, sans voisin plein = trait de section coupé
    par du texte (« ---- Exercice 2 ---- »). Le filtre pleine largeur (0.92)
    ne les voyait pas puisque chaque morceau est court. C'est ce trait qui,
    passant à 10,5 pt du tissu E, déclenchait la fusion en chaîne."""
    bars = [r for r in solids if r.height < 16]
    others = [r for r in solids if r.height >= 16]
    drop = set()
    used = set()
    for i, b in enumerate(bars):
        if i in used:
            continue
        row = [i]
        cy = (b.y0 + b.y1) / 2
        for j, o in enumerate(bars):
            if j != i and abs((o.y0 + o.y1) / 2 - cy) < 6:
                row.append(j)
        span = sum(bars[k].width for k in row)
        if span > 0.5 * pw:
            rowrects = [bars[k] for k in row]
            near_full = any(
                fitz.Rect(rr.x0 - 3, rr.y0 - 3, rr.x1 + 3, rr.y1 + 3).intersects(o)
                for rr in rowrects for o in others)
            if not near_full:
                drop.update(row)
                used.update(row)
    return others + [b for i, b in enumerate(bars) if i not in drop]


def _is_text_band(r, words):
    """v2.7 — bande plate (h < 32 pt ≈ 11 mm) contenant du texte : ≥ 2 mots
    de ≥ 3 lettres, ou 1 seul mot couvrant > 30 % de sa largeur = cellule
    d'en-tête de tableau ou titre décoré, PAS une figure. Validé sur la
    source aires : filtre exactement « Figures | Calcul de l'aire »,
    « Exercice N », « L'aire du triangle », « Définition N » — les bandeaux
    gris qui fuyaient dans les fiches. Les figures plates légitimes passent :
    tissus (lettres vectorielles, pas de texte), lignes graduées (chiffres)."""
    if r.height >= 32:
        return False
    n = 0
    cover = 0.0
    for w in words:
        wr = fitz.Rect(w[:4])
        inter = wr & r
        if not inter.is_empty and wr.get_area() > 0 \
                and inter.get_area() > 0.5 * wr.get_area():
            if sum(1 for ch in w[4] if ch.isalpha()) >= 3:
                n += 1
                cover += wr.width
    return n >= 2 or (n >= 1 and cover > 0.3 * r.width)


def _is_underline_of_text(words, t):
    """Ligne horizontale = souligné si du texte immédiatement au-dessus
    couvre ≥ 60 % de sa longueur.
    v2.4 — tolérance élargie : les jambages (descenders) du texte souligné
    descendent SOUS le trait ; mesuré sur les sources réelles, le bas des mots
    est 1,7 à 1,9 pt sous le haut du trait, et l'ancienne condition
    « y1 <= t.y0 + 2 » ne tenait qu'à 0,1 pt près — d'où le soulignement du
    titre « Remplir un tableau… » transmis comme segment par le serveur
    (métriques de police différentes). Fenêtre désormais : bas du mot entre
    10 pt au-dessus et 6 pt au-dessous du trait. Les vrais segments à mesurer
    ont leur texte le plus proche à ≥ 14 pt : aucun faux positif."""
    if t.height > t.width:
        return False
    cover = 0.0
    for w in words:
        x0, y0, x1, y1 = w[:4]
        if y1 <= t.y0 + 6 and t.y0 - y1 < 10:
            ov = min(x1, t.x1) - max(x0, t.x0)
            if ov > 0:
                cover += ov
    return cover >= 0.6 * t.width


# V2.11 (chantier 6.8) — plancher des segments courts étiquetés : 7,0 mm.
# En dessous, rien ne passe, étiquette ou pas ; entre 7 et 14 mm, seulement
# sous les deux conditions du bloc 2c ; à partir de 14,1 mm (40 pt), régime
# inchangé.
SEUIL_COURT_ETIQUETE = 19.8  # 7,0 mm en points PDF

# V2.12 (chantier 4 — tirage essai_10127) — DISTANCE LETTRE → TRAIT.
# La V2.11 exigeait l'étiquette à moins de 12 pt (4 mm) du départ du trait.
# Mesuré le 29.07.2026 sur la page de mesure réellement utilisée (évaluation 3,
# « Partie 2 Géométrie », six segments a) à f) alignés sur des tabulations) :
#
#     d)  85,0 mm   →   7,6 pt      (seul à passer la porte de 12 pt)
#     f) 114,0 mm   →  16,6 pt
#     b) 110,0 mm   →  26,1 pt
#     c)  30,0 mm   →  27,1 pt
#     e)   8,0 mm   →  45,4 pt      ← le seul qui AVAIT BESOIN de la porte
#
# Les quatre longs entraient sans elle (≥ 40 pt). Le court, lui, était rejeté
# pour 33 pt de tabulation. La porte était calibrée sur une hypothèse, pas sur
# une mesure. Elle passe à 60 pt (21 mm), ce qui couvre une tabulation avec de
# la marge.
#
# CE QUI REMPLACE LA DISTANCE COMME GARDE-FOU. Élargir seul rouvrirait la porte
# aux soulignements. Deux conditions la referment :
#   — RIEN ENTRE LES DEUX (ci-dessous) : aucun mot ne s'intercale entre
#     l'étiquette et le départ du trait. Un mot souligné a son texte AU-DESSUS
#     du trait, pas une lettre isolée à sa gauche avec du vide entre les deux ;
#   — la condition de série du bloc 2c, inchangée : il faut un autre trait
#     étiqueté de longueur différente sur la page.
TOL_ETIQUETTE_PT = 60.0

_ETIQUETTE_SERIE = re.compile(r"^([A-Za-z]|\d{1,2})\s*[).]$")
_LETTRE_SEULE = re.compile(r"^([A-Za-z]|\d{1,2})$")


def _etiquette_de_serie(words, t):
    """V2.11 (chantier 6.8) — vrai si un mot-étiquette de série (lettre ou
    chiffre seul suivi d'une parenthèse ou d'un point : a) b) e) 1) 2.) se
    tient JUSTE À GAUCHE du trait. C'est la signature d'un segment d'exercice
    de mesure ; un tiret, une puce ou un soulignement n'est jamais étiqueté.
    À gauche seulement — une étiquette au-dessus ressemble trop à un mot
    souligné ou à un numérateur de fraction : cette porte reste fermée."""
    # V2.12 — une étiquette peut arriver en DEUX mots (« a )» écrit « a » puis
    # « ) », relevé sur la page de mesure réelle). On recompose avant de juger.
    etiquettes = []
    for i, w in enumerate(words):
        txt = w[4].strip()
        if _ETIQUETTE_SERIE.match(txt):
            etiquettes.append(w)
        elif _LETTRE_SEULE.match(txt) and i + 1 < len(words):
            nxt = words[i + 1]
            if nxt[4].strip() in (")", ".") and abs(nxt[1] - w[1]) < 3 \
               and 0 <= nxt[0] - w[2] <= 6:
                etiquettes.append((w[0], w[1], nxt[2], nxt[3], txt + nxt[4].strip()))

    for w in etiquettes:
        wx0, wy0, wx1, wy1 = w[:4]
        # à gauche : le mot finit avant le trait (chevauchement de 2 pt toléré),
        # à moins de TOL_ETIQUETTE_PT de son départ (V2.12 : 60 pt, mesuré)
        if not (wx1 <= t.x0 + 2 and t.x0 - wx1 <= TOL_ETIQUETTE_PT):
            continue
        # V2.12 — RIEN ENTRE LES DEUX. Ce qui sépare une étiquette de série d'un
        # mot souligné, ce n'est pas la distance : c'est le vide. Si un mot
        # s'intercale entre la lettre et le départ du trait, sur la même bande,
        # ce n'est pas une étiquette de segment.
        bande0, bande1 = min(wy0, t.y0) - 3, max(wy1, t.y1) + 3
        intercale = False
        for autre in words:
            if autre is w or not autre[4].strip():
                continue
            ax0, ay0, ax1, ay1 = autre[:4]
            if ax0 >= wx1 - 0.5 and ax1 <= t.x0 + 0.5 \
               and ay1 > bande0 and ay0 < bande1:
                intercale = True
                break
        if intercale:
            continue
        if t.width >= t.height:      # trait horizontal : son axe traverse le mot
            cy = (t.y0 + t.y1) / 2
            if wy0 - 3 <= cy <= wy1 + 3:
                return True
        else:                        # trait vertical : le mot à hauteur du trait
            cy = (wy0 + wy1) / 2
            if t.y0 - 3 <= cy <= t.y1 + 3:
                return True
    return False


# ═══ V2.21 — LES CADRES, LES ZONES VIDES, ET LA FRONTIÈRE DU NATIF (F1, visage B) ═══
# RÈGLE ÉCRITE (arbitrage de Catherine, 10.08) : le natif ne remplace que ce que
# le code sait redessiner sans rien perdre du concret — cases, grilles, bandes
# de nombres. TOUT SUPPORT FIGURATIF (dessin, photo, carte, schéma) reste image,
# toujours. Un wagon n'est pas une case. Mesuré : la carte d'hydrographie sort
# entière (raster autonome, v2.8) ; rien ici ne touche les rasters.
#
# Trois signatures, MESURÉES sur pièces (10.08) avant d'être codées :
#   · cadre d'exercice = trait sans remplissage, ≤ 10 segments, lignes sur le
#     pourtour et petites courbes DE COIN (chaque courbe ≤ 15 % du petit côté) ;
#     une ellipse-membrane a des courbes en quarts entiers : jamais prise.
#   · CADRE ENGLOBANT (la colle des captures trop larges, F6 résolu) : un cadre
#     qui contient ≥ 3 autres primitives ou rasters — il ne lie plus et ne part
#     plus en figure. Panel du 10.08 : 11 → 0 figure à texte de tâche, gâteaux
#     ENTIERS (l'amputation reculait par le même geste), aire intacte 23/23.
#   · CADRE NU VIDE (rien dedans, ≤ 2 mots) : ce n'est pas une image, c'est une
#     ZONE — déclarée (zones_libres), le frontend posera la zone native (B7 :
#     une image ne sert jamais de zone de production ; vu : 177×138 mm de blanc
#     servi en image sur l'évaluation insectes).
# (voir JOURNAL BACKEND v2.22)
def _v21_rect_au_trait(d):
    r = d.get("rect")
    if r is None:
        return False
    f = d.get("fill")
    # un fond BLANC de carte n'empêche pas d'être un cadre (mesuré : les cartes
    # « Qui suis-je » sont des rectangles pointillés à fond blanc)
    if f is not None and not all(v >= 0.95 for v in f):
        return False
    items = d.get("items", [])
    # jusqu'à 24 segments : LibreOffice rend un rectangle arrondi pointillé en
    # 7-8 lignes + 12 courbes de coin (mesuré sur les quadrilatères, 10.08)
    if not items or len(items) > 24:
        return False
    petit = max(1.0, min(r.width, r.height))
    for it in items:
        t = it[0]
        if t == "l":
            for p in (it[1], it[2]):
                if min(abs(p.x - r.x0), abs(p.x - r.x1)) > 3 and \
                   min(abs(p.y - r.y0), abs(p.y - r.y1)) > 3:
                    return False
        elif t == "c":
            bb = fitz.Rect(min(p.x for p in it[1:]), min(p.y for p in it[1:]),
                           max(p.x for p in it[1:]), max(p.y for p in it[1:]))
            # V2.24 — le coin se juge AUSSI en absolu : un rayon de coin est un
            # trait d'auteur (~20-23 pt mesurés), il ne grandit pas avec le cadre.
            # Le seul relatif (15 %) refusait les cadres bas du PDF source exact
            # (22,8 pt > 20,5) et fabriquait la capture collée du tirage 10366_8.
            # (voir JOURNAL BACKEND v2.24)
            if max(bb.width, bb.height) > max(0.15 * petit, 25.0):
                return False
        elif t == "re":
            continue
        else:
            return False
    return True

def _v21_classe_grands_traces(page):
    """Renvoie (indices_englobants, zones_libres) sur les tracés de la page."""
    dessins = page.get_drawings()
    rects = [d.get("rect") for d in dessins]
    rasters = []
    for img in page.get_images(full=True):
        try:
            rasters.extend(fitz.Rect(rr) for rr in page.get_image_rects(img[0]))
        except Exception:
            pass
    englobants, zones = set(), []
    for i, d in enumerate(dessins):
        r = rects[i]
        if r is None or not _v21_rect_au_trait(d):
            continue
        contenus = sum(1 for j, o in enumerate(rects)
                       if j != i and o is not None and r.contains(o))
        contenus += sum(1 for o in rasters if r.contains(o))
        # LE VERROU DES TISSUS (incident attrapé par l'épreuve directe, 10.08) :
        # un contour de tissu est PLEIN là où tous les cadres d'exercice mesurés
        # sont POINTILLÉS — sans tirets, pas d'englobant. L'aire garde ses 23.
        tirets = bool(d.get("dashes")) and str(d.get("dashes")).strip() not in ("", "[] 0")
        if r.get_area() >= 8000 and contenus >= 3 and tirets:
            englobants.add(i)
        # V2.25 — UNE ZONE S'ANNONCE EN POINTILLÉS (miroir du verrou des tissus).
        # Un rectangle PLEIN et vide est une FORME d'exercice (un rectangle est
        # vide par nature) : mesuré le 10.08 soir, la règle sans tirets avalait
        # le rectangle long de l'évaluation (47×9 mm, absent du tirage 10367),
        # 8 formes des fiches 7a et les réceptacles du CP (7 fausses zones).
        # Panel § 2.2 : taille-plancher, voisinage, cadre-parent mesurés et
        # écartés. (voir JOURNAL BACKEND v2.25)
        elif contenus == 0 and r.get_area() >= 2000 and tirets:
            import re as _re
            mots = len(_re.findall(r"[A-Za-zàâçéèêëîïôöûüù]{2,}",
                                   page.get_text(clip=r)))
            if mots <= 2:
                englobants.add(i)          # il ne part pas non plus en figure
                zones.append({"page": None, "w_mm": round(r.width * 25.4 / 72, 1),
                              "h_mm": round(r.height * 25.4 / 72, 1)})
    return englobants, zones

def _v21_est_receptacle(t):
    """« … + … = … » : un réceptacle de réponse, jamais une image (B7)."""
    import re as _re
    t2 = (t or "").strip()
    if not t2:
        return False
    reste = _re.sub(r"\bégal\b", "", t2, flags=_re.I)
    if _re.search(r"[A-Za-zàâçéèêëîïôöûüù]{3,}", reste):
        return False
    return t2.count("…") >= 2 or ("=" in t2 and "…" in t2) \
        or bool(_re.fullmatch(r"[\s.…_+=\-0-9]+", t2))


def _collect_page_regions(page):
    """
    Régions candidates d'une page, en deux familles :
    - solids : images raster + tracés vectoriels « pleins »
    - thins  : lignes fines légitimes (segments à mesurer, lignes graduées)
      V2.11 : + segments courts (7–14 mm) étiquetés en série, bloc 2c
    Filtres : fonds de page (> 85 % de l'aire), séparateurs pleine largeur,
    micro-tracés, soulignés de texte, bordures de tableaux (non isolées),
    familles de lignes identiques (lignes d'écriture, grilles).
    v2.7 : + rangées de séparateurs partiels, + bandes de texte (en-têtes).
    """
    pw, ph = page.rect.width, page.rect.height
    page_area = pw * ph
    v21_exclus, v21_zones = _v21_classe_grands_traces(page)   # V2.21
    solids, thins = [], []
    micros = []          # V2.26 — recueillis pour étendre, jamais pour fonder
    courts = []  # V2.11 — candidats 7-14 mm, promus au bloc 2c ou jetés
    rasters = []  # v2.8 — les rasters ne passent PAS par le clustering

    # 1. Images raster — v2.8 : chaque raster est une figure autonome.
    for img in page.get_images(full=True):
        try:
            for rect in page.get_image_rects(img[0]):
                # v2.33 — une image collée a été posée EXPRÈS : son plancher
                # est celui du résidu invisible (8 pt), jamais celui des traits.
                # (voir JOURNAL BACKEND v2.33)
                if rect.width > 8 and rect.height > 8:
                    if rect.width * rect.height > 0.85 * page_area:
                        continue  # image de fond pleine page
                    rasters.append(fitz.Rect(rect))
        except Exception:
            pass
    # Déduplication des rasters quasi identiques (recouvrement ≥ 80 %)
    dedup = []
    for r in rasters:
        dup = False
        for o in dedup:
            inter = r & o
            if not inter.is_empty and inter.get_area() >= 0.8 * min(r.get_area(), o.get_area()):
                o |= r
                dup = True
                break
        if not dup:
            dedup.append(fitz.Rect(r))
    rasters = dedup
    # v2.9 — composition des scènes multi-rasters (pièces, billets, tas…)
    rasters = _cluster_rasters(rasters, page)

    # 2. Tracés vectoriels
    try:
        for _v21_i, d in enumerate(page.get_drawings()):
            if _v21_i in v21_exclus:      # V2.21 — cadre englobant ou zone vide
                continue
            r = d.get("rect")
            if r is None:
                continue
            if r.width > 0.92 * pw and r.height < 20:
                continue  # ligne de séparation pleine largeur
            if r.width * r.height > 0.85 * page_area:
                continue  # FOND DE PAGE (v2.3) — absorbait tout en v2.2
            if r.width < 10 and r.height < 10:
                # V2.26 — un micro-tracé ne FONDE jamais une figure, mais il
                # ÉTEND le cadre de celle qu'il touche : les cerises (7×7 pt)
                # étaient ignorées et celle du sommet mourait hors cadre —
                # l'élève comptait 5 cerises là où l'énoncé en promet 6
                # (tirage 10368_4, œil de Catherine). Panel § 2.2 : plancher
                # abaissé (inopérant, mesuré) et marge fixe (arrose tout)
                # écartés. (voir JOURNAL BACKEND v2.26)
                micros.append(fitz.Rect(r))
                continue  # micro-tracé
            if r.width < 10 or r.height < 10:
                # Ligne fine : candidate d'office si assez longue pour être
                # une figure (segment, ligne graduée), pas un simple tiret.
                # V2.11 (chantier 6.8) — entre 7 et 14 mm, candidate SOUS
                # CONDITIONS, vérifiées au bloc 2c : un segment court qui
                # porte sa lettre dans un exercice de mesure est une vraie
                # figure (segment e de 8 mm, essai_10117).
                if max(r.width, r.height) >= 40:
                    thins.append(fitz.Rect(r))
                elif max(r.width, r.height) >= SEUIL_COURT_ETIQUETE:
                    courts.append(fitz.Rect(r))
                continue
            solids.append(fitz.Rect(r))
    except Exception:
        pass

    # 2b. v2.7 — séparateurs partiels (rangées de barres) et bandes de texte
    # (en-têtes de tableaux, titres décorés) exclus AVANT clustering : ce sont
    # eux qui provoquaient les fusions en chaîne et les fuites de bandeaux gris.
    solids = _drop_separator_rows(solids, pw)
    words_v27 = page.get_text("words")
    solids = [r for r in solids if not _is_text_band(r, words_v27)]
    # v2.8 — un cluster vectoriel entièrement contenu dans un raster serait un
    # doublon (annotation déjà visible dans le clip du raster) : on le retire
    # en amont pour ne pas générer deux figures du même contenu.
    solids = [r for r in solids
              if not any(ra.contains(r) for ra in rasters)]

    # 2c. V2.11 (chantier 6.8) — promotion des segments courts étiquetés.
    # Un trait de 7 à 14 mm n'entre qu'à deux conditions cumulées :
    #   1. une étiquette de série juste à sa gauche (a), e), 1), 2. …) ;
    #   2. un AUTRE trait étiqueté de longueur différente (> 5 %) sur la
    #      page — les segments d'un exercice de mesure ont des longueurs
    #      toutes différentes, c'est le principe de l'exercice ; les lignes
    #      de réponse d'une liste numérotée font toutes la même : dehors.
    # Les promus subissent ensuite les mêmes filtres anti-bruit que les
    # longs (isolement, souligné, familles) puis le filtre v2.5 des pages
    # de mesure — la porte ne s'ouvre que d'un cran, pas en grand.
    if courts:
        candidats = [t for t in courts
                     if _etiquette_de_serie(words_v27, t)]
        if candidats:
            longueurs = [max(t.width, t.height)
                         for t in list(thins) + candidats
                         if _etiquette_de_serie(words_v27, t)]
            for t in candidats:
                L = max(t.width, t.height)
                if any(abs(L - L2) > 0.05 * max(L, L2)
                       for L2 in longueurs):
                    thins.append(t)

    # 3. Anti-bruit lignes fines
    if thins:
        words = words_v27

        def _isolated(t):
            probe = fitz.Rect(t.x0 - 3, t.y0 - 3, t.x1 + 3, t.y1 + 3)
            if any(probe.intersects(o) for o in solids):
                return False
            return not any((o is not t) and probe.intersects(o) for o in thins)

        thins = [t for t in thins if _isolated(t)]
        thins = [t for t in thins if not _is_underline_of_text(words, t)]

        kept = []
        for t in thins:
            horiz = t.width >= t.height
            L = t.width if horiz else t.height
            same = sum(1 for o in thins
                       if (o.width >= o.height) == horiz
                       and abs((o.width if horiz else o.height) - L) <= 0.05 * L)
            if same < 3:
                kept.append(t)
        thins = kept

    return solids, thins, rasters, v21_zones, micros


# ─────────────────────────────────────────────
# V2.10 — grilles et formes composites
# Import tolérant : si le module n'a pas été déployé à côté de app.py, le
# pipeline garde exactement son comportement 2.9.3. Pas de panne, pas de
# demi-mesure silencieuse : l'état est visible sur /health.
# ─────────────────────────────────────────────
try:
    import grilles as _grilles
    GRILLES_ACTIVES = True
except Exception:  # module absent ou illisible
    _grilles = None
    GRILLES_ACTIVES = False

try:
    import formes as _formes
    FORMES_ACTIVES = True
except Exception:
    _formes = None
    FORMES_ACTIVES = False


# ══ v2.14 — LE GRAS DE LA SOURCE SURVIT À LA LECTURE DU PDF ═══════════════════════════
# Chantier ouvert le 30.07.2026. La règle de mise en évidence (frontend v10.151) dit que
# le gras de la source est reproduit par défaut ; mais quand la source arrive en PDF, le
# serveur lisait `page.get_text("text")`, qui rend une chaîne PLATE : le gras était perdu
# avant même que le modèle le voie. Aucune règle de conservation ne pouvait donc être
# tenue sur ce chemin — la perte était silencieuse, le pire cas de la doctrine.
#
# Ce que fait la fonction : elle relit la page span par span et entoure les passages en
# gras de deux marqueurs — ⟦gras⟧ … ⟦/gras⟧ — que le frontend sait reconvertir en <strong>
# (filet v10.158 b). Le choix de marqueurs à crochets blancs plutôt que d'étoiles ou de
# balises est délibéré : ils n'existent dans aucune source scolaire, ne se confondent avec
# aucune syntaxe, et s'ils survivent par accident jusqu'à la feuille élève le filet du
# frontend les rattrape et le signale.
#
# BORNES CONNUES, écrites ici plutôt qu'oubliées :
#  - le chemin « grilles » (texte_avec_grilles) construit son texte lui-même et n'est pas
#    couvert : sur une page à grilles, le gras reste perdu. À traiter en v2.15.
#  - PyMuPDF signale le gras par le bit 4 des drapeaux de span (valeur 16) ; certaines
#    polices déclarent le gras dans leur NOM seulement (« …-Bold »), d'où le second test.
_GRAS_OUV = "\u27e6gras\u27e7"
_GRAS_FER = "\u27e6/gras\u27e7"


def _span_est_gras(span):
    if (span.get("flags", 0) & 16):
        return True
    nom = (span.get("font", "") or "").lower()
    return ("bold" in nom) or ("black" in nom) or ("heavy" in nom)


def _texte_gras(page):
    """Texte de la page, gras marqué. Retombe sur get_text('text') au moindre doute."""
    try:
        d = page.get_text("dict")
    except Exception:
        return page.get_text("text")
    blocs = []
    for b in d.get("blocks", []):
        if b.get("type", 0) != 0:
            continue
        lignes = []
        for l in b.get("lines", []):
            morceaux, ouvert = [], False
            for s in l.get("spans", []):
                txt = s.get("text", "")
                if not txt:
                    continue
                g = _span_est_gras(s)
                if g and not ouvert:
                    morceaux.append(_GRAS_OUV); ouvert = True
                elif not g and ouvert:
                    morceaux.append(_GRAS_FER); ouvert = False
                morceaux.append(txt)
            if ouvert:
                morceaux.append(_GRAS_FER)
            ligne = "".join(morceaux)
            # un marqueur qui n'entoure que des espaces n'apporte rien et gêne la relecture
            ligne = ligne.replace(_GRAS_OUV + " " + _GRAS_FER, " ")
            ligne = ligne.replace(_GRAS_OUV + _GRAS_FER, "")
            lignes.append(ligne)
        if lignes:
            blocs.append("\n".join(lignes))
    if not blocs:
        return page.get_text("text")
    return "\n".join(blocs) + "\n"


# ═══ V2.20 — LA PORTE DES CORRIGÉS (chantier F1, visage A) ═════════════════════
# Les banques de ressources livrent l'évaluation ET son corrigé dans le même
# document. Jusqu'ici, TOUT partait au modèle : figures en double, texte des
# réponses — la règle de prompt v10.258 était la seule défense, et le tirage
# 10366_5 l'a démentie (4 figures du corrigé servies). Le test v10.258 écrivait
# lui-même le remède attendu : « un filet déterministe... côté serveur ».
#
# Le critère est celui déjà arbitré (v10.258 : deux endroits portant les MÊMES
# items, l'un vide et l'autre rempli), rendu déterministe par DEUX preuves
# positives, mesurées au panel du 10.08 sur 5 sources (3/3 pages trouvées,
# 0 fausse alerte — vérité écrite d'avance) :
#   (1) JUMEAU STRUCTUREL : deux pages se couvrent à ≥ 70 % avec des volumes
#       comparables, et l'une a PERDU ses pointillés de réponse — c'est elle
#       le corrigé (« l'un vide, l'autre rempli », mot pour mot) ;
#   (2) FIGURES JUMELLES : une même image embarquée reparaît sur une page
#       ultérieure — la page tardive est le corrigé (convention des banques).
# SANS preuve positive, RIEN n'est écarté : le doute ne détruit jamais une
# évaluation (K.5). L'exemple résolu n'a pas de jumeau vide : il passe.
# L'écart est DÉCLARÉ (champ corrige_pages) — une erreur signalée est
# tolérable, une silencieuse disqualifie. (voir JOURNAL BACKEND v2.20)
_RX_MOT = None
def _corrige_tokens(t):
    import re as _re
    return [w for w in _re.findall(r"[a-zàâçéèêëîïôöûüù]{3,}|\d+", t.lower())]
def _corrige_pointilles(t):
    import re as _re
    return len(_re.findall(r"[.…_]{4,}", t))
def pages_de_corrige(doc):
    """Renvoie l'ensemble (1-indexé) des pages reconnues comme corrigé."""
    txt = [p.get_text() for p in doc]
    toks = [_corrige_tokens(t) for t in txt]
    flag = set()
    # (1) jumeau structurel + pointillés perdus
    for j in range(len(toks)):
        for k in range(j + 1, len(toks)):
            sj, sk = set(toks[j]), set(toks[k])
            if not sj or not sk:
                continue
            couv = max(len(sj & sk) / len(sj), len(sj & sk) / len(sk))
            ratio = len(toks[k]) / max(1, len(toks[j]))
            if couv >= 0.70 and 0.5 <= ratio <= 2.0:
                pj, pk = _corrige_pointilles(txt[j]), _corrige_pointilles(txt[k])
                if pj == pk:
                    continue          # pas de « vide contre rempli » : on ne touche pas (K.5)
                flag.add((k + 1) if pk < pj else (j + 1))
    # (2) figures jumelles : même image embarquée sur deux pages — MAIS une
    # ressource décorative (trame, bandeau) se réutilise aussi entre pages.
    # La preuve n'est donc complète que si les DEUX pages sont AUSSI jumelles
    # par leur texte (couverture >= 0.70) : c'est le critère v10.258 entier,
    # mêmes items des deux côtés. Attrapé par la batterie AVANT livraison :
    # sans ce verrou, une texture partagée faisait écarter une page saine
    # (aire p.3, numération p.4). (voir JOURNAL BACKEND v2.20)
    vues = {}
    for pno, page in enumerate(doc):
        for info in page.get_images(full=True):
            vues.setdefault(info[0], set()).add(pno + 1)
    def _jumelles_texte(j, k):
        sj, sk = set(toks[j - 1]), set(toks[k - 1])
        if not sj or not sk:
            return False
        return max(len(sj & sk) / len(sj), len(sj & sk) / len(sk)) >= 0.70
    for pages in vues.values():
        if len(pages) > 1:
            k = max(pages)
            if any(_jumelles_texte(j, k) for j in pages if j < k):
                flag.add(k)
    return flag


def parse_pdf(content: bytes, filename: str):
    """
    Extrait d'un PDF, dans l'ordre de lecture :
    - les images raster (photos, dessins importés)
    - les figures vectorielles (tracés regroupés puis rasterisés en PNG 2x)
    - les lignes fines légitimes (segments à mesurer, lignes graduées) — v2.3
    - le texte complet
    Chaque figure part avec sa taille physique (w_mm/h_mm du clip rasterisé)
    pour l'impression à l'échelle réelle (class="img-echelle").
    v2.7 : clustering par composantes connexes, marge 10 pt.
    """
    doc = fitz.open(stream=content, filetype="pdf")
    images = []
    full_text = []
    grilles_inventaire = []
    formes_inventaire = []
    idx = 0
    zones_libres = []          # V2.21 — cadres nus vides, déclarés en zones
    n_receptacles = 0          # V2.21 — « … + … = … » écartés de la voie image
    n_c11 = 0                  # v2.34 — cadres sans aucun dessin (C11)
    corrige = pages_de_corrige(doc)   # V2.20 — figures ET texte de ces pages restent dehors

    for pno, page in enumerate(doc):
        if (pno + 1) in corrige:      # V2.20 — page de corrigé : rien n'en part
            continue
        pw, ph = page.rect.width, page.rect.height
        page_area = pw * ph

        solids, thins, rasters, v21_zones_page, micros_page = _collect_page_regions(page)
        for _z in v21_zones_page:
            _z["page"] = pno + 1
            zones_libres.append(_z)

        # V2.10 — grilles de la page : forme, contenu, comptage des cases.
        # `zones_grilles` liste les tableaux de texte, qui ne doivent plus
        # repartir en figure découpée : leur contenu est déjà transmis en
        # clair, et l'image du même tableau faisait doublon.
        # V2.10 — formes composites : une figure faite de plusieurs
        # rectangles ne doit plus partir comme un bloc plein. Sa
        # décomposition est décrite, son aire réelle mesurée.
        formes_page = []
        if FORMES_ACTIVES:
            try:
                formes_page = _formes.analyser_formes(page)
            except Exception:
                formes_page = []
        blocs_formes = [(f["rect"].y0, f["rect"].x0, _formes.decrire(f))
                        for f in formes_page]

        grilles_page, zones_grilles = [], []
        if GRILLES_ACTIVES:
            try:
                analyse = _grilles.analyser_page(page)
                grilles_page = analyse["grilles"]
                zones_grilles = analyse["zones_grilles"]
                texte_page = _grilles.texte_avec_grilles(
                    page, grilles_page, blocs_formes)
            except Exception:
                texte_page = _texte_gras(page)      # v2.14
        else:
            texte_page = _texte_gras(page)          # v2.14
            if blocs_formes:
                texte_page += "\n" + "\n".join(b[2] for b in sorted(blocs_formes))

        # v2.5 — les lignes fines ne partent que si la page parle de mesure ou
        # de tracé : sur une fiche de français, les lignes de réponse aux
        # longueurs variées traversent le filtre « familles » (28 fausses
        # lignes sur la série de conjugaison 5P) et noieraient le modèle sous
        # des « segments » sans objet. Mots-clés calibrés sur les 9 sources :
        # les pages segments/périmètres/aires/quadrillages en contiennent
        # toutes au moins un, aucune page de français n'en contient.
        if thins:
            ptxt = page.get_text("text").lower()
            if not (any(k in ptxt for k in (
                    "mesur", "segment", "périmètre", "perimetre",
                    "centimètre", "centimetre", "millimètre", "millimetre",
                    "gradu", "quadrill", "trace", "construc",
                    "géométr", "geometr"))
                    or re.search(r"\baires?\b", ptxt)):  # « aire » en mot entier — pas « faire »
                thins = []

        # v2.4 — les lignes fines ne sont JAMAIS clusterisées : une ligne qui a
        # survécu aux filtres anti-bruit est un segment autonome à mesurer.
        regions = solids

        # Clustering v2.7 : composantes connexes sur les primitives d'origine,
        # marge 10 pt (l'écart réel entre deux figures distinctes du corpus
        # est ≥ 12 pt ; les sous-éléments d'une même figure restent < 10 pt).
        clusters = _connected_components(regions, margin=10.0)

        # Re-split (v2.3, adapté v2.7) : un cluster quasi pleine page n'est pas
        # jeté, il est re-clusterisé plus finement ; en dernier recours, ses
        # primitives pleines sont gardées individuellement.
        final = []
        for c in clusters:
            c = c & page.rect
            if c.width * c.height <= 0.85 * page_area:
                final.append(c)
                continue
            subs = _connected_components(
                [r for r in regions if r.intersects(c)], margin=5.0)
            for s in subs:
                s = s & page.rect
                if s.width * s.height <= 0.85 * page_area:
                    final.append(s)
                else:
                    final.extend(r for r in solids
                                 if r.intersects(s) and r.width >= 24 and r.height >= 24)

        # v2.4 — chaque ligne fine part individuellement
        final.extend(thins)
        # v2.8/v2.9 — les rasters ne se mélangent jamais aux vecteurs ; depuis
        # v2.9 les scènes multi-rasters arrivent déjà composées (et re-scindées
        # aux bordures de tableau) par _cluster_rasters.
        final.extend(rasters)

        # Filtres finaux : figures pleines OU lignes fines assez longues
        # V2.11 — les segments courts étiquetés, promus au bloc 2c, passent
        # au même titre que les longs : la porte ne s'ouvre que pour EUX,
        # le plancher des 14 mm reste la règle pour tout le reste.
        cles_courts = set()
        for t in thins:
            if max(t.width, t.height) < 40:
                tt = t & page.rect
                cles_courts.add((round(tt.x0), round(tt.y0),
                                 round(tt.x1), round(tt.y1)))
        keep, seen = [], set()
        for r in final:
            r = r & page.rect
            key = (round(r.x0), round(r.y0), round(r.x1), round(r.y1))
            if key in seen:
                continue
            seen.add(key)
            # v2.33 — même doctrine en aval : le plancher des 24 pt a été
            # calibré pour écarter les MIETTES DE TRAITS ; appliqué aux images
            # il jetait les pions (23,6 × 35,1 pt, à 0,4 pt du seuil).
            _img33 = any(not (ra & r).is_empty
                         and (ra & r).get_area() >= 0.9 * min(ra.get_area(), r.get_area())
                         for ra in rasters)
            full = (r.width >= 8 and r.height >= 8) if _img33 else (
                r.width >= 24 and r.height >= 24)
            slim = min(r.width, r.height) < 24 and (
                max(r.width, r.height) >= 40 or key in cles_courts)
            if not (full or slim):
                continue
            if (r.width * r.height) > 0.85 * page_area:
                continue
            # V2.10 — un tableau de texte n'est pas une figure : son contenu
            # part déjà en clair, l'image ferait doublon. MAIS beaucoup de
            # sources se servent d'un tableau comme d'une mise en page et y
            # posent leurs dessins : une figure logée dans une cellule reste
            # une figure, et perdre un dessin est une faute grave.
            dans_tableau = any(not (r & z).is_empty
                               and (r & z).get_area() >= 0.6 * r.get_area()
                               for z in zones_grilles)
            if dans_tableau and not any(
                    not (r & ra).is_empty
                    and (r & ra).get_area() >= 0.5 * min(r.get_area(),
                                                         ra.get_area())
                    for ra in rasters):
                continue
            keep.append(r)

        # V2.10 — le quadrillage part d'un seul tenant. Jusqu'ici ses traits
        # réguliers étaient éliminés comme du bruit et seules les zones
        # coloriées survivaient : le modèle recevait des taches sans repère,
        # et l'élève une figure sans quadrillage à compter. On ajoute donc le
        # cadre complet, et on retire les morceaux qu'il contient — ils y sont
        # déjà, en place.
        cadres_grille = set()      # v2.32 — cadres nes du module des grilles
        for g in grilles_page:
            if g["nature"] != "quadrillage":
                continue
            cadre = fitz.Rect(g["rect"]) & page.rect
            if cadre.is_empty or cadre.width < 24 or cadre.height < 24:
                continue
            keep = [r for r in keep
                    if not (cadre.contains(r)
                            or (not (r & cadre).is_empty
                                and (r & cadre).get_area() >= 0.6 * r.get_area()))]
            # v2.33 — LE CADRE SUIT LA SCÈNE, PAS LA TABLE. Condition : un objet
            # (dessiné ou collé) touche la bande par le dessus ou le dessous, dans
            # son emprise horizontale. Effet : le cadre l'emporte. Le jeu de 12 pt
            # est mesuré : à 30 pt, deux bandes voisines fusionnent (témoin T5).
            _haut33 = fitz.Rect(cadre.x0 - 3, cadre.y0 - 12, cadre.x1 + 3, cadre.y0)
            _bas33 = fitz.Rect(cadre.x0 - 3, cadre.y1, cadre.x1 + 3, cadre.y1 + 12)
            for _z33 in (_haut33, _bas33):
                for _d33 in page.get_drawings():
                    _r33 = _d33.get("rect")
                    if not _r33:
                        continue
                    _o33 = fitz.Rect(_r33)
                    if _o33.width < 4 or _o33.width > 250 or _o33.height < 3:
                        continue
                    if (_o33.intersects(_z33) and _o33.x0 >= cadre.x0 - 5
                            and _o33.x1 <= cadre.x1 + 5):
                        cadre = cadre | _o33
                for _ra33 in rasters:
                    if (_ra33.intersects(_z33) and not (_ra33 & cadre).is_empty
                            and _ra33.x0 >= cadre.x0 - 5 and _ra33.x1 <= cadre.x1 + 5):
                        cadre = cadre | _ra33
            cadres_grille.add((round(cadre.x0, 1), round(cadre.y0, 1),
                               round(cadre.x1, 1), round(cadre.y1, 1)))
            keep.append(cadre)

        # ══ v2.17 — UNE FORME À MOITIÉ DANS LA ZONE Y EST PRISE ENTIÈRE ═════
        # Défaut mesuré le 02.08.2026 sur SIX documents sources (151 zones) :
        # 20 formes vectorielles sont AMPUTÉES par le cadre de découpe — un rond
        # de la démonstration CP perd 47 % de sa surface, une autre forme 81 %.
        # Un rond incomplet, sur une fiche qui enseigne « j'en dessine un de
        # plus », n'est pas une imperfection : il empêche de faire l'exercice.
        # Cause : le cadre est construit sur la PHOTO ; les traits vectoriels qui
        # font partie du même dessin ne sont pas absorbés.
        #
        # CE QUI A ÉTÉ MESURÉ, ET POURQUOI CES TROIS NOMBRES (panel_marges_2.py,
        # dix-huit candidats de familles différentes) :
        #  · MAJORITÉ 40 % — une forme dont 40 % est déjà dans le cadre en fait
        #    partie ; une forme que le cadre effleure appartient au voisinage.
        #    Le plateau mesuré va de 25 % à 50 % (20 formes réparées, aucun coût) ;
        #    à 20 % une voisine est avalée, à 60 % trois formes restent coupées.
        #    40 % est le milieu du plateau, pas son bord.
        #  · TAILLE ≤ 75 pt — plateau mesuré de 60 à 90 pt ; à 100 pt, six
        #    voisines avalées et +5 % de surface ; à 45 pt, cinq formes ratées.
        #    Sans limite du tout : quatre mots avalés, six voisines, +7 %.
        #  · JAMAIS UNE FORME QUI PORTE DU TEXTE encore dehors : une étiquette
        #    « 3 + … » ou le badge « CP » est un rectangle fermé comme un rond.
        #    L'avaler imprimerait son texte DEUX fois — une fois en clair, une
        #    fois dans l'image. Cette garde retire 2 mots avalés sur 2.
        # Résultat : 20 formes réparées sur 20, zéro mot avalé, zéro voisine
        # avalée, +0 % de surface moyenne.
        #
        # CE QUE CE REMÈDE NE FAIT PAS, et c'est déclaré : trois formes des six
        # documents (une flèche de la fiche CP 16b, un polygone de l'évaluation
        # 5P) sont dedans à 26–28 % et restent donc fragmentées. Les prendre
        # exigerait un seuil de 25 % qui est sur le fil du rasoir : à 20 %, une
        # voisine est avalée. Choix assumé : le plateau plutôt que l'optimum.
        _MAX_FORME = 75.0        # côté maximal d'une forme absorbable, en points
        _PART_MINI = 0.40        # part déjà dans le cadre pour qu'elle en fasse partie
        try:
            _mots_page = [(fitz.Rect(w[:4]), w[4]) for w in page.get_text("words")]
            _formes = [fitz.Rect(_d["rect"]) for _d in page.get_drawings()]
            _formes = [_r for _r in _formes
                       if 1 < _r.width <= _MAX_FORME and 1 < _r.height <= _MAX_FORME]
            _elargies = []
            for _z in keep:
                _out = fitz.Rect(_z)
                for _r in _formes:
                    _i = _r & _z
                    if _i.is_empty or _r.get_area() <= 0:
                        continue
                    if _i.get_area() / _r.get_area() < _PART_MINI:
                        continue
                    if any(_r.contains(_w) and not fitz.Rect(_z).contains(_w)
                           for _w, _t in _mots_page):
                        continue
                    _out |= _r
                _elargies.append(_out & page.rect)
            keep = _elargies
        except Exception:
            pass          # une page illisible ne doit pas faire tomber l'extraction

        # ══ v2.16 — LES RANGÉES DE LECTURE S'ANCRENT SUR LES ZONES, PLUS SUR
        # UNE GRILLE FIXE ═══════════════════════════════════════════════════
        # Défaut mesuré le 02.08.2026, fiche i-profs CP 8a (tirage 10275) : les
        # deux panneaux de la démonstration « ajouter / enlever » sont côte à
        # côte, à 2,9 points de hauteur près — y0 = 249,8 (droite) et 252,7
        # (gauche). Le découpage en tranches FIXES de 24 points tombe entre les
        # deux : 249,8/24 = 10,41 → 10 ; 252,7/24 = 10,53 → 11. Le panneau de
        # DROITE part donc avant celui de GAUCHE, et sur la feuille
        # l'illustration « j'en barre un » se retrouve sous « Pour ajouter 1 ».
        # Les deux images du tirage sont identiques OCTET POUR OCTET à celles que
        # ce tri livre : l'inversion naît ici, pas dans le modèle.
        #
        # LA CAUSE EST LA GRILLE FIXE, PAS SA TAILLE. Deux zones d'une même
        # rangée peuvent toujours tomber de part et d'autre d'une frontière,
        # quelle que soit la largeur choisie. Agrandir la tranche déplacerait la
        # frontière sans la supprimer.
        #
        # LE REMÈDE : la rangée s'ancre sur la PREMIÈRE zone rencontrée, pas sur
        # un multiple de 24. Les zones sont triées par hauteur ; tant qu'une zone
        # commence à moins de 24 points du début de la rangée en cours, elle est
        # de cette rangée. Chaque rangée est ensuite lue de gauche à droite.
        # La tolérance de 24 points est CELLE D'AVANT, inchangée : seul son point
        # d'ancrage change. Tout document dont les zones ne tombaient pas sur une
        # frontière garde donc exactement l'ordre qu'il avait.
        _TOL_RANGEE = 24
        _tri = sorted(keep, key=lambda r: (r.y0, r.x0))
        _rangees, _cour, _debut = [], [], None
        for _r in _tri:
            if _debut is None or _r.y0 - _debut <= _TOL_RANGEE:
                if _debut is None:
                    _debut = _r.y0
                _cour.append(_r)
            else:
                _rangees.append(_cour)
                _cour, _debut = [_r], _r.y0
        if _cour:
            _rangees.append(_cour)
        keep = []
        for _rg in _rangees:
            _rg.sort(key=lambda r: r.x0)
            keep.extend(_rg)

        # Rasterisation 2x de chaque zone
        page_words = page.get_text("words")
        for r in keep:
            try:
                # Lignes fines : dilater le clip pour que le trait soit
                # visible ; w_mm/h_mm décrivent le CLIP rasterisé, donc
                # l'échelle réelle reste exacte à l'impression.
                clip = fitz.Rect(r)
                is_thin = min(r.width, r.height) < 24
                if r.height < 24 and r.width >= r.height:
                    clip = fitz.Rect(r.x0 - 3, r.y0 - 6, r.x1 + 3, r.y1 + 6)
                elif r.width < 24:
                    clip = fitz.Rect(r.x0 - 6, r.y0 - 3, r.x1 + 6, r.y1 + 3)
                clip = clip & page.rect
                # v2.4 — GARDE FINALE : un clip de ligne fine qui contient du
                # texte n'est PAS un segment à mesurer (soulignement de titre,
                # ligne collée à un libellé) → on ne le transmet pas, quel que
                # soit le verdict des filtres géométriques en amont. Seules les
                # LETTRES comptent : les chiffres d'une ligne graduée restent
                # légitimes.
                if is_thin:
                    letters = 0
                    for w in page_words:
                        wr = fitz.Rect(w[:4])
                        inter = wr & clip
                        if not inter.is_empty and wr.get_area() > 0 \
                                and inter.get_area() > 0.3 * wr.get_area():
                            letters += sum(1 for ch in w[4] if ch.isalpha())
                    if letters >= 3:
                        continue
                # V2.22 — LE DÉCOR SE SIGNALE, NE SE SUPPRIME PAS (F1, visage D).
                # Signature mesurée à l'œil sur 7 sources (6/6, 0 fausse
                # alerte) : petite figure ancrée aux marges — badge rond de
                # niveau, cartouche d'éditeur, pastille « Fiche N », ruban
                # latéral. Une barre fine de titre (32×4) n'est PAS prise
                # (hauteur >= 8 mm exigée). (voir JOURNAL BACKEND v2.22)
                # V2.26 — les micros qui TOUCHENT ce cadre (≤ 3 pt) l'étendent.
                for _mi26 in micros_page:
                    if fitz.Rect(clip.x0 - 3, clip.y0 - 3,
                                 clip.x1 + 3, clip.y1 + 3).intersects(_mi26):
                        clip.include_rect(_mi26)
                _w_mm_ = clip.width * 25.4 / 72
                _h_mm_ = clip.height * 25.4 / 72
                _marge_ = (clip.y0 < 0.08 * ph) or (clip.x0 < 0.06 * pw) or (clip.x1 > 0.94 * pw)
                _decor_ = _marge_ and ((_w_mm_ <= 45 and 8 <= _h_mm_ <= 25)
                                       or (_w_mm_ <= 20 and _h_mm_ <= 80 and clip.x1 > 0.94 * pw))
                # V2.21 — un réceptacle de réponse ne part pas en image (B7) ;
                # son contenu est déjà dans le texte, le frontend pose la case.
                _t_clip = page.get_text(clip=clip)
                # v2.32 — UNE GRILLE DECLAREE N'EST PAS UN RECEPTACLE. Condition :
                # le cadre vient du module des grilles. Effet : la porte des
                # receptacles ne s'applique pas — une bande de 12 chiffres (piste
                # de jeu, frise) n'est pas un « … + … = … ». Source unique : la
                # nature d'une grille est jugee par le module, pas deux fois.
                # (voir JOURNAL BACKEND v2.32)
                _ne_grille = (round(r.x0, 1), round(r.y0, 1),
                              round(r.x1, 1), round(r.y1, 1)) in cadres_grille
                # v2.33 — symétrique de la 2.32 : une case à remplir est une zone
                # VIDE, jamais une image. La porte v2.21 jugeait le seul texte de
                # la découpe et tuait toute scène tenant un nombre (les grenouilles
                # sur « 47 », le garçon à l'ardoise « 9 + 4 »).
                _imgr33 = any(not (ra & clip).is_empty
                              and (ra & clip).get_area() >= 0.5 * min(ra.get_area(), clip.get_area())
                              for ra in rasters)
                if not _ne_grille and not _imgr33 and _v21_est_receptacle(_t_clip):
                    n_receptacles += 1
                    continue
                import re as _re21
                _mots_clip = len(_re21.findall(r"[A-Za-zàâçéèêëîïôöûüù]{2,}", _t_clip))
                # ══ v2.34 — UN CADRE SANS DESSIN N'EST PAS UNE FIGURE (C11) ══
                # Condition : le cadre ne porte aucun objet dessiné.
                # Effet : il n'est pas transmis.
                # (voir JOURNAL BACKEND v2.34)
                # UNE GRILLE DÉCLARÉE N'EST JAMAIS UNE FAUSSE FIGURE : sa
                # nature est jugée par le module des grilles, pas deux fois
                # (même principe que la v2.32). Défaut attrapé le 12.08 par le
                # témoin T3 de la batterie 2.33 : une bande de nombres tracée
                # au TRAIT n'a aucun objet dessiné au sens du filtre — chaque
                # trait fait moins de 20 pt² — et disparaissait. Sur le CP-50,
                # une bande reconstruite perd les couleurs de ses cases.
                if not _ne_grille and not _v234_porte_un_dessin(page, clip, rasters):
                    n_c11 += 1
                    continue
                pix = page.get_pixmap(clip=clip, matrix=fitz.Matrix(2, 2))
                if pix.width < 8 or pix.height < 8:
                    continue
                _png30 = pix.tobytes("png")
                b64 = base64.b64encode(_png30).decode()
                images.append({
                    # v2.30 — chaque figure porte son cadre en PLEINES décimales
                    # et le sceau de son contenu (sha-256 du PNG, 12 hex).
                    # Condition : la figure est retenue. Effet : dire QUOI est
                    # découpé, pas seulement combien. (voir JOURNAL BACKEND v2.30)
                    "cadre": [clip.x0, clip.y0, clip.x1, clip.y1],
                    "sceau": _hl.sha256(_png30).hexdigest()[:12],
                    # V2.21 — une figure qui porte des mots se SIGNALE, elle ne se
                    # supprime jamais : une scène concrète (bulle, étiquette de
                    # wagon) vaut plus que la règle (arbitrage du 10.08, K.5).
                    **({"mots_taches": _mots_clip} if _mots_clip >= 6 else {}),
                    **({"decor_probable": True} if _decor_ else {}),
                    "index": idx,
                    "page": pno + 1,
                    "data": f"data:image/png;base64,{b64}",
                    "w": pix.width, "h": pix.height,
                    # Taille PHYSIQUE de la zone rasterisée dans le document
                    # source (points PDF → mm : 1 pt = 25.4/72 mm). Permet au
                    # frontend d'imprimer la figure à taille réelle
                    # (class="img-echelle") pour la mesure à la règle.
                    "w_mm": round(clip.width * 25.4 / 72, 1),
                    "h_mm": round(clip.height * 25.4 / 72, 1),
                })
                idx += 1
            except Exception:
                pass

        for g in grilles_page:
            grilles_inventaire.append({
                "page": pno + 1,
                "nature": g["nature"],
                "lignes": g["lignes"],
                "colonnes": g["colonnes"],
                "case_l_mm": g["case_l_mm"],
                "case_h_mm": g["case_h_mm"],
                "zones": [{"etiquette": z["etiquette"], "cases": z["cases"],
                           "composite": z["composite"]} for z in g["zones"]],
                "comptage_refuse": g["comptage_refuse"],
            })

        for f in formes_page:
            formes_inventaire.append({
                "page": pno + 1,
                "etiquette": f["etiquette"],
                "membres": f["membres"],
                "encombrement_mm2": f["encombrement_mm2"],
                "aire_mm2": f["aire_mm2"],
                "aire_refusee": f["aire_refusee"],
            })

        full_text.append(texte_page)

    doc.close()
    return {
        "filename": filename,
        "pdf_mode": True,
        # V2.20 — pages écartées comme corrigé, DÉCLARÉES (jamais en silence)
        "corrige_pages": sorted(corrige),
        # V2.21 — zones de production déclarées + réceptacles écartés, comptés
        "zones_libres": zones_libres,
        "receptacles_ecartes": n_receptacles,
        "cadres_sans_dessin": n_c11,          # v2.34 — C11

        "num_exercises": 0,
        "num_images": len(images),
        "images": images,
        "text": "\n".join(full_text),
        "exercises": [],
        # V2.10 — inventaire des grilles rencontrées, pour le journal et pour
        # le vérificateur : nature, dimensions, cases comptées ou refus motivé.
        "grilles": grilles_inventaire,
        "grilles_actives": GRILLES_ACTIVES,
        # Idem pour les formes composites : membres, encombrement, aire
        # réelle ou refus motivé.
        "formes": formes_inventaire,
        "formes_actives": FORMES_ACTIVES,
    }

# ─────────────────────────────────────────────
# DOCX : rasterisation générique des images non-web
# (EMF / WMF / TIFF / vectoriels) → PNG via LibreOffice
# ─────────────────────────────────────────────
_WEB_SAFE = ("png", "jpeg", "jpg", "gif", "webp", "bmp")
# Fragments de content-type non affichables par un navigateur → extension LibreOffice
_NONWEB_EXT = {
    "x-emf": "emf", "emf": "emf",
    "x-wmf": "wmf", "wmf": "wmf",
    "tiff": "tiff", "tif": "tiff",
}


def _is_web_safe(ct: str) -> bool:
    ct = (ct or "").lower()
    return any(w in ct for w in _WEB_SAFE)


def _nonweb_ext(ct: str):
    ct = (ct or "").lower()
    for frag, ext in _NONWEB_EXT.items():
        if frag in ct:
            return ext
    return None


def _autocrop_png(png_bytes: bytes, pad: int = 6) -> bytes:
    """Rogne les marges blanches d'un PNG (LibreOffice exporte une page entière)."""
    try:
        from PIL import Image, ImageChops
        im = Image.open(io.BytesIO(png_bytes)).convert("RGB")
        bg = Image.new("RGB", im.size, (255, 255, 255))
        bbox = ImageChops.difference(im, bg).getbbox()
        if bbox:
            l, t, r, b = bbox
            l = max(0, l - pad); t = max(0, t - pad)
            r = min(im.width, r + pad); b = min(im.height, b + pad)
            if r > l and b > t:
                im = im.crop((l, t, r, b))
        out = io.BytesIO()
        im.save(out, "PNG")
        return out.getvalue()
    except Exception:
        return png_bytes


def rasterize_blobs(jobs):
    """
    jobs : liste de (key, blob_bytes, ext).
    Convertit TOUTES les images non-web en PNG en UNE seule invocation LibreOffice
    (rapide), puis autocrop. Retourne {key: png_bytes}. Robuste : en cas d'échec,
    la clé est simplement absente du résultat (→ placeholder côté frontend).
    """
    result = {}
    if not jobs:
        return result
    with tempfile.TemporaryDirectory() as td:
        names = {}
        srcpaths = []
        for i, (key, blob, ext) in enumerate(jobs):
            fn = f"img{i}.{ext}"
            path = os.path.join(td, fn)
            with open(path, "wb") as f:
                f.write(blob)
            names[f"img{i}"] = key
            srcpaths.append(path)
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "png", "--outdir", td] + srcpaths,
                timeout=180, capture_output=True,
                env={**os.environ, "HOME": td},  # profil LibreOffice inscriptible
            )
        except Exception:
            return result
        for png in glob.glob(os.path.join(td, "*.png")):
            base = os.path.splitext(os.path.basename(png))[0]
            if base in names:
                try:
                    with open(png, "rb") as f:
                        result[names[base]] = _autocrop_png(f.read())
                except Exception:
                    pass
    return result


# ─────────────────────────────────────────────
# ─────────────────────────────────────────────
# V2.15 — TOUS LES FORMATS BUREAUTIQUES PASSENT PAR LIBREOFFICE
#
# Le serveur n'acceptait que odt, docx et pdf. Une enseignante qui dépose un
# vieux .doc — le format le plus répandu dans les classes — recevait « format
# non supporté » et refermait l'outil. Or LibreOffice tourne déjà sur ce
# serveur pour les ODT, et il ouvre nativement doc, rtf, odf et le reste.
# Le commentaire de la v2.14 le disait déjà : « le même mécanisme fonctionnerait
# pour doc/rtf si besoin un jour ». C'était une ligne.
#
# Le gain n'est pas seulement l'acceptation. Ces documents empruntent la route
# ODT → PDF, donc le modèle VOIT les pages, au lieu de ne recevoir qu'un texte
# extrait. C'est le meilleur des chemins disponibles.
#
# LE DOCX N'EST PAS DANS CETTE LISTE, ET C'EST VOLONTAIRE. Le faire basculer
# ici règlerait le défaut connu des zones de texte flottantes (invisibles à
# python-docx : ordre des jours, barème par item et titre perdus sur l'essai
# des saisons). Mais c'est un changement de pipeline pour un format qui MARCHE
# aujourd'hui : il se mesure sur un document témoin avant d'être fait, pas en
# même temps qu'une ouverture de formats. C'est le chantier v2.10 nº 1.
FORMATS_BUREAUTIQUES = ("odt", "doc", "rtf", "ott", "fodt", "sxw", "wps", "abw")

# ODT (et formats bureautiques) : conversion → PDF via LibreOffice
# La route ODT→PDF est volontaire : les ODT contiennent souvent des formes
# vectorielles dessinées nativement (quadrillages, figures géométriques…)
# qui seraient PERDUES en ODT→DOCX (DrawingML non extrait), alors que le
# pipeline PDF/PyMuPDF les récupère comme figures via get_drawings().
# ─────────────────────────────────────────────
def convert_office_to_pdf(content: bytes, ext: str):
    """
    Convertit un document bureautique (odt, doc, rtf…) en PDF via LibreOffice.
    Retourne (pdf_bytes, "") en cas de succès, (None, message_erreur) sinon —
    le message contient la sortie de soffice pour diagnostiquer sans deviner.
    """
    with tempfile.TemporaryDirectory() as td:
        src = os.path.join(td, f"doc.{ext}")
        with open(src, "wb") as f:
            f.write(content)
        try:
            proc = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", td, src],
                timeout=180, capture_output=True,
                env={**os.environ, "HOME": td},  # profil LibreOffice inscriptible
            )
        except subprocess.TimeoutExpired:
            return None, "timeout soffice (180 s)"
        except Exception as e:
            return None, f"soffice injoignable : {e}"
        pdf_path = os.path.join(td, "doc.pdf")
        if not os.path.exists(pdf_path):
            out = (proc.stdout or b"").decode(errors="replace")[-300:]
            err = (proc.stderr or b"").decode(errors="replace")[-300:]
            return None, f"soffice n'a pas produit de PDF. stdout: {out} | stderr: {err}"
        with open(pdf_path, "rb") as f:
            return f.read(), ""


# ─────────────────────────────────────────────
# V2.15 — LE SERVEUR ANNONCE LES FORMATS QU'IL SAIT LIRE
#
# La page portait sa propre liste d'extensions, écrite en dur. Deux sources de
# vérité : le jour où le serveur apprend un format, personne ne pense à mettre
# la page à jour, et la capacité existe sans être offerte — c'est exactement ce
# qui vient de se passer avec le .doc. Le serveur est le seul à savoir ce qu'il
# sait ouvrir : il le dit, la page l'affiche.
FORMATS_ACCEPTES = tuple(sorted(set(FORMATS_BUREAUTIQUES + ("pdf", "docx", "txt"))))


@app.get("/formats")
async def formats_acceptes():
    """Liste des extensions que /parse sait traiter. La page s'en sert pour
    construire son filtre de fichier et son message d'erreur — elle n'en tient
    plus sa propre copie."""
    return {
        "formats": list(FORMATS_ACCEPTES),
        "bureautiques": list(FORMATS_BUREAUTIQUES),
        "note": "les formats bureautiques passent par LibreOffice puis par le "
                "pipeline PDF : le modèle voit alors les pages",
    }


# ─────────────────────────────────────────────
# ENDPOINT : /parse
# DOCX / PDF / ODT → structure JSON pédagogique
# ─────────────────────────────────────────────
@app.post("/parse")
async def parse_document(file: UploadFile = File(...)):
    """
    Reçoit un PDF, un DOCX, ou tout format bureautique lisible par LibreOffice
    (odt, doc, rtf…).
    Retourne une structure JSON avec blocs texte, images, tableaux,
    type d'exercice détecté, et positions relatives.
    ODT : converti en PDF (LibreOffice) puis traité par le pipeline PDF.
    """
    content = await file.read()
    filename = file.filename or "document"
    ext = filename.rsplit(".", 1)[-1].lower()

    # Formats bureautiques (LibreOffice) : conversion en PDF puis pipeline PDF.
    if ext in FORMATS_BUREAUTIQUES:
        pdf_bytes, conv_err = convert_office_to_pdf(content, ext)
        if pdf_bytes is None:
            raise HTTPException(
                422, f"Conversion {ext.upper()}→PDF échouée (LibreOffice) : {conv_err}. "
                     "Enregistrez le document en PDF et réessayez.")
        result = parse_pdf(pdf_bytes, filename)
        result["source_format"] = ext
        # Joindre le PDF converti pour qu'il soit transmis à l'IA comme document
        # natif. V2.9.3 : le plafond passe de 150 000 à 4 000 000 caractères
        # base64 (~3 Mo de PDF), aligné sur PDF_B64_MAX du frontend v10.62.
        # L'ancienne valeur écartait SILENCIEUSEMENT les documents illustrés
        # (EvalSerie1maths6P : 1 474 ko, 1 341 % du plafond) : le modèle ne
        # recevait alors que le texte et les miniatures, sans jamais voir la page.
        # Les deux plafonds doivent rester égaux : c'est le frontend qui décide,
        # le backend ne doit plus filtrer en amont sans le dire.
        if len(pdf_bytes) * 4 / 3 < PDF_B64_MAX:
            result["pdf_b64"] = base64.b64encode(pdf_bytes).decode()
        else:
            result["pdf_b64_skipped_bytes"] = len(pdf_bytes)
        return result

    # PDF : extraction directe PyMuPDF (images raster + figures vectorielles)
    if ext == "pdf":
        return parse_pdf(content, filename)

    # Parser le DOCX
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"Impossible de lire le document : {e}")

    # Extraire toutes les images du document.
    # Les images web (png/jpeg/…) sont encodées telles quelles ; les formats
    # non affichables par un navigateur (EMF/WMF/TIFF vectoriels des DOCX Word)
    # sont rasterisés en PNG via LibreOffice — solution GÉNÉRIQUE, tous formats.
    images = {}
    raster_jobs = []
    for rId, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                img_blob = rel.target_part.blob
                content_type = rel.target_part.content_type or "image/png"
                if _is_web_safe(content_type):
                    img_b64 = base64.b64encode(img_blob).decode()
                    images[rId] = {
                        "data": f"data:{content_type};base64,{img_b64}",
                        "content_type": content_type,
                        "size": len(img_blob)
                    }
                else:
                    # EMF/WMF/TIFF/… → à rasteriser (une seule passe LibreOffice)
                    raster_jobs.append((rId, img_blob, _nonweb_ext(content_type) or "emf"))
            except Exception:
                pass

    # Rasterisation groupée des images non-web → PNG affichables
    for rId, png in rasterize_blobs(raster_jobs).items():
        img_b64 = base64.b64encode(png).decode()
        images[rId] = {
            "data": f"data:image/png;base64,{img_b64}",
            "content_type": "image/png",
            "size": len(png)
        }
    # Les images non converties (ex : WDP illisible) sont simplement absentes :
    # le frontend affichera alors un placeholder étiqueté « coller ici ».

    # Construire les blocs
    blocks = []
    img_counter = [0]  # compteur global d'images

    def extract_paragraph_images(para_element):
        """Extrait les images inline d'un paragraphe."""
        found = []
        for blip in para_element.findall(".//" + qn("a:blip")):
            rId = blip.get(qn("r:embed"))
            if rId and rId in images:
                idx = img_counter[0]
                img_counter[0] += 1
                found.append({
                    "index": idx,
                    "rId": rId,
                    "data": images[rId]["data"],
                    "content_type": images[rId]["content_type"]
                })
        return found

    def process_paragraph(para):
        text = para.text.strip()
        imgs = extract_paragraph_images(para._element)

        if not text and not imgs:
            return None

        # Détecter si c'est une consigne (début d'exercice)
        ex_type = detect_exercise_type(text) if text else "autre"
        is_consigne = ex_type != "autre" and len(text.split()) <= 15

        block = {
            "type": "paragraph",
            "text": text,
            "images": imgs,
            "exercise_type": ex_type if is_consigne else None,
            "is_consigne": is_consigne,
            "style": para.style.name if para.style else "Normal",
            "alignment": str(para.alignment) if para.alignment else "left",
        }
        return block

    def process_table(table):
        rows = []
        table_images = []
        # v2.35 — une cellule fusionnée est rendue une fois par case couverte
        # par python-docx : les références sont TENUES et chaque cellule ne
        # compte qu'une fois. La sérialisation lit la cellule AU NIVEAU DU
        # FICHIER : ses paragraphes directs (zones de texte ancrées comprises)
        # puis ses tableaux imbriqués — c'est la lecture mesurée au panel
        # (voir JOURNAL BACKEND v2.35).
        _refs = []
        _vues = set()
        lignes_texte = []

        def _texte_tc(tc_el):
            parts = []
            for p_el in tc_el.findall(qn("w:p")):
                z = "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()
                if z:
                    parts.append(z)
            for tbl2 in tc_el.findall(qn("w:tbl")):
                for tr2 in tbl2.findall(qn("w:tr")):
                    ligne2 = []
                    for tc2 in tr2.findall(qn("w:tc")):
                        _refs.append(tc2)
                        if id(tc2) in _vues:
                            continue
                        _vues.add(id(tc2))
                        v = _texte_tc(tc2)
                        if v:
                            ligne2.append(v)
                    if ligne2:
                        parts.append("\t".join(ligne2))
            return " ".join(parts).strip()

        for row in table.rows:
            cells = []
            cellules_ligne = []
            for cell in row.cells:
                _refs.append(cell._tc)
                cell_text = cell.text.strip()
                cell_imgs = extract_paragraph_images(cell._element)
                table_images.extend(cell_imgs)
                cells.append({
                    "text": cell_text,
                    "images": cell_imgs
                })
                if id(cell._tc) not in _vues:
                    _vues.add(id(cell._tc))
                    plein = _texte_tc(cell._tc)
                    if plein:
                        cellules_ligne.append(plein)
            rows.append(cells)
            if cellules_ligne:
                lignes_texte.append("\t".join(cellules_ligne))

        # v2.35 — le bloc tableau porte un champ texte : ses lignes
        # sérialisées, que la page transmet comme tout champ texte
        # (voir JOURNAL BACKEND v2.35).
        texte_serial = "\n".join(lignes_texte)
        ex_type = detect_exercise_type(texte_serial) if texte_serial else "autre"
        is_consigne = ex_type != "autre" and len(texte_serial.split()) <= 15

        return {
            "type": "table",
            "rows": rows,
            "text": texte_serial,
            "is_consigne": is_consigne,
            "images": table_images,
            "exercise_type": ex_type,
            "num_cols": len(rows[0]) if rows else 0,
            "num_rows": len(rows)
        }

    # Parcourir les éléments du document dans l'ordre
    from docx.oxml.ns import qn as oxqn
    body = doc.element.body

    def blocs_zones_de_texte(p_element):
        """v2.35 — chaque zone de texte ancrée dans ce paragraphe devient un
        bloc SÉPARÉ, posé après lui : le paragraphe garde son texte d'origine
        et ses consignes restent reconnues (voir JOURNAL BACKEND v2.35)."""
        blocs = []
        for tx in p_element.findall(".//" + oxqn("w:txbxContent")):
            z = "".join(t.text or "" for t in tx.iter(oxqn("w:t"))).strip()
            if not z:
                continue
            zt = detect_exercise_type(z)
            zc = zt != "autre" and len(z.split()) <= 15
            blocs.append({
                "type": "paragraph",
                "text": z,
                "images": [],
                "exercise_type": zt if zc else None,
                "is_consigne": zc,
                "style": "ZoneDeTexte",
                "alignment": "left",
            })
        return blocs

    def bloc_sdt(sdt_element):
        """v2.35 — un bloc structuré de premier niveau est lu comme un
        paragraphe (voir JOURNAL BACKEND v2.35)."""
        z = "".join(t.text or "" for t in sdt_element.iter(oxqn("w:t"))).strip()
        if not z:
            return None
        zt = detect_exercise_type(z)
        zc = zt != "autre" and len(z.split()) <= 15
        return {
            "type": "paragraph",
            "text": z,
            "images": [],
            "exercise_type": zt if zc else None,
            "is_consigne": zc,
            "style": "BlocStructure",
            "alignment": "left",
        }

    current_exercise = None
    exercise_blocks = []

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Trouver le paragraphe correspondant
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para:
                block = process_paragraph(para)
                if block:
                    blocks.append(block)
                # v2.35 — les zones ancrées ici suivent leur paragraphe
                blocks.extend(blocs_zones_de_texte(child))

        elif tag == "sdt":
            # v2.35 — bloc structuré de premier niveau
            b = bloc_sdt(child)
            if b:
                blocks.append(b)

        elif tag == "tbl":
            # Trouver le tableau correspondant
            for t in doc.tables:
                if t._element is child:
                    block = process_table(t)
                    blocks.append(block)
                    break

    # Grouper les blocs en exercices
    exercises = []
    current_ex = None

    for block in blocks:
        if block.get("is_consigne"):
            if current_ex:
                exercises.append(current_ex)
            current_ex = {
                "exercise_type": block["exercise_type"],
                "consigne": block["text"],
                "blocks": [block],
                "all_images": list(block["images"])
            }
        else:
            if current_ex:
                current_ex["blocks"].append(block)
                current_ex["all_images"].extend(block.get("images", []))
            else:
                # Bloc avant le premier exercice (titre, entête…)
                exercises.append({
                    "exercise_type": "header",
                    "consigne": None,
                    "blocks": [block],
                    "all_images": list(block.get("images", []))
                })

    if current_ex:
        exercises.append(current_ex)

    return {
        "filename": filename,
        "num_exercises": len([e for e in exercises if e["exercise_type"] != "header"]),
        "num_images": img_counter[0],
        "exercises": exercises,
        "raw_blocks": blocks  # Pour débogage
    }


# ─────────────────────────────────────────────
# ENDPOINT : /generate
# Proxy Anthropic avec retry + chunking
# ─────────────────────────────────────────────
# ─────────────────────────────────────────────
# ENDPOINT : /pictos  (V2.9.2 — ARASAAC)
# Mots-clés → pictogrammes ARASAAC en data-URL
# ─────────────────────────────────────────────
import urllib.request
import urllib.parse
import unicodedata

ARASAAC_API = "https://api.arasaac.org/v1/pictograms"
ARASAAC_STATIC = "https://static.arasaac.org/pictograms/{id}/{id}_300.png"
ARASAAC_CACHE_DIR = os.environ.get("ARASAAC_CACHE_DIR", os.path.join(tempfile.gettempdir(), "arasaac_cache"))
ARASAAC_TIMEOUT = 8
ARASAAC_MAX_MOTS = 24
ARASAAC_ATTRIBUTION = ("Pictogrammes : ARASAAC (arasaac.org) — auteur Sergio Palao, "
                       "propriété du Gouvernement d'Aragon, licence CC BY-NC-SA")

_arasaac_ids: dict = {}      # "fr:araignée" -> id ARASAAC (ou -1 = introuvable, cache négatif)
_arasaac_png: dict = {}      # id -> bytes PNG

_ARASAAC_DETS = ("l'", "d'", "le ", "la ", "les ", "un ", "une ", "des ", "du ", "de ")

def _arasaac_norm(mot: str) -> str:
    """Normalise un mot-clé : minuscules, espaces réduits, déterminant élidé."""
    m = (mot or "").strip().lower()
    m = re.sub(r"\s+", " ", m)
    for det in _ARASAAC_DETS:
        if m.startswith(det):
            m = m[len(det):]
            break
    return m.strip()

def _arasaac_http_json(url: str):
    req = urllib.request.Request(url, headers={"User-Agent": "Tailory/2.9.2"})
    with urllib.request.urlopen(req, timeout=ARASAAC_TIMEOUT) as r:
        return json.loads(r.read().decode("utf-8"))

def _arasaac_http_bytes(url: str) -> bytes:
    req = urllib.request.Request(url, headers={"User-Agent": "Tailory/2.9.2"})
    with urllib.request.urlopen(req, timeout=ARASAAC_TIMEOUT) as r:
        return r.read()

def _arasaac_search_id(mot: str, lang: str):
    """bestsearch puis repli search ; retourne l'id du 1er pictogramme ou None."""
    q = urllib.parse.quote(mot)
    for route in ("bestsearch", "search"):
        try:
            data = _arasaac_http_json(f"{ARASAAC_API}/{lang}/{route}/{q}")
            if isinstance(data, list) and data and isinstance(data[0], dict) and "_id" in data[0]:
                return int(data[0]["_id"])
        except Exception:
            continue
    return None

def _arasaac_png_by_id(pid: int):
    """PNG 300 px d'un picto — cache mémoire puis disque puis réseau."""
    if pid in _arasaac_png:
        return _arasaac_png[pid]
    os.makedirs(ARASAAC_CACHE_DIR, exist_ok=True)
    fp = os.path.join(ARASAAC_CACHE_DIR, f"{pid}_300.png")
    if os.path.exists(fp):
        with open(fp, "rb") as f:
            b = f.read()
        _arasaac_png[pid] = b
        return b
    try:
        b = _arasaac_http_bytes(ARASAAC_STATIC.format(id=pid))
    except Exception:
        return None
    if not b or len(b) < 100:
        return None
    with open(fp, "wb") as f:
        f.write(b)
    _arasaac_png[pid] = b
    return b

# ══ V2.18 — LE DICTIONNAIRE UNICODE DES ÉMOJIS (CLDR, embarqué) ══════════════
# Condition : la porte /pictos reçoit "emojis". Effet : chaque émoji est nommé
# par son nom OFFICIEL Unicode (annotations CLDR, 1966 entrées fr et de,
# consortium Unicode) puis résolu par la chaîne existante — le serveur sait
# nommer n'importe quel émoji de n'importe quel document futur, sans table
# entretenue à la main. Données embarquées à la livraison (zlib+base64),
# AUCUNE dépendance réseau nouvelle. Un émoji hors dictionnaire vaut null :
# le frontend garde son repli (l'émoji d'origine). Voir JOURNAL_BACKEND v2.18.
import zlib as _zlib
_CLDR_B64 = {
    "fr": (
        "eNqVfWlzI0l22F+B16ds0av7WB/yrlZSWHZICmlj/cHyUQCKYHUXqjB1cLpbloN9kexuAiAB8GY3+yBIAuz7vjtiEPoP+oz5Ad6f"
        "4HxXZlYVQPbEREwT+V5m5fny3fl3P/q7H/3kR06lEvpO1S2F6XzkBIn7o1//0a/22x8UqOE6aaniO17EhR+lsB5edIPArbtBkkH4"
        "JAj+aFAbHUeEUXek1c+6AV30RYpmw6AyGkDpeAHwksjxklKcxpXIS1Tp/9JlqrNxmPpejXqgYDMGNl7YTgMvDGbqoRfE2NpqEYrl"
        "XShXfYcW695MxalGTuIRrKdhmeJ1VVx2osgtzYWRdykMEseHLn9/FSdMfVINdzSoek5QOu8kznkncBT41xVw3otqKSL/44r5WXIi"
        "p4wtLFy2Sr3qaBDWIqcx532TAvjfSfMzVju7+ULd2k/Uf1U3vTCDYJiGf6L7B5PgXlDLphaBpuK7e9OAJS+Yd6N4NFBYf5RBUv+6"
        "URTWdBtfpoPtVv5x/xQ86f544ZPGwt4wyqyC/VuBIN6h/IpxW6Rxww1i6tH3C1d0I4X5/O5Nfr1U4b+E89AI4yQKG3NuqepUkot+"
        "ptp4YSuLU3PSyhxBtnO1o9BLCLIzCVIqO3FM8D0Fr6We76v9nBDYbvj2BDD+H6F3JkBV0zDGkwxIt/jdh2IVVfxrMCNOpE7U3Og4"
        "dkuzblRngvDfFagShap6InRClf4Pq5SRVenf20TFauP7hWXAn3PnI1jpQO1XIBxWe98v3JiIYdr+fuHmaW2UqmFapvO4cOu0pmzE"
        "FWsUqsdqBhLPLalZ8ICAZbrXPB3V7mfrq1otVSJ1SLFC+6vatiqsWhV8wKvkJiMIvQgx16ZhSqMas2dhukFpNKg4sEsyk7A+Bcce"
        "/cZZfTPD2Dyrc4L63ZHCjN2K0JvX6mfdidS55Ply8KDCov4n2IJRqE4B/PrX8CtORoPIi+kU/xha8p14Tv39t38LYPVdKfgX6rer"
        "yEjkp26CO/efAvrFejlUo6x6cDZU4T9H+pFGFZcp0VMuKNVhOrDorioKy6Nj/glUljZeySq9D/XSCv04MsTRLVXCtJFGOLx6SB95"
        "DvDIqxP2C9OgKXyJt53XyBS+hVlWZ5ta+ZCZumg0mFWT4wbUh8tXrPnC6v+bTjRc5WqK55GEvDRljldLVdH/NCUVL6qEwazvXkDc"
        "Y+imp9rjqfvuKd5OtQgvhPE13AJh42Lk1eaSUmMuDLI0dxG2+qw/OlabpAQXScl3LPq4uDMVXIK7mpmKxc4ENEOkF3engjOtrBXQ"
        "FC/gpDixi90JQKLF48W96b0Mg2rqywduT+9HBu+OhTe6W4LbvoRbx40nTdL+mejWOO5+RdvW1N37irZlGu5Pnwa1tSolt94A1hJx"
        "H0yfigLugYVLbF4pcdQGSiZNRv8MZGsqDs9s15qIozPblWk4tjDhWzMKoNiXqMREa7x4csZElcO0QiRk8eEZ82ShPrJQqbWZSZvr"
        "yeSNDAT/kle75NQQ6+mUs0BfpR0wafqf/ZB61vQ+n9KrMz734gdUs7728oxZraiahXV9ddq3rBqZDr62KsUoiigCytdgrPjZEtxP"
        "SuDAArUECm8WK745q+KkSm+tSkEYVWfC1I2T0rlU3Y5qz6oDrLatHwY13rtY6Z2pFMs2FN4R+G51XyoGV20fHipVpTl5b33Q6qDi"
        "BdwoKY4Q63z4ijo2Poimc47qQzD5sDhRWV3LYVA84J9+aE1Z6s/5ivYmzlfL7KwvP6ym2SlLC8XOZg761GEuXf6hNXmYS1eKndUb"
        "+rRRLl39QRWtQV6bfvCAQk7nBJau27s0M7VFEry0WCCFJSedqbqxkiKBL9KE3AxpCUUkj3iy3IcyHVk+BdFalBuntmd9+Oap7cmI"
        "bk3dyjBvE1cf661M3RsT6lmjtPky5kMnYDUNl3o2l7a0Nq1Na9o6079rjao9/btFvm6pO/27Mr29s4YLuyzbh1axD9lb1+7CxvQu"
        "GDqNmJtnYAre1nS8ODVo26ejme/anDbLGBMWfHc6ljU1p/DD6rYC3dsM8C1zpKxbOoUtnoR+51QWgweJmjxE3z/t2i5g2+wx8g+K"
        "lKlFnDQX905Ftfb0/TPatGbuwRltyn49mD7D2cve3PNL/enTPLWOzSKDtB5PnImjaVjWJBxPb8ka/2B6SzL0YRHldD5/aRrjpkbL"
        "c0YN5ImJaWJ5QTQAScglcHdWQZCeh/UhnQ4rB5bhevT8knvBi2lcy3B7Kb7GreMB9KpUCteUF1QUlWBl93gZrqHAKfsO/lpG7SI1"
        "roTv0V0shaslUD22IA2anGVUjYUBFVbCel1JbQPfah+o1qxiVWHEUJmFh+U2ah/CakpKeLUqs47WdiwD5Y5Va6TKzcP/DUBBZV9q"
        "+ClqJ0EKgL+1pn70RuNUQdMy77Emd7RpA+qpnyii4lVE9fzvcYZAh+G5avZpAm48xiszLuUh/0GVjwY1x0csOMpBCGosKvmP8KW0"
        "kWvrCbeVh3wHklcAdbkr/0f9RsME7A41cWQd+L9oTPB5QTt6MNpEcRnYXfeC4qBB46WLl7s8RYrvtUqNESIs+x5rcazpGl++Nglj"
        "NnJEeTZehhkNYTBOAqPJ6sfGy1tZcKSEQyzfzpaL3moZLgbVvKcGVTE32vJt2jCNMIIPBy7O+fIdW5+mtgmaAdR+8xC6j/q4mu8a"
        "HffyXSnDXw9QQ5cAFcJRQ094bvukw3YUQWT92jLQHbiXw5rok5YHuGESdX5dMyNAMrRtaPmEUNTSRg439NDc5TnI4wwyqe4CRQQi"
        "hL6kPkGZEpBUH/iUvaLyiluSfr3GWWzAdOHvN5npw6KP2QUwu+qT7LFS5Bpry3gZpJQYVImVudEgC7txFTV9F+uNJExgctBahydB"
        "Nv51oi1ReMGDcz0/AQVZ3ciFrVMy52oFaRkpr4t1Oqc2C7tdlddcXtUbwBjBFixVHSX2mY/A5arIZICdl7Zhd8SK7agk0uo3qTfv"
        "+IY23jjIEgz1OdNmP0sAMrDDPKFB/jgzsqM8/SiiDFFBOzoukqbxjRMNKxKhhzSlxIPceDGhn4WBvse1GQ0qo2Oeyg9YjY7VTZAJ"
        "A+DxPH053LyCVFxRZX0TUc2bV7FmVChfLFSwR3tzKV8tA10XAldxo4pP2/ImLDdSu0wp8ryhly3cOo0UCib1M0+6UMtut7VTpHkZ"
        "+O5UkplBuyNDcgMwHXPpvh5Tphh22xSO5OZjvsFhZDVXL9BTe03BEB6rU+3LBTS++YxXuAi6Bcq/wOMVUH8AlXIDR3MBuIy1KEwb"
        "KEfUadvfBLIURl7NC1Q1XvV3aEKAY6Z67EZ1D7iJAAnfubSWyqR/IWME0fNIUVtNxW/B5lO7lulykZ+4dSVHpSegXBWaPQF2fcJ6"
        "I2AxCxgNktDjOkDLzgEimF0CB/6R5joZ6p+5Om5tEUMVg5U5UnwZEZ7C6b61PRGvcNJvPcZZg73Opm2ya1eRbxrfepIHswWboE/z"
        "UAWqhYGTXtBbDPWQKA2Btl4z1reefWVV5otNzY27ZF9S21Kxdy45S2zcM4VixBtv3M8VlqB13BVK2org0Hk4jo0HeUSzW9XFX/jY"
        "gUGfcypzKf5hfDSEt93oT8DTvJrGOjRY36TquKKEhZCjCfWraW42p0gQbiAzSptq4/j0xrharjFr2geTuln8DNwc+vBZc/YoU27W"
        "6HGmXGbHqvjUsFkDdJAA6mHAz+1Tj8UF6rbxwsahBcZNB+tbQEYWyk0UUcs1O73Kq2KVs77y5swu2WLoxtuv7FSm0ruv7Fam0nux"
        "CvMXJknFGx8M0lTReePj9JVh0Xnj01csjOB+/up1kRpfvnpZuMbm5a9YFaNm2Lzy1atiVbr61atiVbo2dVUspOtTV8VCghvJB9mv"
        "Zp2izSWrVJ/NzeV8aZYuFptBnxrgTybiM8Sg3zTt4+9bprqoJDZXTJkbWCooWrCmgU6nsJstjQXSSQB+AxEwAtRG27RhurZqClEv"
        "p/gpJWAMhEzSIiDm2hRMps4WZmcKprm8M+jdKejmDs+g93LoSKSd7J1PHltmlOtfUUcuBlmQDXSoAPsfVbRnBBE2cwj5j26xEwq6"
        "AqkruVT2naCCoO3c1T5x/5j9uYe+NxUUgvlbILT57ikDwYq3f1BFrRvYvHN2PYsTM/X2f1A9q6N3c3NpeoXgezmwaQTB9ydyxIUu"
        "Wmvz4MwaheU+mFyluINMlf6ZVfL9OpQ9Q4QNy4A1VeyASE7WFn1R4Pt4c4NyMpq8tG/ypOv0Q4R1vnzd/ALu9lD7nsHP46OMipI+"
        "5il+tOKiSwviPLFUoAqeKA6TtFrHHy0IKoP0mg8eZEUdcZFBwTVGqQt8nIn0DT4VlAyig4C/vaDiewHxosPmRH1E7CkhiK3lGWl7"
        "uCOqT2QRoQuoq0T1YkEKGb7XgqQqAx1pRsvDWNvAQ5ZHg3qIs7D9hHy9yNHrV/tbC+jUHIMyR+ZDlV41pWr0sFpimHFLuH3i0kWa"
        "cYV8zSCzakRXUSMAPGo6JvTLk9qehLg4rROI2fC9OKYrTeFen4YLBoFamCaJS/7HMsQDVEOpQYepr7Y36P/A0Zkc7sRffuuKPTvS"
        "Ud+J6i59d9tC4M+haz0cGOgGt7MN85mMhgnuLdCnkpoPgSd3TRvfpB742Vfp6zcMYNbxYqHoPmv+/6Hr+YR489SZsud0aRqmA2xA"
        "6FN/+0+noVX+oavbap76VbWDZ0M1UUi/oRZN+yA3YVU3FgkfhSfFeDV8B0+0vce2TDU1deFFnouyp75MGJsTJ0sQxlvvi72lijvT"
        "K1IXfe4JerzqHbdtDSS7eyd9X83p89O2KO4pahgYurrn1OnXbnZzgB0V5sgHl2musHcaCpi8ZyG0YuLOORjqjQkGCJnL26e2qOcj"
        "cwQP1vC+roHGhC9qRiOjOxp/qccHm7kdL/0b3ffx7lIojwoHyy3VHdV/IHbZ5k7u53Dxq+SJFMPXp9V7kO1G5NacqIo+3eieFZeq"
        "oVdLeHgnxZOO4AkNH/Ts3SrCB7jx+u4F0kmpr9+bPAmx46cJ+Ch7iY7bOVjNfZw+hhvSRV3YaIBhPoR9nMP26RKLKp6PCkd3Xu9i"
        "q+HATROhe2vWXoW1dC/A3RJL37de5+D2rPazg88IH1B1vND81f6KNZ16twSp+slo7eJ0M0VVyym7ZKtjsBRDTGYTRW6tGwnoOwYU"
        "ZMjJQ4NQi7y6MzrS5OAx9q9jzaEavBcJ/MAaHtg61WaRK+XEqhO7lTCV7lxRTY4Xe/qoMdCdoP/S2OtF7LyMpjq7kqEpvlPjT25Z"
        "OxADbmapeGivTVWRZ48GdZDdi2VyIlcVXtofqAehL58+mUTI59XKxMwx6fNPDb3JodcdtmWqr3dysAQU1mF9dCz78WA9v/8VC0KU"
        "F6x4rFzG6SJ8iyIETgpuj9yPg8fZkc6HdY8X9igLUZeSGwXcxf6rzI6qIVmtzDkp3dT91/Y+dFMkuoqLZag1iYoRU0d1TpbplQ2B"
        "oVQ9BtAuPJlwyahDFcOG9PmoHDzRWwX6rXarHzJ3d3A3N2+qyxjLhxbQb8vhRergC+sKdOOEd0v/rbXVFVOTenLAtlp5CqPIcwLe"
        "5GjxUPSOr5du5rDUXPc8lh+tFk93PQzCCvMfW9Z6Q6BCypzSwWQymNOnqvr7BtEL1Ky4fLYs3jMbAakWBaVmlDu23hm0DATIw+Qu"
        "sN8r4TyxV7XszPGa2gxAAtw8kHwCvcjdRGqLMehkmFkEYKl4EfoWR+PV1aJrGtW3+h+5rAuAc2lxrluHk8Y4fUzW6VBcQIgXP0Fs"
        "Sjk7GzkXBfDUrnLB0/LhJHZ8q28LEeBxA7d2nKNq97PTpC7WVG6sR5NAwFDiKc2N/VkWeS6NyCugobuzkZl11H9FiXTjgU1cwVg/"
        "D+Yi2SF3MlVHR8xTdTPFdbVaPpynU6ZkYG+jhnX6TmyONfFqctH0rXGVR/fVras7lT85s2kdrpQ4jBK5HQMnUmwKz5DFnKjdUNV9"
        "upu53CuhD0eIzvrDCRcCe5HEfMdnuKSt5elCJqgcI+nM2pdJH52I2llAjcHoPun3sLuava2T24bCGqBVF6lVJQoTuTXuZahVxQ+/"
        "JTZjDc5TWOOBrr1HU3EQA5d4zmkosuNxR9EpXg1j9JKZ+TVQFqiFdrmhz1ZV9qxx44ZT4Q6QH0U55EV7jyGSToKbFKeIAW+LADMd"
        "BZnvnY3NB0N9/Fzo8UJ8EISiPA7zbYl8HwUTvyYMGS6KsGSf7K8JU2vEoO0FgWe27taXYicJ8NkG4F6kdpZRGRQg4wrDDdRdzgzw"
        "9o0CDNi0oCri+M0CvCqsdgdYKl9dZhE5FDr1kPd+B63IMBl4vZOTIUFuawjOS5SWHZqNzoaGqPvdU1REH8nOpgahRkVRVA3qalDZ"
        "SRJdfEeKYxLLVM8CuZg66xxRbq9XZ5/8G9UJ4fhtLcPD8bj9QH9nWuS5aqNnOhMxDRrfPkD+pNc3UxLMqmr1DHjwLjsvzLXFhKJh"
        "yFDhtwZvrELmYY7u6TK44Rizs6tLzzlpwIU7uhBuLirbNt1Xu4o+8ypTVjKiU2fPtOApwsWsYUuX1tEEjqUbZgSB7LzBa2tRmSoc"
        "NK2vkdJefegW5TyokMIJWUrQfoyO5yOhZU9Quwp+oZIjQBXez/rm2RT4BCqgDJWHdPrkvqgYQ72sJ6hrJJ4zjo1M0QHGgO6jOHsh"
        "oUwEu8mpq2sPi9bJSsxnuwNXQDn1fYn4FX3S2mXcD5vQAOge6DoAkV5j/+zniIkouSaMbl/Bn2i43APaeqo68ChT2wV/PoYcoHOW"
        "EXnWbtkuq+pmAiUBrdcOepLg4FFzgPIyNbOxim41upA1BBhkDfwBbu29W1kkqgmUgGT7+dSvyMdOngluwRisgM/zQGsmTl4IsOGk"
        "9byHtYK/nAxna7pCeCMISHKVyDbJ+Knw3p6CZ/V2DShnSAz+AfytiDVxFjxJ58w2PmiTK3YFBSkvqLrqykKHQ5CNaBahiV/y4sx7"
        "lSTU2pA7uEDYZgW8YUS3cPJUeiqmARe1T/ShDOrBvl7+c4rjBELv8Z19QN5CwFMY9YmtCdrWVVWPXV/JneCcDWHRaD2YVGdteULP"
        "tO1ae5IQ7o0zcO05XzwNt7DkG+syc0jq58Dllw/32tLXtMQ2f1QQnvWxk1fTkMBPP/EU++to3claU2+IQkNrrQKM+7F3k633NTpq"
        "hG4KQaepmatdXZyd74M9G2BPbpu8YH0lQAiJZLX/Ck9WzEec999rXUyK/AiFyzSwFOBrqxqH5SveVZg5AE8qyGJgpsYfvM8Lk3Jw"
        "m3pdC5AMuYZ+bbf1F0D/Fnl8DYz3mvZBViRL+MkOWk/cKPBQBR+isxd9Bb1YXX+WWcUOSKNlr+I2YuqmUxGp6BAYtXLkoLdbxQk8"
        "Ue8cAmt3zqmX3QLklUDoJ0xfw3NJa7J2Bb0aXfYXVOAPpqCEURMuxdUQ+lX0lb1EPMNdMkzOuw7dTCcLcgfTz8u0o4AVp7bfIPfE"
        "V+AhkM6QTy7U1Gqstct8g9EvmDdLE792jaieppxwlfKVTqw6CiH6O2uvyTZX5ik86uAlNSsM3xpexU40OuJLew1jdPSMHKHKvZr6"
        "smExYkPtFDjXitkIA2aW1+BWnQMXUqrXy+A5UTkVAOge0aiF2ATjtvMgX+xUaxjv45rWEXNBCu0PcCsWzHTyqGv3anS/ppf2pf4o"
        "llLhK92GhbrdtNuYVawaanYzGnBGzA3nbFzT6dNwW3YHciJLK/fNiWDzmRz4emZoxsSkGOrENbmaql5F8+7b1/Oj/CHVrAF/XbXF"
        "M3sI+UcaiWNV+ZoOTql1Rv/ytTqX7e5NYYo6l3MdOh3PdGEK3vat/H4AepBlOLdvTdgXU9Gy+6OIdtS2vwibU85XO/cZhNkg07RV"
        "bXspP4IYFx+MiYIxof9TkLK9zyMdHNrfIkbK4alFMxv4q4B+hrFz3/3KCqYPZ1Z4kx+7Wlhgf9FYPVDV9EX5ZsIcnIGcnYtpyEdr"
        "gLyzTt4rs2SFC1yfOWPY96NB6V8lYHKpuYG6an6Naf5X1GP6PRnTd4u42Jtf7be6JD2mVVQXTfl6Hs98LQ9x7cbbJ9asgytM7Co2"
        "R9c/Nkh50GASyLWnEaSwc6Ju4CmSspK5IXlKNMS63aiPK59x8ZQ06UbTh29j2T005VpE53ab6IaVerEXnNpyFs9u24bkW+8B/8BM"
        "WAXUXlPbz2PaX8jDguzqPUJnqvno1P7bWJm10+X5vndA9+iCneMiOIqC6iACU9bUT0ypYH8tg+LmcGTOQKkRwzgTb1azrjJNeWB+"
        "D01CKWymzgfyE4OInuSMhSniZgZUgGYXp3VA2lQ0Wp9ybG20zKm1AHG25WN0/0g4WNo0ZZUXJycDzc3LeA8jFz0/tFu0SwsH1oYV"
        "ZnlngfKmRRAwk+1kDlTsZx6h2HgHJYp6Q+/n40L5hGZtqGlzDYyM4ewsLCAatxuh71UEpi8aLNXfe6zvFCrX52f9FantEgiImpcy"
        "3YqBWP1j+EIWbnrYuUKSSVV+6ubIhca0xKAFDTKN9N9guFtwzqERvMmQ/WoKNqzE04rJNXPFCnmZhrNgE5Ei1sn1ApuRYmhcGoVy"
        "YA5eUia/QOb9rf4d885fe2E3A0ku0kgU/2svcvzARLDhAHLg51mepoLOEmAhrMx5QejJ0VvPcjKzYCxllu7gVa5zcT087wU1BhZ7"
        "NwGe6Z4NX3tqN1531EQPuA/zHKsHSJPYswI8x40Z+MFT6/MV15Ob/OSq1W0sp+Jr2QFbNQ6eoRXOd7zE6Lnluvxovo8Y4rhnxHK5"
        "cj+akUzBlEv8Y2Z/TWv2o5b+S9qu0cKtidv2L8LRCbtRoF4pVyh09xrqlYOEDtHhWwo7dqOZudEgYiXG4Vu9FDbMOqaMsZDFGD3h"
        "03D4TgPmwXc54NJcswSzm32Xa5YxLPq5jUqpmvzSLWJ8uaWz2NYt1Z0sR7ZjZhshowE4SstNfbSjm1QQNxGfHS5f4HIqQ29LR9Fj"
        "4T12dWUutru0q1sQoNWrPbJEDzBIteR8kzqJ6dOebjaJvIQFVC7FyfKi0bFc3aDpc/1Z+aWrQpndndu6OkKsvmAK4tEg8OS3IddY"
        "ardyx9BrglntgNL8UlgvSzv7uh0qtdvZ1+0wzGpnhaxGPu3kk+XMVM27Pk9TZ7EgerpkUq5Dama2Si1Okj6n4uUE0Al4S9M+Cule"
        "3QglNEhcLDrBztIpHTi9zuTOTKmz8zrfMUV4RdW483pCJ4rw7Act+FFGcVZ1yyErMI7yirIizDRqwzKKMNC4u0HIIznKK8Imgk2z"
        "OTARvsMnky9xh3SfTqDtKccGP3OnFhAHBtG6kYqI0oOPuSt2FnjD1PMhG3AYaX8d6cBH+7qdhjswuPrqnYIr3fg0rRvi+1p3gtT1"
        "M335NLEvkyoMTIVihyZUaF/Nb1PgqxKGTdilBXB2kxowj3eA2YkdJd9j1ijQL/tiqe9ctWDEyb7XJWyyNtpk1BLOzoKbNdpK/NFg"
        "3rM0hmv2FouNhw0bIWJSjzQ8wc4O7msr6OGeWeFoI7PlIQ69XudAgaONPLc1CWwxW1nwZn7RwC26wdt9c8KqFeHZZbPgB+8xH1AF"
        "kl3zIrSX0MaaRmoF57xGQzL/7GKGqfM6EKmNvEUArqogPEjpCqaBSCmvUKkW+rNcbu41VWbQzZWGxbw12tcm6aurEGoVcYPXLAYn"
        "Mg1es7iayDS488BuMHIok712dcCE3oyo21VYOIzpWAuCFbsT0do37a8CCyT6sPZN/Rkolt7f1M1iKUUG7r4rkFM4V5RAY/fdJFa+"
        "AM+x8hquvkmZL32Iugsbc8Lit/Oa76kopukiyg6c6MrFiq91DzsvdbtSbrEmDF2woYY52ckLTr/8xS8YUBSacrCMwCSwg7eTtpnv"
        "APHkA/J2ihVkAlLR6JFB+pihWHpPU/hOqj0dP+oPYqH4i3H5gi6PRWOduWPU0SMnwdK3TqKY+kboh4ym2zXHczLOguAQDc8jfZ78"
        "vTknqMK2YiQzbVyuo/s+648YkDjFv8vTunMhJ70BoDUCKJUG31m9xnKRwbcyjBO7ucJbIimpHvwwYff1o608K/UV2BZzdRr27kLG"
        "uBXp+O+yCGu7K7nLw/cy1+rBbfwHbcqUIF5vI7VE7ERuG5rWHgkm315T0E7MoUhKbHmajPlQGuT7czJap40+COJQ2lmjq6Qh8fjG"
        "zE+R2ToyouFEono4QBnLnwtJq4KKX6247GeAsYayk8cVMlxERr+whtksdWojPaXyT8eASz8pkcH+1/W/lsE/X+/rq09sZnLFSYhn"
        "42fqrT0Q78WY3PR4ahZJnU5TU6p6NS+RkIzDp/zeC3onCy9/+MwqxSz44qB1+MKGAPPN5c9JA52yJnP1le3ijU69VN4RH1+q1sTb"
        "PzLeJEfoTeIEtRklsrDebfW13VhlTrTrq5jzQf88fCw/S7XUYx3o6joxyu81DNnLOFb3iyNaxFVKCBR4Ikqsvsd0HmmDGr6JiX0C"
        "UeId3sa7PwkVb6r2Nm/r1WfZXjq0KKvL7CnNv4B8PNwV72ntonp42a7tC6+7+sQuTjzxdl+9jnkL9U9MhTIahA3p4+rLbG/ceYfF"
        "+hYeFJ/n9gSOLTvmq1otXl5GPryGiTnR7Y4KwD54aXRc5g8fck4TYswOHzIN4M4/tvsw7+jZhbNaVvSAaq1eRffQ2VneA6sYJ2iw"
        "32RGEnKuYgVAL27r9ybp3Gq+aN9XP2FeBXGCX22Tas/A19DxPxVVzyqmSplD1yAqABJSjcK6U9VRn6sntHh1sUgdbqO/UZ1VfbDD"
        "a17ksGZodYvmG1wR5aAePkB9WR0+PUcl0LUItMgVSz8I2wT48LARJg4zQquPMicLfP0Je/UyZRTUv4lD5X33Du/cuqKXPPSnmf2m"
        "hZlVTJ9jfn7B3lfUzZaye9Uhjt+JlcxLv8FtaQ5z+snCHy5pajBj9whdxYSrWf0AR+H2NS4Da4qZYzBAnQ8dn+Z0FfkncFOn5jld"
        "Kaj8hWId0oGV0MTDY1rZ2VlXuKtDWIjzECWVimf0ITixl/GzvJarn21aGROxCLy6I99BSdYLqrJjeuSyJlsXQwTCb+jHA4IpWqN9"
        "+/zRS16QAwOkgr6FbXtIrsLQwLggXcTHjjwYB1ddx9e9FJOmXfYw/6gn3NPhIpE3Tb8wmdzFmpxpej4sTM1UHaBDapW4vuFrSs0o"
        "ymPKI+3U9W5GTa8ja79HDE0UcqJr1QCG24hcsHqoSSCNyhDB4RcstPA4mKExB0pPWoLVt3ioQfOkr43Vm5ySJKzqz2DGmDBKxKux"
        "RSTykqaQTXSijBpielh9bh+JagShiwS4gVRA/z5cR0KTQopZvmYOgRb9YuavXe7jC0ywSRHmMWxD1MdAOosYsybP64/eMphUgOyW"
        "kypBiil0E69E7Z35COcjlN+r/DCgPnyrd01BCZ7Og+BkgtwzkBlfwjMOl/F2+0Z20+o2+Xummgju0KbG6bbiXGnJ3pJbtsPUYfiR"
        "vBCqwokfUuyXI1vzDvHZei/Se2yKodCbDzlHBzYYExggwnPp6LEc7tUV0lw4UY1jsg7Rn0tRLS3artIV6wZmj2DmfHX4OdRgFTMn"
        "lY0n6/A5ugTAY4pieVi9Q1cMZGjWLAo9veaZ7U05emZlSOtv0HrukBMwFb3Fzej5rN3KAA9RkaHmsCG3/iGlPVXcrbZNDJ9SkdyK"
        "w2dI8OrlyKPQlcO7mKBCHQKe6c4qub3qfq7QCzYs/6lL24vlIuw81iBbozp8gx8Raab9CgPhYr1UK++4gAzHC7oRxZLK6Fbe4yWm"
        "uIIKN7PyAScDoqTikHbNykfqa1RzeQpW8NZPfY+1VcMP6JSsFiLhzq3ADJxzQe1C0Q+EhwZiOG2c54DH/pxyiHqz2gC/8gKXqUzx"
        "qLNuSqq9imKwOZ+awqEk1L72FFrBGCWnIlOCLkMNTmVQijxydF75gsKV0Kcx5kmBJDzC4jQXTAF8HcxQaKeiThDOZdMI7hpOfkrA"
        "KxYQIii5+KpVPK+utfnQT7Sz1fAdWvarkl1KlbznEhTRgCPjb19jHqeh9imvyPCzni9MLmHGp/CXkNfw5DJrYmCFK4exSQ+2xMno"
        "WPZz8yayQOpKj4TyNTFnnpdIJFfzFl4A+wNdbOLHmhjJoq40qdvEVOfqPxpA/xE1Lz7gTQpWACnXxLU127rQhFQ3V4lI8i5prqHl"
        "dyhHosmZ7zzRwTTxBRQYO9vBMZLgYpSYlewDpTnvfUt0p4mse6if8j2BFkNffED6fUyt711gbpfu4D76lM+HFebnmnClO2nZjWoy"
        "ff2eHg2ICm7EI+ivUwoxHVC78gl5uoZH8QmjJ7yVyc8fYyPMHJ2s0XTMy6r0O3SSKkB7+AsP6WqtpfXsWbICKfqHGIegxuPTNBxd"
        "I6aAfmDWYbPX+nt0zBxF9o+1iL9F7+0ocYQDPlZIeh3dT4CecoetnPqghQZZU/f0BG+FMI45Ei3xEkzaK7ueNtxJZvPbwYUne2ak"
        "M/QRqnmHHhri3vMbo1bsaH8D92vN8FgnXfS98XVBH2PsIje5xCab/hOqIz/pG0HFOS8+qEvYRDorZvEF4osTerYg1Pb7JsXNq9NG"
        "7rRpUGIRo7nJbok8I8DJco8HeOFEFU5MQdUJ1EUZga+7Zo9EizJQb6KSTXwXRnKzqd/oxOpdukT8+8oj5AHgOVlyXOkPSXKrfqsk"
        "cEJB0dHh3b9C4YURGOVp8jBri1NnxqYPLMt5t+yUaR7wiSrFQcy67P/f38EAERY40RcTfoG+TRFvHvAWkXLI8qMEwoDNeU2K9oNU"
        "3bxw6N4Dz/7xbDzAu8tnNRdsfX3H9EkkxTd9LRJDbz43ZkCuph4DvSyrm1NW8goF+FAjJxStjoFOFNUdxBBFQ409w88r7ggzGOCF"
        "yKuwZXaYfTs1t+myLQJ2eCR2GZwY9QOmSkmBEaUmae4x5myqxYAmULm44dTmVLtM6pAVdRInwQeSWKBp3mdSMCcdsNnXJmaAS2NO"
        "z9E8wJmB402Pu9qsYrNPB+2bdFJL/cfoLXE/4R2sOd7mvXwHymFQlqCj/n5uehrQBIHukn6jAun92OeLAPfMIigCqARIuJJ5GTBc"
        "CJOngd7NOH82jzjroydu3M1jQaWfA0qeXpsL2HzfxNSLYXieZaTWFXuIKKsGHkQ7aEm6+dTCiBX/IFbqy3jwG0BNaAhHeKT0Hj3R"
        "TwQo0RCepU6Eo2uilgdnjH4/whWrmGV5TJ4jx3XrVQ0Goe02xKWse7zB0TGr7KnLTFZulyLfOB0D5WMZY/qZMgs2lD+EDiRQt0RJ"
        "iMa5svmKUknwndx8jcfwvEzzZw6DpaA10tbV1c6VhXljf35eeBqSdCrnE7m1mu+skqyo1SR9Y8jca/OD/CS9telpn72NArWHab/0"
        "r+qvJyk8/cDBcZtUrE4+pjNUkoFwcf0D1KKWXbVl+XVOxyi9b/HsYPYIeJ9Zbg9Uu5EcQxv9nNg4KNRVBxfcpF2p4936963LjCn8"
        "Jxau+Ey56C+Dm47yGMZK1tD74+U0ZGr+Gt1JFh3t0b5P5SiTIz0vP5Ezmfk2zLxazblQ+HzUMftqeig1A6YdjjDc9aezOHxMhad+"
        "/glI9FypNbmSWr+f1tmPjXmQ9vTmY0+3/VMlzEFuGW5+VVeitFvQZFWRBmpz8z2xbAmak+o6Fm/zs13+55DkhFboEW3oOJbQ/naP"
        "UpoksKlLQHI8jsYb7z61QGSDIhUl0jztiLkC+0aJDxVWFG9+4GqlP03PEW1uE2dZb4gDanuDrlDJNQO57kXEa++R13IsLHwb7ovR"
        "Yx+vrVhfju07/L4PPCThQfIoKsbI8MThmWjvUm9SZFcrPuwvbXnf5PcikijlFzDoDbOsu/HRM73/Sb45RmFW7UlWRr1ickNztIFK"
        "CNnC7S0K2o3l1mjv6IKSAxk1kINg5LsaRr/vWbgwT+cUL8VUpg1ny1NMfEpvxXDoA6mV2g+ID+JTACbHXK6b9kERgwB91NK+bMAd"
        "RCWHJELpaTsmjEQ8hgYo/KMDui4b6udz9GluU8qKit56D3UyY8WAOiISYghJqk9t+4mw7tDV7BieWiBST3SIB/OEod38yCoEGOQv"
        "wvMXiS3c/ESbJNEPo/xXuFWIkI13h5TH1xdRbR3jxcP4m1RWahfz+0HoCTjAqEWk76+jqvBi4NTCGnvEDIhjrSQpvRqssJk3Xb+F"
        "ema1CRDxOXGL6oqg5Ee77/ETIpOvXMZdKMrVFYyXBvaCsq/qLUR5TYWor6B7MWZ0g1B3Uj9E7FgPRErON9Oo63l0Kl6URpG7iyCd"
        "leJt5BvsDAQHIpdAbgXl+ZCUHgE/i7R9zPaHCnpIYnovYlLuaqMSSZzB6JjXcfc2rmM5rNWYzLTuye5xjf9E675FsuJSBNoeTryy"
        "bILK2OssKouypjUkQU7tedlJO+hcofiZeqijLnZgzr/V2t0dmNtf/Nkv6cd1Mql6Abo7Ur/mQbAR56JFQaCfS6T+VMwA/V7WvzkD"
        "MRXfoMALbuMmtVH/1iH2euc2jdeoVXfu6H5gyiyLeu/cku6j3G83g0oSvtR3mvQjY1vYaYnLr3tRI6ICwws8XYC6ByXf+dpgihEy"
        "ivLzE0J2hM0OZlQLvYTzWVlBMju9SaBsh/AAOhfogsFIP/iVxdk0zVDBVrZdg4mnZjSY82BXlzC7RULntSFZx3bxARCvcn6Grb07"
        "O5mx+d58ZEj3zi4KZnVvJnLrRim/s8eWdx2Z1W5l+0RedQTChA8hk4vdV6T9DUVeIJ/OU1w5yQt1svMpUZf0vGIdKeXLznOaAvbg"
        "2X1JC64ktEBTcHIyI70vbPOIhR+++cj4ltARyMJ2MBdHFI2GCftjUJUHqJ1SHQslpGv3gOYDs+NGmt4SJxl5PglacE5IZbr7iWOw"
        "SKBSx0323u4dvKdl5lFkugiOZ/o49UkJChkLwSoW2RldCeOwgCFZXOkT+CZnEvKGOKLFnadreIw7XE0Wf26X5UUOVYzB6pnQRTXG"
        "xYXYGTkdu6/JFBeOTqjtA1SFRPFMOWR9yi7GL8HjZpRxboxrBrNGp3q3j3OG96aaGrX+sud27qPz4Dw/7o5hd+oXb128sui1Fqvw"
        "RHDgCsCcZL4vt+zuQw3k/OhOAnpF0GsxyvAKP2lUmZOV7mACfE/T+J3LZB8GN4JGoln6nX1N0eSZLiqn5IBkMLcjRnbuGSpZcjDp"
        "Po/hKd7FikfxRUeDgYGzqRY7d9/S3YQ+/CXF1Wqd+C56PPghvlQLxxRedWRz5BFZ8Wq8mCt49pUITYs5br8wvznzHujnxT2a9vHK"
        "Dl9dtCbI20Rg/0PCPm6Tp40SfE1i3XGb4swiP+QpXKfALjVFfDevI/urJIYfK0KdssfY+lG2sIQakzqLHOur+OgbPDsndHN9zy7K"
        "oWuHMwTSuV6/nS3MVenod9DsOndypblKcJjY/GHX2s8X56ph+jslNWcq3c0W5qqg1dbLDuhepixXAXODuY0kU+N+tjBXZROlBC9b"
        "5UG2MFdlCzPkpLOZKgfZwlyVbXRCyI6knynLVcA3GoNL2Sk+zBZmq6ysoRlC0TJgD7XqbKWDTqgB5+2K9Tla6dJ7phh5r58nsbVu"
        "Kz2pWvPKZfQwzbdBz1qi0dxU2yhUAzKVq0kP9UTB9K9vSTMTam/b+vpstZ38NHA8pMlGgWRh2siL2Hun9LSIjTQvm3d7jI+D2Nz4"
        "7ey0FVu5Y9j/PGy4ync/2FmAzENYDV90j1Y5GaH4VKzsmwJIO+kbOrpy1wLNehZgRS5+H7NS0RAuS05EvOAs8SQrxujU8+PdZZ1F"
        "EYfAmfUxj5yfyqY9mNSQkvzp2pMPKsT+FMQaxF9aeIdT8HQ+R+vjR9keWpDjLASUMQwZTBwVwYbkwBKIzgOT85cjdHsR48sKpvfV"
        "KaLtRRdH+lC28jK93DTjBjMVUb2uyP2NveWnNHCFsgBKioaAXgZg5UmO7fkY7z5jPDV9ksOFdEFjvMnnMGI/gSz4FK2AnmCzIJvh"
        "T3zsphwGEt6mZ228e20ShIzSGgmfv6ngqWEVInOFxKMc6WyalF2VxcMVfLLFEatx66qYno2pmyN7GySJmcDf1iLzlPL6+6yIPS0U"
        "pNUxh5zeTs2Emx6jaXmQ0GKO947pCGE+WtHStpbJP0gto5pmyOYqLgAtdNsKE/R5YWRyBAvA6Y2VKa1btpfD/D90WZXUQnFQSXUh"
        "s32QipbMYy1O7NeADHpaNSP9aaNnnhcooT0iKa21SlpJ7EgJnqsRzqyFrjyKVYPM2hC0KhEPEkOLj4wHiriGjYZtmW+h76LO19tC"
        "Y4U6B7xMrU0BozQMrYu9o4XcOKwW581VnJdwga0TDaPfG+yr5JAbReZxjzYJ9iEwoVxyPYOOIqS22i9lYPBldo/sLxcg9BwLQW9k"
        "oUBW4U7GY/LJLD0+AxMmHGQx3kELhhO7Ouqi30JKNZvokvYC4Zzn0bZXSaGrRP0Zg7SM6ilqueSgHlpnQW3fwGmulUnwaH0mzVUg"
        "Rq82shz4thWx1ISGXHP4ra+1wm1ymKronqwhf6t+X0TKUHMuiWa0k4NoW1i/SXe06F3aXfaCnFHbk7/zFgdcrYMqji0VaLrgS70c"
        "XuC2btFAUlkOeAsYjqzoWfvXcb8xuXlB3n/yQgPHtY1RptByselp64H2GOGXk3ztnHDwBRWZMb/c3FBLq5PktT6RVlitd5LTGLS+"
        "UPwdSSzolCUnpb9C3npsrKqkkZ731hOOsSzR5hY3N0ydGLLKdHiZHbhnrCXsYaonT0nAcIxA/paD90xso2TGMhdR77ENqcBbASxM"
        "D69Z1qIMAWw9Jg8u8EdSZIO4UnohO9A2KabYLUoiC0ZXnFZ4y4QAz9muQBQFhVx1TZNeqZFeuiR2sLfiaAxuYa5vPOmuU5VnDidr"
        "GA70UMCZxakYfcdw0SaNWg853kZDt4RLbpv822isIRR5G5nmcrz9wHIBs9DQnVOMFvAoiUv+fVfb2vbz5+F5VicsXNOeqZytuu7M"
        "nZMD0XrJ8wuxq+iYx1OGBtsGPPJHbw1qjTA+dAvvrPPF7mkXEkwJ1DDOBA2XnvumuSVTSZl9bmbFF/QeuSGTOxo8+hCmkuz96DXl"
        "zCREVFyoCePMlV3m1YWpmOcNs/46B7C44P4nC0YlaFb2QxAZbK9GjAWpeT45FMTo3a5N8Gs9cledF/estXX07I3nPFYYrmH6MZev"
        "pqMHZMvUB+HogEkOb8g+7WetgT46ZN/4ODZdXdukJwb4k+ip7oECmAbyAVkOvm4Hz7J0AkLJMPiUtj3hoCur75ECa4CKijnReK6h"
        "wdQ77wWccxufWgO/Z/IwyWQDGj5CpgRYTNFFr1HKXrUTZ+qQ3IO53LU97GUFlVI6eIzy4ZKhl8hXk9Bo/ebUxa/pVYd3uGZxBqit"
        "lO28dkfPXMrPAFipm/azQFvF2/+ch0VZA13/SwYBnIC41bsZgBqZomphYCVBvseeSw4RPjNxg6e6qmLP0KMcwt2pFqo9Qz5K1mQ/"
        "Rj0PJbGeZRPCGkfSmSRIax22a/CNZBpoDZizJoZBHq/DrH4JuMmojYL+IRajc0RPy0sqI8MDcJ82NDjDIY131wygyu9hRTrWt/tF"
        "+qiTAFwjve2ElPpNuSHoZ8tcaQ3FvCnWWZzJe8hmiUNzb5k4GkzIjLkCKhy2Met4sod6N6Yh1cOLzKf2bk7DGQ1MBuvufZ7cGQxZ"
        "ZKVr94E+DUrKJf6h+4Sc4tH+ytdgr6cVjfR7Hb26HHyLUV1fDiTz0rzAR/bk83TWg9Yr1ETwmympMXe3XjMgLkC2xSUcFwiSOtKW"
        "aqGJA9gLsqDCnpAbfpd9Y3jTRmAw1ZfAgTRIP4/0DqBhf0DXS/0VzBrgXAghBzs38J6Y6npD04PhTS5SEp1cEvCAIUOBMDoVNZfA"
        "IfBMYABIqnai8Pxov/AcJpetD2jWDMX7eLhOhu5zTE0vk/48cSOmXEOgt9Vzbl1yQA0p63rkVNideIhssj/6IJ2+Qa6o7C3SfWbp"
        "rEN0SLeoWPc5JSzHpJkJv0g0xse6rDqESvEG9TK8VZvLYE8IJN7oF5G6d9EWRNuud4vy7HEfT/hX5jRgvN7oflnCcmc95FPVNq8k"
        "RtfeoaTeVS/AXPDZ8Wz0s8BsfsaNY3RgaMD7faQDGqP+o+I782IF2XimQ9VYVJLGCPw8w0sqGq9dTzufjMwxoWbnswbL/uoAIfpj"
        "esaiC6zvz39JP44eY/yHo09Ln+x+4GZPBeiJAlYFbRPvfqJsrOdctPSVKrBOdScJa0o+mDMH7yENmK7f7nu90vOeTnjSfUO7AvOl"
        "l9Qisz2w+7ZQzivlOzE5x3bfSU/jijgsdYlzI6bCZqh7TQ5ldUvw6DKkjHVzqfZ7rckoVtL99Se0JDXJOH+PvKx05FvvkLRV0mb7"
        "MQVFIitrBPthDxfoIs1wt0c3fiB+axXSMwlf2V3H8NP5yNJMqdINXarVUqp0U5easi1dpl+Z6W6b2ubhmu6OlNKB73Z1z+h3x/RU"
        "zERcExPr0CNYYBwVEt7dYz+mObfOzFD3GqPSr6eUayESFyd8kp0LStQUIa7JyxbMU/Y2iDyp7Tajm2u/Ye2RtfeH21oaolV7yiwa"
        "vUif00UM0RQWzs5G6A8CepOYlRidl0arosSni+Ko1nmVKa+C8VAeOXqdAaltzy7nnTcZgDXpnbcGonPuQNwgg19oCagMbvCagfmM"
        "AWzypFvnHfNLZf3uLDw0xZmcUdF2w1Y70RyjJmqmLjxuF+NH3Zgf65O2u4OMworPpXlyqntgOWPTe9IhMdndvgXBaxXSoMvl3j2k"
        "CE7WrHRPDDI+bY0PX8V8CDA2EXk/izcZfl0NJSzoJwK7DyfX4QcwJn3l0dfVyHzl8cQ6tONhRSEOjC+gkMMKx3ukMnIuki54D7MG"
        "JBf9sIRxqCZidONWFkKFN6WQfpKKJKjoK6qpW4ckCeLc1r1Nurg6VcMsxMTYVcNKap4Z6V6ml4h09Fz3iimwSdLmFUz7qxgNCGf0"
        "NG0h7ZIPL5/pJhYpJmcOXyKNxLN3s0MibGWGGDxQfFjBl5vdTEt5cHfJIlpw1qsYMs29WCZHK7q2MgdEAW/kgbikDLxpA7FT6CjO"
        "7d7S0wYdpzJcgTTQQVLdJocg6Qwv3ZbFBtIqLVkl3DZGyY+OdSUyUIHrNLs8YG4hjComlmjzasYxnxeJd981POPoaFmAkdAVUXgo"
        "gCG6VKC9DuujAye2b6Ze1yq3NkKvXcBnhwG9SXurU1Aqcg30sEvya/M2/yopUmhiCnrHdDVodcOwi5wqEajxbpu0S/xzp2Ow0ZLm"
        "afFk924WhJ9CdxMTsbR5DxWM7BKJ3l1qTVE1xq86CeWmEBkJTx/SooR1Fy5gur3ekXEq/4bf7j1+u0V70Ax3OHE3jfdIpiHXt+Gu"
        "DjKdZ6ra45c74RtoA0Y9HPqIgTtRqUqcDw1tj4LDI3cG38dCTIrhQi843BGhV+VEDU/IqyeXkLC3STLw6HFAFprxbpfygfc1wOXn"
        "+2h5uqacBY/bHNg+x4zI0dPMhg7Bp401TShmeDoeZ7hHfrlzOnJ52GaGltZrkwRQJWzQKT1Cg2NaFo8rhwP1TjI3FzppMQT9g6os"
        "Qz8U2dAwpb1HzPm6FqOKnCO8+wih8Rmvnc4NmnRPh20O3lp2OjTx0Xbp3CTZRrhPfNIwBmsjXEFUxHl8vzGmuwF6Jqn+JHNWJwda"
        "aiWaxh5EQyFjtA9xsuKKazIWDu/ggCOdvACD3txgNJQ4dTxtks9p9xYnYVATiEEBjiTiG67RqgtxRLsPGtKN3vkum3lDndTuC8UG"
        "6YNCUi1ELApfNDxA70kUtVllLpk4hh0cc6w9nV7mI2L8UDMnR28sYk0q0w+Rq6ND3vGp4D3xnqIU5bgefaDj5YtHuqM9RYcY6aX1"
        "oh+RN52XSMsjeTOQe9wn4zcEKgKx1rfxEUezhsLKHz3BVxXBjKrlwl3WmEUex4PvoITmgQsuUYzxzlMyhFRS0SgPD40eKlHEi3mT"
        "nWfCtMymAQiLmg8l4046Dy8k69eohg/leS81+j91Ek4Ts4nIIc/a8EhiO2XDDYcm1AOeY4EBayV1+4jEXHUQy6mWwZFxZr9MYLq8"
        "yuil6uVfhepIyxHAUcJSQEYZEujH21+sxyt/WqmMjmPM6OZVYK9yg+8snF/I3uSMYITxfiIGJSEjjA80SPAvMRgE+mhV/plJxo12"
        "zv/2x7SGV8i4reSSlz7p6Rog78IhZbfTq3QmJCgcHQXUAkeQvhfwyZePka9zlAgb58boeKjf8AR2CX1C7Uc4d96yJyhFtDGLP8vK"
        "/vFuDzd0EJf4ESaudZJruPBE084LcdI1VXliHgkLyfio2FXMKHt8PiHjV2P0MhF9CPpyZTvwjBcdopnNwivAG9M6vS6shMVBIu4E"
        "3VdZFVWs9UD5bvaAFv7mH5TA2EgeOfQg4X0hraDdRKspvtJdBbk7Ujs29EMxB44fLpKODe5+UbCPFzdzhUYpMb5zz4KZ4sUtq7gM"
        "z0lr0MOlLAjRt3PoWjEyfnjdApnixY18nyzYeh4mn+kV2rL7/LgAnSEostbMHiwOLCyWcK1PD4tQ8wV8OCEHlUnGRwlyMHmvtHfV"
        "wOS9cFD2B3D4De3rXZuKpk6Hl0G1J/yv3UQecu7tWOV/yiqT3q5V+FNSkRNgzwL8DJVWnFmpd9uC/Illq9ldxevY5QhCn+n+bvbp"
        "XnJzDuXFDWBJwjrKFvcsDzeF+HNn3kO/HXz3mzjItPRzdckgiR9vwdm8qOj+ReJaxnvMy3lKCAa3AfJTOtSF6mzNhVVyUhhvDcVf"
        "0ZVYICp/nO1uw2Gbw3qL34Cpur7IguCmWo6QIeXVfGKNQWH8XhY8BPD5OU60NcbHvDlfWulSWPUcMXDfQM4i1bbsm5SXvi6eg9vM"
        "7FTIUxof5uScdmN8DHAe3aHwZ8tw1fi7bWXDyX51FTmEmpdoG9Q22aAakU5RN97uSJws961rkh4RlW6ZvE9mUy5os0cJNBH4OIAP"
        "ue81O9W7bFAilHMs60zvymQgezw12C423nhd+A56fA9MsTOPAobiRUkLMG4/MkBsqBSnaLIn6JNCiz9uOJw+eNMaViRHbdwe5kvt"
        "rz3Ofw1tcKNBVces9T7qSIIEHqavmaf1xhgNqO5gEKrlEBae3+t9KjbAdBKjB6fUF5S3Vg9lpChAciFHbowxFJjL4DWlGnBM7Lvf"
        "NiBImlCR6JvWIfmQg1qf+oqviaZ1LwhjYMAy1pTFHIzNhSzpvJYQerGkaWZ3l2zlkEVHFI74bnSI0WXlSBsJXuYMS/DqD4ujlOja"
        "IlyONgRzbuwsURNb+RjjWbIVwcaSwhtKYHVSFCHG16eQaG3oZ5XrimR6DYl+uEPO2nTR30GNtBKU6demecTbA5MHDWb/qS4eDWqO"
        "D+F05PPymZ7HIoeE8cJHswXwiWLisy5UICUYn7jx5Rv8PHAS58ElVnchixKFNV3ndlfqFOHaRjG+3TsFS56KH99ez2DZnzc4m1Nx"
        "5GvfLzxF0giTAilFtPmg84zYUzTWEx2pgiaDoc/tJUStO64sBWRX0wpzDNsfLLy/dsEdmL3advbsZGD+Rb4fnlE2Av0WbneXfOpq"
        "oHOtg6mNnUR7T62WxUcUe6m4d8jcinv00ToLVb7x7hzvXTdnr4Ivi+PFxu4ELpcRao9TfNLKrMhdib9aVjNSeOep6IuIRfxitpIu"
        "/X4B8xGq+z8K5NbZe0EviWEaS3wwDuI1rMtxvPfSujoL0NtLpGyhZDLfDazJqYSNi5FXm0PAY225gYlqhKL5uXbfAOAwqinzHD+L"
        "pITiffkKBvuLmaxHurpzFLzLK3mPojgzZRi7Alm6JJEu2HoY9iAD484z7ADnqzHnlFFxkTBzdv0ppVcFwyGoZlLF7PwUIYutCZCf"
        "UaVnRRBBFtfMgv7JLPjdEd1Y7JjyPw45ncBi1xT+WeQkHIR87Z0OQubTy4/uiUyDfDjX+89Vev2PNyuq8evo/BFpPeviusH/C1Dp"
        "MgVHIYDL/0pJDWXxeb7+uTi+v6Qqm6bKX/4XQv5itUJIW6bkb/7yb6hs25T9Mn+XLu5YQMJftjgVr+KB+j4TYL9sMSuYuYz8d6J5"
        "iMYtIL8xyBh3DMGoiq9P0Z0hh/vavpljIPAJBqLl8Z5kuCX4NCjFclg3Vg1WmdKj06sgxfbeZdqrctqHApo1VVO79tyaO5GMCz2z"
        "diq9Z6wY5EJLFpsCz9V7FWdyr14aPPTvaSBDNhpAKtgJzb6wyB6k3UL2wosVioeSebbC9zetXTeLUZj8lkoR09pmsVuJ3AmzY3FU"
        "ZN+g3trPtWQrvLJpPeRVKDbae2l8OoynwP5dU2oZ6vfvmeJzjgQS7N+3XNG11eWVKdTm//0HFqYXinPa/oEp1vnexqi54VI2OY5R"
        "MaxbpZte1e+zd6/ax9YYjkypPYZjU2yNYWBK9Rj2D02hGcPQwrTGcGKKzRgw0S6l42CYjOThXh6iOZfNj1ZT4CanK2EamAzEVPps"
        "QMTc56p+mgzXDWwMtViQ6+rGSR5ipr73Wg/DD4HiZvZL700Bquex91a3OqnmuwLU1CQ/NhZdyKkRGT2gjkVx58NXYIvCpUP7nl6p"
        "4fezCJvasu4H42rXe2FvEz0zz02ptpi3L1vhDKBJd3Ry052BBaL+piYnd2vFQC+IgY8ztry0Kpov2YETplftF/TY8rIFzYelIQ4J"
        "KqY/llgCnwSMrbvZ2AwwfqOQ+h06tRIX+8+QB2Qe+bsH4ugDP+DUkhPo+ArWiNJLrofujeMrDyixZ1CBSzIYHRG9Gl85wBZIPLqC"
        "7t0eReVcgfn7lo7dFZBjwekFf3D8dnARnp10wZ0NizFEzo0JBzlQyhKFv19rPyU0Y4b89Tfkux+rj1I9ysGbNhTj5gVVMQWPr3yi"
        "clbyX0FOw0sqlAz/N8ZXH5CBFDjY8cKTn5R+QxX/ZrH4N1XxbxWLf0sV/3ax+LdV8e8Ui39HFf9usfh3VfHvFYt/TxX/frH491Xx"
        "HxSL/0AV/2Gx+A9hfd+hKagRxvTsChQ9t4uqlD/+uxd2IQar45wd66kFpFcWUy9OsP/vzUfMj6jW7O//P266rlE="
    ),
    "de": (
        "eNqVfWl3G0eS4F/B9F7du811u3vO3lOyLLtbtuQRZfuNd3Z3CkQBKKNQha5DtDg7+6iLpCgS4AHelKgDBERQ90FdlPWe+SfmM/oH"
        "bP+EzTgyK7NQIDXPfk9ERGRWnnFlZOQ//uwff/bbnx3sFYue7RXsXMkOR8qjtlOM7Nwp16pW7eBnv/zZn7YabwVZ2XZdO/e5FUdF"
        "K8jbhNgXiKoTRbabiX7HaNcO+nA/qqKF2Kv0l30v8BmY3jgUPNjM20EYBc5IWcD+t4B87UUaZEhAjjuiTwQZ+tLx4hBLz5kYhC0K"
        "2OeWmy/Z5x07iDRMU2A+64MuCeioZZVEr0bK4rO5YYn742UYqlNWZFUszxKfhf7VRC8igfwloPxq1RJ//fOM+GEFVt4RQ26HOQn/"
        "4/hFgXAKtl8KrGIK+Z8EatiuOhXf9T2oZMOsRMf9Vvx3wq/Vks//mYAci8MgLtpjtmiuDWQ/3RbQuFqyK3Y5iEQdfRT/TRCcDERn"
        "tULvU4VS+H/eMtuVQvfGYV38DiYs8POWVxKw/yggX3FDe+Nt8etT13VqoY1jcik9JkFOEv/0UuDMcf53AiLmxioIsmM1X0ycX6OJ"
        "W4UhkJCc63gVWhRrBhxnlRDrGYhcLFpO/djUt0+YK+B44yh6xYP9chB7pZD7nfPzXOiGKCT64Dr2wda/pNzNf8HHZBN/2hWFPGuE"
        "O5sTmzuk4p/FjtiwVRvH8K2k4h5mk/0i1fCc+GRB5xT/w+Am9kjFKeno/5ku30fxT2mKbI70x/Ep40thzYnGTIJr6ZoyaKaNSuR4"
        "ZlFeT9d2GPGMUe2nI5WKhpxN15TC143CZb/s9pE00lVkUs2ZIyRG0QpEQyu+V7HOmw2e7xurw6ibZsVRXK0VTYqlvvoyiJYzOpr1"
        "uZXs3vaT/tQBFiK4DjAJ+L0HezcfWtFYwnr+O8Cioe8U4N8DSxWMyOPfH8HvkXJw0Cophv/3f2+wO8HuUwT/FtiVXuu/Usy3Etjx"
        "mID8G2hc4I/ZHnG4x/S7CrsLAbcE4AvBC2WB3jhw5u/E2s9bbsX2Evgd5OPFsYNW2YU9X1NsE0Zg2HEFw4gC2/NM5FP8olOlzz1L"
        "xIOCPQdYILZ+IuleCdAnVmBTFW9JgNnfO0Eclh2AXQTWfCwUo+KEVQH4BxCYgXXeAf75E9R4rBJD6f8FvXGCSlwtuvYPgLwnIOcC"
        "GyXbT4+pXCF3zK34Zd+Fuq/A3H/i1y4ETqkcSdYmeJ1g63HB8a246FllavzEXDajE4KiaDtY3cQ6ahTBSLnEwt/2JD6XFEXShQEM"
        "Ua9u46jqEjkyMS/rA26eWduipEDWnUkC0uZbULO8Wlabb2ShtTaA/KDWCbUrR5PP7CtV09YRlFIkTdw6glD7+u0jSJVQnbgjKZEN"
        "eWPp1t3NxGufaqlPaQvivO+xmibGtuSKtSOGF8m3TXKcocHUbZOa18dg+o5JT7M7mPxeRmOElOVyMHSmotub2M3qLRAKLuUKyUnb"
        "435ms/vIHgxcYtgKbZAf4X4eqYxZI5WhdAeR4nFGV/CDmWvuyQdQa59/mjWsg2t/9iHkWvXPBw9X5pp9Mah+IZX7GrPH5kYFeG3x"
        "YC+oCi1Ico6SXRJFvy4L8QvQ0PESFXDi5WElneqgYq/M1p32g8Ko4FniO2Ll5VwhdWy1HsdifTm+zlpcUChNnTvYz2dxv0xirPoN"
        "CJMPHAEs8HZQgXTHkXofhyoQ8s9Ob4zvxJIVXzBZ2bsPo0+m/Md0AbVkdXptTb3/oAJqmUyOp+m1VZjVhcmLH1hA9WHyUrpEsnCz"
        "OzF5+cNKJL24krmT+lZLUuCqISqB62SIRSqKPGtyAjQnKE4si6o+i9+pSdk5OYk2Yc2yhALHhdMfnhpEogb42iAKbYCmB9Eko349"
        "e6lRwzMmEAvNDJjcjFJJl+ayhC5txhQl2CUnBigzR5WdP+IravwWjiDUuts4RLk6svDiEd9JpqL5IQOUEn2T9X9x2zKrWT7i2wmb"
        "RvKVDyD3FfXqEdTDB/sa9doHUGtNWTfJBU8uCqrsxbFxJK02JJtkfXmCwVeEZZLFHm4MItGquTlAl5BWezX3t7GuQ01uDRLfA0vc"
        "QscgT78bDWJmtw8jU/vizmFUWr/uHkaXLOrWhwptpN7OZNADyUH/FVxYGOMDutzJJlCdvZeN17q5k02RdLA7+BODVevJTB2NBsez"
        "oWhGa6dACKPFabGhPgUysya4u4MO72N5oTxHMfowe1MgHe0wZ/8gbGjwGSMQBJprgwf8S1toWgibQAeoUOOq7BCYAgF12sq7Fv6a"
        "QgdoggXJU7EdL8dA6CYiQJaI3+WDlhuJNoY5vRQwqk/RXSUaddxG4YlDPNUg90MhrkSaX3YKOPlwXBUVJND/AKRuHGru3ycMqoJX"
        "PYEfgHp6whG2v4OtOwCW9WXsRk7NdSpWRND/DF0RgwY8UzQYvnvtIch8UU2UMzH/BRa8Cx8oi1HWPtW7BhtQDDsiBeC/gtgNDvZw"
        "TchqH6lqU6ifwBw6bZewTVq1/wecTehrSfn1/y/4Khy3QLO3gC5no++9i3hG4pNQEEsDHLNpmqlFPPMQwJwPpn3NpUOJqSVt4MKE"
        "y1wE5el4EEt+ibQrusdq6EzNDqzIDxAFTP+s45VM6JrUm00wMPG/ja2CgIzGQntGPWkK2Gst8Gt+AENjuWL3I/wmHrGIreSKdlRS"
        "kzEF/FOO2bdiR3FlMEf0U/ON9aaAjZ0TWy0yoNu0r2Cfqa8CL3H9Eh8jfE0beWqHRYATRVW1o4AnfCM2meM5JXB8JZhdPmwoiboR"
        "cF/5vhwdDIvwVByctz0DDKbgSd8tuXICXqCNTHrl1B59Fvef5xiD8pL2GA+lgdqXy8mckneoO4Y1R1hA4FxOVtyPqQMX8IhaRXLR"
        "D+NaB9ZjhReqtciPYLxyamv0rl2lvWIXD1rlYEhDTOkIvQSeTomZCHXgQgY1LeTjoqGSJV1b5lUIdtmAzXubznM8bClC7mLXPSvO"
        "HbT+EDvnLZeZ2LWWyS/og1qrtlM7v4+gnaqANGaNoJOuoY8CFtd5wfBzaca1KxF9vAcW2dgorlzq87PMdqZ7C1bxeT8Ak7zMguua"
        "PJ0pilXIsOmLirWlMZd4d6kNMA1L44zoEULkNE1P6HQoNrWBA+UHqSbTZTOpgH+JgRHrE93HYtmYEqM3vdxHkWaO0yv9JIbsMIhX"
        "U8RBLoN/Tq/1UeGxobHrptf7iPp46PRGHw2yYJPopiQKrHI1axC20gR9Y7CtCyBTFZkG/mQHYk15JYGCuUfwY3PF5ISScNZ2pbTt"
        "TT8xV08af/2BrkHQkoLV6flBVSxLcjeWgliwy/TOmgbmdyZwSg4ICrmuXmOsQCCWBpzd4UnN93EJtCGhAFlibn5AsvdKZIRoFdgj"
        "FWwOLOzTQ2EUFx3BCnMJ9z9TwGOe3vVLOkUuLQiuXzbQmSLh+lV52mxM4PUJKVpN8BSOYCS0QJABWGfFO9ivFVntuw6MUZ0I9jVo"
        "FedUcMwsDnJ9zcCm2ch1mPTzYuzEHoAzPnUm3rv+CPWNwBnzvUjDoeYuYy6C3OcHe2WifywHPBwSJKMHrUCo+WIpOyU60XSsEgh8"
        "8yNPtEJ5LlMEwTCwyPItOrDDA8QwxxoGYm4rtd2E3+mHo/Fl5cVOgSNmUNTBkYoLbPlu1gdgWVMdVaPuVqLlYJ3aiNE5nlUsxjjN"
        "y9spUjXsacJ2ilAI+LJPIxHkTsE5XWjQd1L0qjbip/nEDEnsksQcWr53dHE22tJ2DRbfSRVXs5bZ1vvG4KIXjXfm8gNtmgyEsUbD"
        "3IDij42av2JND0I6qlUkeGoQ9J+WGXxi+ZnWmqNonx9ZM+1Ks9SLI76QVeZlukxfQIVO/aq/XX30WV95feRXskq90b4W5LJPN79D"
        "JzqSv9WM7COJ9/t7Yjpszaa8S3fgMOIfj647q7vvj/pGRqGVi+lC6bAZnfhSf8PS5FnfuHzUN7IKXemfvPRBdzIdK1fTc3cI7YTR"
        "i1NW4CN4UmulAk6lgMhQqLlVte+rCf01rR1ogg4qADgsAe7zs1YcoSxZuY4yruzK4ggE2wTDGlCn0DGz0tI9lH+vgFtEaCUQixGg"
        "ckFV0CcbaIgUD/Zd16x8zqg8mYvPheIE4UmK467MH0GZ8PaVhSNIlc9rZfEISqkmrjQzCau5b3C4XF1gIP2SsbKSklKaGuWSToJC"
        "D79s1QJatLIqjweVqFeUtvdB5KAyOR5OoM1RiUm8i6Yvr6xl6gJ685MJBOcydwa1+HSPbmTitZG6mSZIi+qVrWyK5Bu39APcZM5u"
        "G+Bk0u8Y3WN+0P/Zu4eQJd9uZVKlBmF7MJH2wbbGBwZNzSO0d33c+317+NkAlW/QIn2ZyUgGUeusf9CgrYFNfwxWpOv5xG7udUxX"
        "J35DrMCicukJxinqRdpH6PipCBth6CPUo2rkTkXkPq5foeaDppPAd+4a/iywaFKnDKngjp13NB0QbWY4XFJWRHc208chpsVzVZnE"
        "cOuCycvVsa8jZXd036SsRiIrWmGU07yuvTVQCj8NPMcuBGRkaUtg7REiy8IEMTF/2lodx9XheDLM1A7ZoBW4y5k4HClaUJ5YdCVZ"
        "05XB1K4FHyyY9CDixQhbZfcDC0xgaBp8oKSRjgmCiucUi7Znkl8d3J5hjGU92IJo4iLTt3DzwkwIIz73BTYAbbiCWA7HfWiLsLvc"
        "MVn9JToDUnWeFGsPPGliiXhMs3YJTxyQxD1owVlR3yCvXU55FHXkLttwVRu+2zc/18i1JXh/0I8EAZ7x0cHDO3lIgc9txwX7FPaL"
        "Q/Tbjw+lD9RAzR5CVxZ0MopEb0wLjCVc6jDmkVPiTq2iazgMxayP2kExY0xAwFUO9sO+9dxbfZPdECq3nlmO1jr4T1wfUOaIrQ0u"
        "k3xHL7ENppUcfLPAOVw3VO91uRILjp37whE8SUgGG+zuvmZvpFahGFArDvE8XBRARhmT/0MQb2avlyMLtoBFO0GQ/viND/s4nreP"
        "xSEECHvZm7U1n6rqhO+6IPoSTtVaSZHEVQu8oTDCQpWBf4nuAVnBLkYvVaCC/t2xeydVma83CRv8uaW4LXCAL2N064qid1NFMQIB"
        "p+okuhRBiOOKqGJt1Khd9BlWrbLXP4mtJvN40YWMOd4FpSS03Bh0sIzSc6n2nAXOJgYAFq3YKlXZ8Na99HThIdmYX0p6nhcTKFc3"
        "BvLaseDQbnriYbbETBeCeKQSip2Rxu9pX/JFl7XR26ajFXV0ni7ZG5/909bMrlaB4+W+9d2KbBZG3R/sF4tOaHmD+erqAi2DMHKq"
        "1fRn1q5g6DZMtSApBHY5Y2vdJ3FTtWB/g37s2XBU0kf3EJu8oI3uL3MFC07uQyuiCxti8LeR+eyXMuYXSo7BxRlap3QiVtKbe0l8"
        "ojeBBoUQWKIpB/twz8bTsUuM9ZwRNVYzuOpIAtl90mVVX3foUU3huxzyf9AqBugENtYdqrFWPp8lfNjRo47GjP7upNYhTkQAe4eb"
        "/TJNIDQyUVscjYnJqPCOWkjLX8cWTYnA4+2LUecbc62lFBkMkZjzvCV3BTo9D/aDjJlpPUQDNcqQva0OnvvYGWrT9gtW3Z2sgtuw"
        "N8RwZm3lbekfg2txbt+Ir0LFQsP2gY/24WgV7qa6O1xzYP96qIxq/BY0QvuHmusXuCUBjgwhb6Uq+cQfzfsXyjG38hnd74gu6Gt0"
        "+xXtN1B+C7zoV+uHiP5hX+jlguHwxQvx2UX0+gcwjXrNnTSD+9L3fDrZFp9Yos+Ognu6X39qpYqyCq/z89UtHNXQD0p9bOIiooRK"
        "t5+PUsoEuP4KQvZHdp9YfJgpW/RPPqJGF8SK7tuUT3HPhpEFhlIK9wxxwkSI7OhIhXy3S18ZEyXSNW2/oZ5FGZxz+3Wak5Fu4gkZ"
        "LkZipIInVNSgNp1ZhXZGR2CH/CG23PO+2ydAgOWJOkthlMV3HmfNyEDFfRVYq5BSMfEoCyT2IUyPDzky5PDqAzwxi90M1BM63Msd"
        "gzbjMhLWbD/ZMq9Gp3/IV+/yWCldMoW/SfjooBWHdBxnoBdxUAo4A4OGhSh3aAUJKbFXK/bVAxyierDfN+vQwxIYqhndYo+FZ8X5"
        "DOxtNlJ8N2vkbtGQHLTEpu4blNb9NL+6UM37wK8yNK/VqcEW3efCiBA2N62I+ffE06Mj6BbA+D3nC7IK8b4etlZBWPvHG1/gtDjl"
        "+UodXdhBeRuVMbCNOgMD8Ynrj3pDeh/nX2PIDpwKxyyV5jHm3hc9pZ7Nv6VhEIp+yIXeoYYCrgAnKNBVW0L8iHexqrUYHBQ1x3aH"
        "BD8MlbzDoA9RswSg1aMGLXfKEkKJEK/6EYONw9d0m5pJMmzd1beaaaPVRpZgSvys7vM+QR09kRB6897pxpWGWBs3l7de5r22t3X4"
        "j+YK1Wub0m0tUgnH4rLl8to5Brzb0MWu6QXOCD5wVIFpVUCQfkm20CH0CzMYRi+kYCiEo00ieQEMXxhKHFN5E03Ab+hw7QaRQMFK"
        "KMYQ9IUfAipCrOCt+ZEyb6MEsSi1tlIaA7wJfaI4fJp1v7CENqXtGMAtblXBrviBinC4cVc2llZD6n67KNdE4yAf4DZLGtC70ULt"
        "prktKxA2wUm8cOpp6J3X+mB8o5Q8gQevtx/pXdoBJbMGtrAG7MAOphAOvfcb2Cw3bwDXKYxo3zOg4BLIu8KI0j/1grM2pBELm+i7"
        "d60E1AKVCawwo9pl02Wc1LyHjTBrbc1qXtfkW9fZdWLl0WwjICgiH//qVxSNw+sGBOTBlaDEDGT3EW79aq04ClYY0cBEfMcWfRj5"
        "LHlQ+2T9VaivBGtTkGCf2EaDaVioOfmk3iVQcL7w0aUpKEA4DNfAVSxGLiR+Iow1h5vLzP4iTv7KPTbpYG1ohZAGkXpN0s8uUI9S"
        "KAwylsIDVmkotkPirxbNeoDMumCBsZ58Y6FFHUWbh1oGQ47uFtw0yuLAOKegDM5yAmP8QNnxZJeW5/AYQiDY/SSa56AEIg8DLvrN"
        "6+lqqCyqILFbsTxO0PBZEPP87D4ZcJs2Kb77NPv6rkbxTG8c/FEkBm4GhwvC54cS8pmLoHtpflK0Grif+dFXqZZnEc0D5zxzSgtm"
        "FYMNsMT3JNUnHkgiaWAcqeeUQGzTCcTmDB4gjUR+cEFBBelNc2KUZpDMi2jr40yiKp31FslHBGLghBVXZTOBY/5upDzkQq+GCpAq"
        "RfvsKoYICm7phiU1my3gNWfj4lAVqC1vCL6ps9P5qUFXsbWWEOW1gdes+0gnBtyg1nuXXs9LKjWIXtPkoJvW5jd76Gf8kE8KBZ1d"
        "qLsv0ENVzB0X1ptF8dBjHGGV0Se8dIWzgU4xAtYTYAAtxLZsTuu77qQVS32NrpvlLTid0uAtkB3407yeITCbJkbjMPMNPEWyopDV"
        "KfJvkjoxIwWusfd5a/EKxIutWOaXqKagUMxDEDOvDGAwZJbqdbee8sqVtaV2ausG40lJoPY0jA3m2qWor0W9Tc62IdhYas8uXMXr"
        "FqKg0B1JKLWeYXYdsVSYsYIdiyK5ZnkwD7njzphdo6Fqg2p3LKjWhHgXdiSVaIMeeFzogSb0BUPpF4zRSbmX5uGkRihzhHpLP5TB"
        "UALzhecZoxolz++Qp6TsyK0/LjWQn58JSpb3C4KCG+EL5U1vA7/7zipzO4BL6obF/Lj0TbJ445/06yoKXlnV/BW82CGZZJsXPR4Z"
        "5PKoAnAt0NvjVv4CNRujHRzJNqHY75M6wW7/8qBVUC3q4L1CoXL7XNkT9fu3ubzrk65oWbyZQNZ+aXlctqnRHrcCduk0wWWJx2NA"
        "qWHmn6QwWD+hdigBUqxXMc6w/ioSTFJFZxH9DnCAbOf0Hj1P4EFONX7+hUauvrw2i5dHwLMeQWSrVs/arGp9QqBVyPhxs4Kk5nrK"
        "TtGrrquqNQq97rqqW68iqfyqmglc2GhgKwF2LJCiaO2qMQMfQiuH+nDaCfP7fGajjliV+3Ftor8BRxEbLRhIvHARb6AUfTFAYIAF"
        "kBcM0pk54u+fO7xdFy6q7w+i1enGD6FjXrN23ey65NxVGB6m6O/zQCqjs/1UHYqcEuZxAEcz+iLqNNRnEgJtDTF+3KwgWUKT6ohc"
        "GFd2XOpbpJPJIk2R6St1Mlmp6crUt1ptibY8tJxhFiHuIUx/tNU2P5pFr32dyccPrT5pxkvaTzFM7VhMZ7DG11/q+9Ig0z/6Ut+c"
        "ZmUJS5sHovUllGhj0UfCCBuLeAnN3zNwzBAVTCfFav60VSd3Ovpp1NrGakxcUpUJN2trgGn3hS0keJCuTEPpdWlgvareOl1LQp3M"
        "rMpA6T3UwGarZn5EsRbbwUfHD1qxQnPLErTeMILq5Fzb7DP0yo2UPzp1sCfsC6MuhdSrIqBGzDU1UYYKhU4YYZX+EUuh9RpTqNQk"
        "PECVMR84lT6uJSejn8SYlH60+YmFfRyc/cBHjYuirX4eGN/IpDE+k0lhDNB9vJEIuw34ZhEuSfSNUgaNMVQZ+FRnQJH73bmhT3+o"
        "gev5I/WnObUpMqMjBiq1wdDgh1OM/rbrOGODafBUbfdwKe17YdZYmFijRgNj7LRNsP++clw/tf81hLbLEqDZsPVxSiwW+J4V93ES"
        "E6s3zMSk6gQF9CS4w0cFg6gKLvnRUFGyQVlzH41Ru4FNysLZ21e+64w5odbWh4pHS5wOH9fg3M4l0PxO2JFdiZzzqp6lF6oeidPh"
        "4xqc61kA2m/BK84/VQUADH23YEU6ZtzAyEgrkEGnHe97i1r9krhaPwtIRJGG1lHjJorr301phqfElLFB+Jxy1XnkQ5x/JX/aYSj3"
        "0DOz8LkY3K6M6ldq+tGGNqOjn+rlgMPaIXq0qrnPeaV1lvSyIMGjmBl060XSLKeaG676FccrMcpsVjZaNSuFnn9s9pd87nKYH/f3"
        "OIvA6LNB0FK3c4BFaMrB7mUDoekWu1fMIrpy8gTlORyxG7tndp8bQO4BMAwpgMm1MAy/5OQjQ/jtc48EPdjfA+mlJN1Xg3TkF5A0"
        "jpjv1q/Spe4yeB6iUG371p4BLxo6kyh2xUDXkiFov8JUCgJStt2C2ivtV2qmFFJHkCHvFqS34DUKtL3QHlUHNe3XibEnODoG+MLg"
        "ZFKNk08+n03UWcPxhcuZqoWdNW0tleT6YOi4gkrGuk5WewBWe9FmYbue0cLcSROLa1GCwFv1jVWtOfy1jcwu9lOkuqcTYLJE0KvB"
        "Egx5e3Y2NUaIOGWCbCYdtO3g+9grJZN9A7NxFH9u8xDdULUIqA4aJxD3Cry2vPDEFpHcv3Mzs3MDCVN9zKADJ+53fjXPjqvOVuYH"
        "+ilSNesEM5ibUxjztNVBRpctiyIdvoXoI5IeE6n9BmegyBAheg3OL/J2xa9WWdaYBjbt6Q8pMW5wjcNLTCZsUgibKvqHbBVExzQm"
        "tzyCTjHNQXTr0qVXSill63vqQxpaR42bKLmp0G0ZUXSgYUcnfh6J1q3oxMujCidGHl568py+KhP/DmP1GhPfjix6MsUB249MuQQp"
        "bMPIZ58q83EkUpIpRbKTkCjZZJLIT+1rstXL2a7QeAIWzWfFUg2juOwan91XsvZw8p2EnGTvoeSyOe+M5ghOEkNGoEFteae1ZSDt"
        "TkIrGzKItnEZAx8wvZExo43LakYlWptSxo7rhdNzuvMY9TXx0QjCgoTaJVfmAnw0srwxs9zSGw2sfWxJBhbSkRlhQLM55o3FrNck"
        "64cv3HxuCfbiQ4ADEyS7FbgZX5nJppJ7NbuqzrIxYyfgcDm0Ys9ibEo/G0Qxrs1OmmIFff1Bie7Va1J1JdGRE7SOGjdRPOCtNygm"
        "Ta9FY5IyWQUFiKGkSJjeBqjnwxWnaBnOkgaqGh5EUlpAnyCAxX/mC2Glw1QrCaNDxxWU29a4glpMYNRwRdNvAlXDFVUDQbmG9buG"
        "tnw2Fg3M+2wert/t05czCXSN2SRoTHMMgAMpwpM2TidtlEgdMa4jHBrc1yaXg81B8H7FO4UzdG7GiU9dxxMcVMXKEHmatC5xyep4"
        "HTeewsnRfI53ZwupFbD+XFWokDpiXEfIqsCM+dKPvchyvLzhS1rXzBidQEeOp5FyNb+ib+WSaB/DbHilt1Qj0j2ar/RG6xUlvtN9"
        "TkCktbq1r6qmiG9RgNmJjh/X8SeTSwmtd0pjhKNWjH/Ta3+XUis1Ip1gPItAjs2PfMiZ+YEf1QdSJDp6vB8tK4f1+3sfLu7FWq3J"
        "6pU4HT6uwaVmsmps2C/8yBcWecRRTKt9GzaTQN+wJsEGoPJWIb0wNoBXCerjQiDlXMiOldZiSHC1buA/eHCoO9nlyYwMf8tJgUOy"
        "44E8tsazZ0Wkzf/8ruwUBFBg44lMP8aev59dj77OFhqyadoFrq/kGebCvIrBq0kYXe7FQKFUvPx8i9L9l/0YIqpzdPvzOAaHE8G2"
        "SVBFr7iiYTNeHWCoi02p4emi77bquI7hS5X/LCTo3+Z+nfs04Bg/+5e5j3PqtDdd6Miyv8ayUkZmfetjjf7nwS8yP7fwIeWMT83f"
        "ZZ0cAtYgaocjO8DUoXAbPZKt/TgJ8ZMz1n6Ct3zxbRMdrN95TqBPMcWhCgydA87bF5w5t8BQKjOLohiuMJDa0e5gIiahNA99HbEf"
        "aw6DMeB6glHREkOpHLgtj7vQea8sgXNLpIS+QT9qCNmuvLEEu4N+kwLfg5gDqm99toLb0zh2Ymjp5w3idyPl/EGLujv3RAb0mv2b"
        "kmD+CYzi/oYWfajFzrYxruFgb5SJH2EyvZJ5e2PuqoTST0wSbfs1oQQR4LnSofRiu5gy02V3wS7s1k9D2de6LELNAPXnU8cr+xyF"
        "0YZTrO/sfMCTghmbnUD6Htr38eUpybLmKN1g2Wg0hoPAyqSflxN5ATfPZDuuUEn68VLFOXrmTC8nCAKs4LmFWwh14DuteIAsiT+C"
        "0WlOQW6LuXkitGim5+Yw/b/NwRtzXXw6xK/aBV7Xc7sUvsnVtdfweZMqDw1eonIgJQeXX6V8p0WLDwbbd5HhVqvscm03KP4lGewN"
        "dE5DAmA5ow/IRIhDcxwuYk0xLck5Uhyk228OQ3etqopjn6MgujC1PCcZSr/e84UhR+xSOK32VMhKexfnWOprbYhC+V1JjsEk+oPs"
        "AridZHve0qEgf/0trPsbtLDCZNPgWYtv8W6fQztYyHkezG0K6HIj6axrwxn2mUj2qU0Rrw5lTCEQXmgGp0MckNLUvo1hZ3LrzoE2"
        "cg73qckH25fRMxRHZRlJNNdEliJ/0Vtq8tddeZVvv8ZX8k4d7Mt7gnMtOkNSv7fV71Rc3Rz06BtfjuQcJmoQzDiW5ybT2NyY345r"
        "w+Y/VpDqUXsC8+DynLfv8JJnRtmGIMRPY74N1m5hctCCj7+66Nqx5SZo75BHryo+TQRtSoXNmmd7E2emJjRSh8qjJY/3L7nl7T7O"
        "FuSSjnVhaX1meWFCzOHnX4mF5mHqPAF/hRqULznL3DSmfPErfoHj8+e4j2K3V/DaFjUPmZtTsBMOMyuDh2UY1txT3MkWuRi1LXBN"
        "wakuECTDVgz5TVn/aQPHOTd01uZWPqOYdA6E+Nbipl1HpsY/MBUrhmcSw51FQTbCK/sBRuTZifDZQiGs+n2LvJQ12ysmwNvIHUtQ"
        "qWK9U7goaU7mgBWdqUR+Te7Bdc4ni1fdy3Jrdl/hrgsseVGwu0+Z1OTv9jhl2smztMMg3VhZk228P2iJJSzp51EkQdCp1rQuvcqn"
        "GNCM1hiiuE6QKuhwgSsPhOY2KFNSjecNlt6xKqcxEr9vICeyeQN1n9KGlxb53E3kioEgqCTQ9grdPVY9xjgDSxhwriVD5VHe1CC5"
        "Cf1+pX6L/1mPb6NtXfGDmiOPQ7aQOSs20n2MG0MJkC7oBd/GHErUvoWBu5XA5w29gK82ujGG/Qvtm8MnZ3CKUMTmXQidJ+KHGrFd"
        "rcm99RKD/KM41GgbL/DeUmhLiTDzmgF0BEoucrdi57BCInmDgj3vhBVeKjPAxfliaUK2L69viLLqgzPQinOxy9PWhZKfX7DGxLLg"
        "nTXzhG4GCLuISJ7j3bBasVZ0wcVGRE9RHgo1LG/FNGgzeIlB8MAEgjqO5coGvaDZjGSrgcefFSsG0vkyyXtc0kKIctwxZb5wbVts"
        "5Ig4wSy6H1yYy7Di8gnOLIaLgnTWCC9xg+jXZZyVgxasYjD9vpW6eve1TJAe5k7bfKiBCWjgF72i48go99krauF4Ncel9daFnlQs"
        "zDF0XPZ+dhK5A94dJAByAd9zSkIrkUfNs9e0c6IEOk32ntBmHAm6jhpPpA6pZ68jd97aQZutqlbQ7Ax6Tz1L0gFPOyZ+WjTs2w+o"
        "bpYgs3W2IOCtSqnizTbULR8DTI+XSqk4Sw9viPXvMDeZXVAbQnZ6EXPwFPIwvsRjaS/BuNjSM7p9A8uNEoecRd3Zr1rco12o9Yzr"
        "nOdNgWJaSJvQi3mNbkNLjp33RywWnLMTyNbgSrUcwG3MI2AFka802e0lAvlq95GXXOws35XTNrOHPoeaLUae2rc7T2IWoomoHtQn"
        "44CZy/Z96mEUleyq0GkZ2sZk6n5F1E31dK5QDDTcj5IHhBjoPAomODdxk8ZPdXUXdsRxyC7BzXuMfQiF3cX28e46hmaWRuVNpg2s"
        "Ik8BQ75cJldo+eyqi1dB7iu5oHc3iR2OQUKGpAuzN6kLfHUWJvKTwHfCUKrN28u49kqxWo27i6SyFODOuyx4D6uxx2QfH1Ep+ZOe"
        "z7A8rxInwdiTuE/U1HXGSaCEYosCK6D2LZOGC9KtbCktY3aFHCNlVJRxbcp79Ns7SQmtou1FbNEIS4/ZJqnp+VhdD5vFJzsgJUGY"
        "KwaQa5jAS6gZjo3RwphB144fFXwSmttd1Fq8wqjcMTMPUXMcoWU7gyMRB6JCAuw2EQ2ZkKkCUB5OCFVMOibQLWDB5Svu0DrqWNSY"
        "ZyQbIZu77cBxwfHAimowslzbqhxpG4//pLI1+5RZPjG/UJ637mLgie8VYi5/l1QDstvo5Aze1JRtwacCROO41vc4ZLURaUN1prC3"
        "ylDoXCJ6WoPbaMH5nuCMvPFnn+AZhRf5Q8d9UvJmV1mGjATw+hPPzRoDhQXjugkrWmcwXpWTrSbUhkRh3gcxAXRjYHbT6CA438Ro"
        "iIXSyifV4oswQucui0XPA3+LXgA+2AshYwE8X3OwVTFYzyxoyGfkEe4sDiS/aSl+tihZRkTpB3O6Fje7LXXQIvCfgi2Vpe2H6PX2"
        "CtqumSWrSvL67S3KuV6K8AyfYLc0cWrzJZ3t27jsKrY1al0YArMbE9VRnXiz0S9GlEtPADpKhCnQPTIhR+CBEYLsoF3jsT09Cxvh"
        "lPxcnU4bxPaKBB8rhXoPHuNSlPGJHZDzn8S1EakPbndQ3U7od2nGhHhAphOpnTGLXNn38nJTP8DrKVFyx3z2IfmWClLHnX0kJXYk"
        "nZ4Y4gNhPUVXG8MNcqJBQnXy3fQwQcfn4G0r2XhfnBjLLtootl1Re3D2BUEiGEDKGvR5wp9m93BLVOwhGUGE4WJAS/gfkxgPym3q"
        "cXZsiP8PpGU7+5LihLySy1rALL6s6o9UIstxE+hrDUqQN6THB5WAT0xn3yoI2JU0ALhz7UpUcoEl03Buk6ku7H5V/+4K26Nh6JTg"
        "FQ/IFlSip3WpTAu3OMaDYJCY0HcEK3c5xP86Mox83rVz52w+8lRfqR20RiqJwLiGojyyh87JKJ9pWpSQ9UaxqDvoYUtv6Xc4J67L"
        "rfjShsWN2/8zS0rm2ecYq5OB2b5CDkK12ZtdcizAB6pYgkZjjTx4VbG71BA10O9ZrQldlnUS9La6fj4m7vQp2LwWfvJYUakgM3WT"
        "6hicHUpcI4ULHU72BYafMDAdqUvMmYRfijoKjiVVjZU39DpnBKyMbxjD+vu6GggdAF0mv7dq7NboPKDHemsWKy6NJmfTEMq8XeBE"
        "aXDijIeVj/n4mc5bVknpCXB3UtNg6r/B+7304bfo6v2eGGZjCd8Jrtbkzm0sk49MXkcl4whuz7KJ10ANJzmraNygtxFCq2rnfudJ"
        "wdC4iXYOvUVTswLaxY0tusNdkOZlYwPv6kVOBfNuCLkTs1OisYIWjRyUJ9Jt6cI5OwG799DJ40op232BXIfF4MYyKi/7UqFqrKLH"
        "u+yV0dXIB7PrlFiCUuOjf5+NrAZlfJW/bvMvHJTPYBq5A3dk5HrJaD3Ipe9hTjll6Fd+GJkU+NhXH3Sbe607NBptMkx4DO+RaiQH"
        "eoeuYOBRQALt0ktHTlUMPR+qM4OXDrPGfdR/42LynSTYn34/SnVimFLREPIxe6AkYGEB2zVSHrMd+uIKvpbjVy74uXO+VKpXKEbG"
        "dvBtGSG1ItSJehtdsoCkaFiaQakMHybABgoGYXvGmm9gCf1gFzyr5NOjDRs4HHCX2AfxG8i1soTpDCwrj890bTwlU12s+nwQq0iM"
        "N3gOwLmKZvClFmE/c2KCmcvykQkx0BBiwCO9hi6TC640l2auKJ+CGNyS0CE4Nw3sUskvrg6gwfw1Ng/VDNpiEIGDSelieJnHw+A7"
        "2ELcyklVE945xrr6qa6h4YBuHExXehr0EtWL3to9zDgBR1zwXp4K/63fIjsvBi1Oc9lv3GAF2a3BCRixGw7IvY2H+BBqGFhEXb8j"
        "M/YGeelgXkAVFgLLg7Ds14aGxR7ghVnvyhfHwzE5GesYQw8xO65QUKp+JG3a9cskm4SVJKoeteQIr1/h1YjJOCHBBMnNcCwuUWYv"
        "fA4YRoI1gvWrg0sQwQTepeUfMO5fDx2X/VmH/gxrv3HExa+yT/rP+jQzVhDZahjWb5Akw/QEno64yXw9gVxPVaB1Fj0XvIvXZ+lH"
        "yhO/XucHgfjZuzAv6efIU+R4CjKfMCHtI8YNCitmQ2t9Ud6GsB2NuJmCpluzhHryDySF8MIW/EpTraAzQn5plX+lqdaYsVf9PCtf"
        "G2/RmhypDMU1olmnA3qxgrVGbqCOCGkG1Bxvku+pQk/IiO1dRwPH81SHG7PImASBXN8bLxQEbtGySxb4nIwBDFIxgBixqEUqpvHI"
        "iaCrghXDXSVq21Nkz+Xku+hgDISanHwWA56GK0J9w7gxgmFcDXxAcH/SCtYbtEgw1gHvWzGTvcsfVstuo8ULGXJH0/rYwIwvy26R"
        "NJTeBmp9YrmEySbeuCl7wI1dv5esCFcZzOvbqcdw4IkesbyENiFZ/HrbeEmnn2ADTziFoVEL8eCESnXo6gr1zbPyEOMf8Ls/uGaP"
        "edLs3djCOY5wWwSkCfVwUodB4VBhaRt7yA09OjNab0mHv6uRUK5DiG92MOeUU8Q8YThVJxN/7ca2XDJJVJw8+qESeL3qpBuXxmxe"
        "myja5BMTJmYXBUXRlb/vowNXiEohDEx6dk1fIu+Hqyl2C28w6CQaQ5ZOTcKHqeM8vjQTy8PR9S159Jy3E+50S4ZBphG3JULMe7Kq"
        "HtPWE2PHysI6HfBWbFbXNsDOKqqYIcyvHinzbYO25ZhdwwfGKnC6Ih1cz9AmRqMGBnJmg11HMaZr6DWe6SG0GmJmnVJD5DXax6i4"
        "s4Ok13giV5rEP0WNwYnoMnfMCSGWyLUID0ZKyAbKNGHYsKM99/Gvf/urX8Hz54SH5SpAv9FAwJc/Nog2EaLTwNI3K7qBEJ0G+PZv"
        "DJqbCNFpYEP8uUGzhRCdBlj6Xxg0txCi0wBb/0uD5jZCdBpg9n9l0NxBiE4DjP+vDZq7CNFpQBz8jUHTQohOs4b5qgyibQLpVDDx"
        "H5tD3SaQRjUzjwcbcVUmh5hZ4Myb8B6cAJ6nRw0It4hZTjy7XKVTVHgkIinZ5JwsmUWX8MjcdRNyGLECeLWyyGGorPygL63iEQ1l"
        "2+wrusaOrxBzKxNsPekl2TraWTJupaSAjtcS1cxsDqbSktbMoKvMSL/bw9cFUJslkhvaUPS35qYk7kN15ziSVPAyD5PL9h7MaS9R"
        "VDG3kHzkApV6qnJLz4FnYG5J7R8scXhUj8BopYDvKiTNDHtxkWKrKmSRKF2fUhwFuQQ1JQkpyFjwwjF0L5y1laXQSpeWr0V8q7Kg"
        "zWynaWTPNJp2mga/kaRSm+kYbdFacM9AoMTjIjv9zSdEF2OQhHIPOg3LF8zRfVqdz8zQbWy5TsdVCfo9hfwdMoz7qiGUhiLJh8p9"
        "UDKsh28spGFNE4bR5ty53sYTyoCjvROScnj0UH6BWYDZhFiBwEgfGokiBAzStF9WwKq8JyoHq7dxxUSiv1KvAduDb2Kc8mUScAwg"
        "OamygC50OGxCJoybmSa/kjScLtNrNK4wtiE8YD/Is8aWuu6oDprrE4lOzy6aOp4XYXbkCmRoJJl6T6cLKwd7QY3zW93TdwxtGG7N"
        "FLqJixFENbO7p36NjiaK4PHP+1WOEqhP62AZal+/TmqlsNkja0gdD9dnKYYIrzVp6SIFpm74K3JfwTZlVINy/hYcoMYHwDVCWIYh"
        "6AvcnLmU3wPxYO0Sej5dF7HWojwC73TUyfDXVS31VX1cppA0M2/WL1I2VzxfIwi9XQB51kODcovMxQiEQVhJ3FD1XQxRET2gJtSX"
        "iS850UEr8AN5XNLAe39+haN4Gsia4J1jiP9T0SPbk3RjpGBCpygoGd+U1eHkW4BnuTR4b/0dh+PyfPbWKblGaKsJ3q7zgYiCNOjC"
        "eSi6kcDmWAjYFxIYKsdCc04g1zC5XSmfENUxFE6Ybk6YfLKB3icBYR5Wf4aHG6MqPKfRQOM3GcrGPHkZC2XYpnS81Vggr4MG2p5F"
        "nXsEwj/YwbaIc4K5maERBHyFPSxUwUPFrvtpbMIPZUqUxhba9nWZRjOECbLUxattOuTH3YfmBtztKcJDEQh5xW45WCCqsjqadaAi"
        "c+p76R1uvaf4u5Fykim//o7CP4qh4IG8nN7THSXSvV8mn5CnDzN4uoRBTol/tv4InTYBC/wu3UsY+j2H6F1MAtaI2+2q47CaE0a+"
        "5Gp4eAkh00GhEkstv/mQpikUVqObgLtXMFca2CoCQwEsdcyvblXtGpuiS3Sp4kIYOXxsXX8s/Vlg5lf53LX+VEL1k48OBVKPjQm1"
        "SqbV7bzCpVYoXFCxn11MkuAcPBHMiwB4kCdG3teaO0FnLRgSiIcOvbVbWA5NlzWZvxZ/tDnmAn/cpRC2mFBbMm9Njo4csfLLDeym"
        "4OLJscM4JoCzyt/7XmloOJmp5ybnpCCspFidQl9CeSy2TCFS8jXk49LqxuQiX0IdNctVoQydF2htSibUBYl6GqIKLBmN3tkjTcKt"
        "eJBnhwfnDkV5qMStGFuQBLkt7SXBWwl0mwOj42hMh+7jms77QQVVP/LCkMUL/hd0Ro+qBJrzGLdgVwuc38IaVZHH8+i5GhpWe2Me"
        "WO3vbRn72eGjfWaxnVaSAjFmF2lnmwKJZDs6pJupZzjmV9gjx3nsMJTLETowH0Zj0JoV0G7ceULnLxG9uARxVAU74RY7uIZdh9xg"
        "OxjYUPbllpxfw71VcThAZ36dIvCYg3XpQMCNipiake9Z0BNCbiEvtAAmnN/kXmrn4/M3MLt5HPFy3Jgl2VUBzhJhvlxaMwvyXCLA"
        "i6jU0OdoWEPUCQRz8SzO36T9EKCCxrxtfksKMY15/oisBJSYEN6Bk+eL79GnYKHHbTipgXT7gz0xA5qPbJ5iAKpmG/Q7ttqczt9R"
        "xNoHuw9xXCyxA6ukvcyTb1XGos0vyGLytYw6sInvLrjJxRZMsnXCB5ekJOrc0WRpVYwl62ldXewL2ws1WPSPSfcWA/+0tfieQgfs"
        "CKStjMDEA1h8QEp0o8hX82fZmKJfGItc0M/hmqAv4KtG6qEeUM6FCIn4/hN1pDlFgkVBOeFtQBlZEQMNZznUvJZFnryQ3Uc/nUVf"
        "9ssZtIt3UrQExcsLdskqsmhefMSHdOr6QrNJJ71SV2wuZfQ9pxHUgfOcBiZG4fQqdKaOXuM4dCqeDC2r7+kg3h5rnGLN8asQWCub"
        "VufgZ3zFxi5J/299A6UqXBhLYC0ZlasKd/jxH7gFITv/lq5mOsRl6ijlrR+SMm/I+qjWpIuuO02jAzFwBEB/faUCCqcs9QpvjERW"
        "wJ7POr4M4FrnHSvia4t1es3AkSHFS3Tc+T2zu4v8XXni1gV+aOEJPosrAztJAdLWCMcFdMmRqoLpu9dwTwasUi8+QW9F3hHWn2ur"
        "iX+aghp59Hv4UM05nf5mAtDHFIPdLBn7tngL+ekPWn7YJkZmgHQJOAqxK5OiWyiJdVRzBk3oaCyMbBW0hSnHvrBqkU88Hp9DP2GH"
        "ECg/wu8+EOIeLo1YejPRsD9n4REsUzzR3opQV13wfe9zECam1OiFd3y8DC9NEORH0m40XgIM5pMT1HFMdfUN/eg8RG+1JYOw6xSY"
        "5VYFk7T5Dlb9JsNCvq5D1bxjaC3wv7fVGU39vqzAtaRjZpHu3wdeaKvd/RKNhsgXRFbAOt7iKxOY8mQsvsa83mI5a41b3E+AYRKB"
        "3kQLIK7Z6czNzbqB0HxgS5grX5nZC7c5pKycV7HCzTbNEjAWFw49yIx4yCZm7itLKKiBayWGd7dJxzMJ9WKTeJAzlheaPqXFZ9P5"
        "vF3NCaGsnrdaXEo5WYR2J2NsF5dVFmgduqI9bJBAV5OHDRLgmnouoQhHv2aJdYqkAG4s1qQ83Vlc1NtOoAXSywPLzbsy18giZrgQ"
        "+xPTfcOph9hlvosvxh1Xoe2LHJcYCOGGp2YEvYImJqRZCMEpy2cmWIzjUxcfs6c/ZjmIjwoneazpUzmdYnGeNz/EYEgRuIzXGUPj"
        "GYsGLMxPIQaGW9ldo0sX8tLAwmNWuUKZ87pbZy0zGosCqYMsgNL0d7Y3BLEiSqgsYGo6fB4xhdjD+1OBnwK/xNjW2Cuk4K9kI/BN"
        "TzIk6YYUH+k/I5uk4GhxTh20+iHNN9G8piQ3lFk9dyoOzpPzEJqd5OHvbZJTQWhSse5BWQS59enQlzLIbvEeKb8lmVMno8gOBloR"
        "JPUayqKMwYFgOktO2zYDoV4FbONcSp/A4m6fKxK/LPhBxClMoAai7R5Ki47AhPi+2mKHVflgAFWqsodyEAlN6xZm6euAnebfWmV3"
        "zGaLqLeJp8JwzsvKX29zwbjwhjfqKG0xiYbreKy577o6cFpe5JIZ3Bk+Q447qQAtz8rkhVXLTdTNxRt0ScDh8EC9PZwp9Bg85qGZ"
        "GosX8fZZQQZ3L5KnGAIJPThayGm4lUt8m8F2giTf4eJVhFKObIJQ8E0AJwYVHYFPvxNDUswHH3knjmTQLk7iyXtk0SklfpUQU3g8"
        "XhwVKmkEMk6y32v41T7wNBkdkADYsUqBtCYWr5OlWK2KhkoWN0Mh2QdbY6PqLtriLB1fjFQSo3uxLkPXKvg0Ds/SJB23ly149kJH"
        "swxuoIng2WzgLs7h4RS+1Y335hWut8mZWuQFlpXLauj1BXlFzigeM7MHdGVerQIhqMiR01zok016AFpzURNPBqKRhKoR24IVKzny"
        "XAo5bKrpTb4/rcPwxXmLjgtS1PfowoMczS5GGvyAe2mjwUEptNkwnoZIcZ0rjDCVb5moYT5/T5sQ+Py8YO0UOb1OryFxKnV64y1Q"
        "ycfvopNRety7eKkrBpVG+j1f0yV5h6L1j6tjl43byqGixVfgS+To14DnEanvHZatWU3tbujIMUdpZc0dtHnwcpneV+zRGr98IINI"
        "VjaTWuBhYor5w/Chby1LvpfwKEmOkOQia67Quq2E+GIcTsgi5ezdNl6Ryp1ijU7gUTtTXs4uPWch3TQdCg4IKhBXoS3oDiUiLXks"
        "LLqbKPMVr+nycoQQBFoJK6SxQ95Nsd519bTTxc0MIUtjKtK5s4sXqyh2SPolOngj9/Qx6ux9afGFFTYLmg/YQlGQxdtaDAZOmDZf"
        "C9c4QlK6FnZekWsI2JL2KBPyeqsmi/GTWsUI7jyqu7E7+3Q7Ud0i3XlD7EjY1ao5O2h/Hux5EZz6qTgeGIBzggvRWlzhEKgkVKp7"
        "M7mrQwC8a2wnT9zhvjsu2ePGdTp0oEDwL5SO3KUoIlktHl+c853Ei9m9RRenS6XEuF8HG+dErATRxkXy7tuj6ppCt8VpDOxiciF5"
        "AU3tEJ5S5Rs8z/FKmLz00XlpeiY1TykmFlWJHDv0MCAzyM7bJFw+0W67d9HxJllSZx8Dm+SR0m6H7typIOpt3nN5MejSH9qR2Sak"
        "4Ok8kgeCwrBOLoJsLGg+Pi3Q7j5Gv5SsQG6tdYr/4bD1bhuvL1t56VPqrT9hLYW+9h6DyceEBnusGrtKTb7PSSDILFuh9A4cUCzQ"
        "mGwAc7xEOufii9ujMpytgXwLuUKuN34TVFw+BeiNb1L7HxoUiVhSBI/lXYpRi28p9NbeG4VUgJ8q81o5M+n3G+kKpJ9vtQXIoH2j"
        "RrhMoypDh+cnNAmXOAat4nuakbNxGbONum4KfCUJl0I1mCUEhi0oeDEvlnRZxc/dQoYbeNLUWX/FD8QIJo6BzH+I4Yoqpv0cTgIj"
        "NpoYQYunl/iKg/JyrQNL+8aAPJPRjgctqNYkf0B36EfKBCbgI4qCE9ssRY25a8BmMadIYF6aeTzNcosvUk4g28Q3b5IXRDTvewjt"
        "g7dZNIIePh2LHjWLU3/38FHX444QrkUoUeABvT+hhK/+glJvYsWE82tbEn3zdhYaC65mFpSPrfXuT5p4hZhYMxH0ElhS7moGGsst"
        "Z5VT/Vjq7x/y3+S7zezvFvROPcz0Mx9VakdqjaPwZlFWCSTrHkaWVIcpzgfSyQ5jOvGBVMmbd83L8qHM4sFeUMXngYDOhuiRr8sB"
        "vkEWhA4n0WtekaNk50r4FipkR8ugw3OdY5+cGkqM3SZecD19QgeBUnbm9J/pINSwzpw5rcNA7zl35isNtAFacwBvjjuQmVpscswP"
        "gMwBKj0m+Kfu4lgCdcLH44hNOlY57xRCjq7q4dOyJ8pWULWGzqKq18Mnr//OoVXyd6Sk9jZvYOgmCAmOzFBHoPiytB8IhaLg/2Ag"
        "unhtxLVIrx1W33xIkQAQZBEaTa3jIYDns7eviXcZ7R/I7CJDHJ6TJGGCDrwy5/Tp4euwlHwp93P8FNf8C8Reo3w+2chpusDuYihr"
        "FsF1mbUkCzkjU2tlIWf5wS5IjJ2Fr0v9OQvZ0FKCZOHn5E1feME3A096FVwvEGpTFsGCkdQ7i2JR3tzN+sBGXUuFg08MS3d7cxzF"
        "DQiEsOoX2N3cxOg8Bz0Zvrzo0ryEsRqCz+dGDVRveU+Rs2zsNehKbq1mw5UY2IspHoEXqU7DmSks8RzWi/BHRl05H9bJVxYdt/ZW"
        "xhUL5IfTsQzehqHfYS3mVqH/9xtfaIcBHvrKT4ie7OueBbyC7Eb8th6W3M1svHqzr/lO80BkFL+fWVzxcYxvUV1CoxIDpvEX3Woo"
        "UqwqQhpSFwukRVFvkyrhKztNtOkqn1TYuRNOVWpIzQnKuS8s+SFEktDew1SU4CpRLHXYKXngYsLzPw9OmeVZ0wbeZvzi2Gkq+wwd"
        "+3mKKdPWzOLz9HEQHY7oiYsTu1LPeZxA1yl6z/JCSo06NKxQm8t4VwT/vgmisuZS9TcR4Xj8Cz3tGHUS5VAOYuO20DuM79ejscCb"
        "I2f/IDSvIinLvTW8bwoPXcPliwpd0euNw2Ip4GyCn7zvieXexWs0QQAekjk0SzrFjUWVSLEP1dQSJ/YhlzRk/4dvrKhq+5B/HH8s"
        "4x4915HZOxaesH96FNNksRn7VDm/M4eFH8MKK7ErjFiKGOqtoSVqj1wYAV6szRPeDvpC5dNsPqGzVmdMuuUXMbkdRgFoFkfzcSqe"
        "8JhYmqhv6i7vB0t4FQZz4VDw4ClOC9DbvKpnoDxoycvavdV5DP4u2eitgxxtrTBSPvQm6hYHrWJECS21gjdmZPyQPuZ1fjiTXUip"
        "+m7KG59sO/Zuvle8QItS/OM4XQIHkQu8mvaS9plN2GUQhwrDhtxLCeXN54yCaOqyayJvYBo+CjUVv3/awcPK2gXBATHs+6eHFOkL"
        "iS8DuzAk9loBz16g8JU7dCws9r8u7He3EjcHC41bZMICg6bNm9M0jtxnEFkNJ1GwWCQjun1oEb5bZxa501fkO0iDxNi7fdhheDUz"
        "0JvebB361ePmB68ikzAanzsGF0VhqVVzfxtbBelzmqj3kXq5Y8ezia8+6a93AOnEfEa9n3wxgBi1AqHQgaX7ie+D2QyUeKSYIl3U"
        "SE8Gtj2ItHcFnSeqnYLcOaTaZkZzf3cCqPEJdq1ePG8w6v2S6z2lMnugCSRbedoePeTDyxkfPv3ZAOqrP/ZPwJkBY7qSUfGZU4Mq"
        "ft9f8VeDmryKVsOw8ZQ2mpJ81gq3bkW/v/7qz7DfVarEMytZN+m/GYav0XGx+bmpixols3VRe8Wv+NIhMnUpkyS0FMHLfgJ8yh1m"
        "zxdcK8zDU9AlRb83mN7Duy0lUBtCRf9oMH1gQ24euIkjqa/NDaaGWwIHrVoNgKr214Ppz0LAvKp5an0wZarFTwdTglvDEZqHasC1"
        "+cHEVgXv5YEHSNX96hBybxQUooT2+WBaCKQS2jFkFKenUlShZ4MLfYcPAuSKYi8S+R+nVwZTfyYGRWgDILUk9doh1KIRTtXTRvHN"
        "YcR8JKmIXwwmRodVX9MF33+ubhckdujWLRVZYYBvy1SGBvSOlqQuATdfyJANgxhkETI8DdbSc58xuLe+a5wVJ/CudvNKq2RbaXj6"
        "3t7qqJ6kEPdkX1LwHS3+xEC0kxAUA96VPTKgST43nZxSvtJ9qjAne2dSbGoUUkvT8JgGQ79xNqAezMxi0mXV9qOikneBs2t710eX"
        "UdtyV7tTnFnP8q5G0V+DWDZ7qvs5I7Ynd9aKZWjUy4QGZ0THvVIfGFz+dULTV/6NWkcn6BiaA5yE6og5HtIG7tsPok98dAu3yNkM"
        "6V/9IO0BaoLoO0PXkIvyUbhmkjI+9wcaK9LK0va0IH2qJ0s/grhxkdLguEXXKrEJu4638Bx1Hb4Oqr3+qDa8PFaSRzGN5/rnTibV"
        "NLQmG2B6A9e8j1dMEfQbt3oV+FzJKt1rEJ3Tiv8EnOgTelT4XytLDRB3ZRAS6v98O+enbQpwgnV5CYsGsWCbdCXi0l3M/Sas+4O9"
        "kLMqCKPTG0FkCy03dOn1Lm1TlLSrGSWX6C4j/X2fA6Pwx2Pyynge5oHpXaLH7ENCPuMkwAEZ75f28DPnAzsX+XHg+ZTVt3eJku/a"
        "YdVxqeBrTDVZoIk+G9cconuHt5joFmTv0ns8LIhGfEcPjvpV7/Jd6Q/5be5XAvKxAflYQH5tQH4tIL8xIL8RkD83IH8uIH9hQP5C"
        "QP7SgPylgPyVAfkrAflrA/LXAvI3BuRvYOZeo3FLnnu8aQlRh6AZ//S0D/PdKEbJ/PSsDwMbFkfmHo4TjdpPL+RhtzZI/+8lPSYj"
        "5uWf/j9Anzly"
    ),
}
_cldr_cache = {}
def _cldr(lang: str) -> dict:
    l = lang if lang in _CLDR_B64 else "fr"
    if l not in _cldr_cache:
        _cldr_cache[l] = json.loads(_zlib.decompress(base64.b64decode(_CLDR_B64[l])).decode("utf-8"))
    return _cldr_cache[l]
def _emoji_nom(emoji: str, lang: str):
    """Émoji → nom CLDR, ou None. Tolère la présence/absence du sélecteur FE0F."""
    d = _cldr(lang)
    for cand in (emoji, emoji.replace("\ufe0f", ""), emoji + "\ufe0f"):
        n = d.get(cand)
        if n:
            return n
    return None

def _arasaac_resolve(mot: str, lang: str):
    """Mot → {id, dataurl} ou None. Caches négatifs inclus."""
    key = f"{lang}:{mot}"
    pid = _arasaac_ids.get(key)
    if pid == -1:
        return None
    if pid is None:
        pid = _arasaac_search_id(mot, lang)
        _arasaac_ids[key] = pid if pid is not None else -1
        if pid is None:
            return None
    b = _arasaac_png_by_id(pid)
    if b is None:
        return None
    return {"id": pid, "dataurl": "data:image/png;base64," + base64.b64encode(b).decode("ascii")}

@app.post("/pictos")
async def pictos(request: Request):
    """
    {"mots": ["araignée", "l'insecte", …], "lang": "fr"}
    → {"pictos": {"araignée": {"id":…, "dataurl":"data:image/png;base64,…"} | null, …},
       "attribution": "…CC BY-NC-SA…", "lang": "fr"}
    Un mot introuvable vaut null : le frontend applique ses replis (émoji, sans image).
    """
    body = await request.json()
    mots = body.get("mots") or []
    lang = (body.get("lang") or "fr").strip().lower()[:2]
    # V2.18 — la porte des numéros lit "ids" ICI pour qu'une requête NUMÉROS
    # SEULS soit valable ; toute requête SANS "ids" suit exactement les gardes
    # de la v2.17 ci-dessous, messages d'erreur compris (prouvé en batterie).
    ids = body.get("ids") or []
    if not isinstance(ids, list):
        raise HTTPException(400, "ids doit être une liste de numéros")
    emojis = body.get("emojis") or []
    if not isinstance(emojis, list):
        raise HTTPException(400, "emojis doit être une liste")
    if not isinstance(mots, list) or (not mots and not ids and not emojis):
        raise HTTPException(400, "mots (liste non vide) requis")
    if len(mots) + len(ids) + len(emojis) > ARASAAC_MAX_MOTS:
        raise HTTPException(400, f"maximum {ARASAAC_MAX_MOTS} mots par requête" if not ids
                            else f"maximum {ARASAAC_MAX_MOTS} entrées (mots + ids) par requête")
    out = {}
    for mot_brut in mots:
        mot = _arasaac_norm(str(mot_brut))
        if not mot:
            out[str(mot_brut)] = None
            continue
        out[str(mot_brut)] = _arasaac_resolve(mot, lang)
    # ── V2.18 — LA PORTE DES NUMÉROS (ADDITIF STRICT) ─────────────────────────
    # Condition : le corps porte "ids" (liste de numéros ARASAAC). Effet : chaque
    # numéro est résolu par la brique EXISTANTE _arasaac_png_by_id — aucune
    # recherche, aucune ambiguïté. Une requête SANS "ids" est traitée à
    # l'identique de la v2.17 (ids vaut [] et cette boucle ne tourne pas).
    # Un numéro inconnu ou invalide vaut null, comme un mot introuvable.
    # Voir JOURNAL_BACKEND v2.18 — famille « un mot désigne mal une image ».
    for id_brut in ids:
        cle = str(id_brut)
        try:
            pid = int(str(id_brut).strip().lstrip("#"))
            if pid <= 0:
                out[cle] = None
                continue
        except (ValueError, TypeError):
            out[cle] = None
            continue
        b = _arasaac_png_by_id(pid)
        out[cle] = None if b is None else {
            "id": pid,
            "dataurl": "data:image/png;base64," + base64.b64encode(b).decode("ascii"),
        }
    # V2.18 — la porte des ÉMOJIS : nom CLDR puis chaîne existante ; la réponse
    # porte aussi "mot" (le nom dérivé) pour la rangée de la fiche.
    for emo_brut in emojis:
        cle = str(emo_brut)
        nom = _emoji_nom(cle, lang)
        if not nom:
            out[cle] = None
            continue
        r = _arasaac_resolve(_arasaac_norm(nom), lang)
        out[cle] = None if r is None else {"id": r["id"], "dataurl": r["dataurl"], "mot": nom}
    return {"pictos": out, "attribution": ARASAAC_ATTRIBUTION, "lang": lang}


@app.post("/generate")
async def generate(request: Request):
    """
    Proxy vers l'API Anthropic.
    - Gère les erreurs rate limit (retry automatique)
    - Gère les documents trop longs (chunking)
    - Retourne la réponse Claude
    """
    body = await request.json()
    api_key = body.get("api_key")
    if not api_key:
        raise HTTPException(400, "api_key requis")

    model = body.get("model", "claude-haiku-4-5-20251001")
    system = body.get("system", "")
    messages = body.get("messages", [])
    max_tokens = body.get("max_tokens", 6000)

    client = anthropic.Anthropic(api_key=api_key)

    # ══ v2.13 — LE CACHE SURVIT À LA DURÉE D'UNE PARTIE ═══════════════════════════════════
    # Constat des tirages (pied de page essai_10138) : « 10 009 lus + 0 relus +
    # 160 006 mis en cache » — on PAYE la mise en cache à chaque partie et on ne relit
    # JAMAIS. Deux causes dans le code v2.12 :
    #   1. le cache éphémère vit 5 minutes, or l'écriture d'une partie par le modèle en
    #      prend davantage : à la partie suivante, le cache est déjà mort ;
    #   2. seul le prompt système portait un repère de cache — le DOCUMENT (le gros du
    #      volume), envoyé dans le premier message, n'en portait aucun.
    # Remèdes : durée longue (1 h) demandée sur chaque repère, avec REPLI AUTOMATIQUE en
    # durée courte si le serveur refuse la syntaxe (aucune génération ne casse, la réponse
    # dit la durée réellement servie) ; et un repère posé sur le dernier bloc du premier
    # message utilisateur, là où vit le document. La réponse renvoie désormais TOUS les
    # compteurs (lus / relus / mis en cache) : la facturation exige des coûts prouvables.
    for m in messages:
        if m.get("role") == "user":
            c = m.get("content")
            if isinstance(c, str):
                m["content"] = [{"type": "text", "text": c}]
            if isinstance(m["content"], list) and m["content"]:
                dernier = m["content"][-1]
                if isinstance(dernier, dict) and dernier.get("type") in ("text", "document", "image"):
                    dernier.setdefault("cache_control", {"type": "ephemeral"})
            break

    def _pose_ttl(duree):
        sys_blocs = [{"type": "text", "text": system,
                      "cache_control": {"type": "ephemeral"}}] if system else []
        for blocs in ([b for b in sys_blocs],):
            pass
        tous = (sys_blocs or []) + [b for m in messages if isinstance(m.get("content"), list)
                                    for b in m["content"] if isinstance(b, dict) and "cache_control" in b]
        for b in tous:
            if duree == "1h":
                b["cache_control"] = {"type": "ephemeral", "ttl": "1h"}
            else:
                b["cache_control"] = {"type": "ephemeral"}
        return sys_blocs

    # v2.13 — le client peut IMPOSER la durée ("cache_ttl": "5m" ou "1h") : le
    # frontend sait mieux que le serveur si plusieurs parties sont probables
    # (taille du document). Valeur absente ou inconnue : durée longue par défaut.
    duree_cache = body.get("cache_ttl") if body.get("cache_ttl") in ("5m", "1h") else "1h"

    # Retry avec backoff exponentiel
    max_retries = 3
    for attempt in range(max_retries):
        try:
            sys_blocs = _pose_ttl(duree_cache)
            extra = {"anthropic-beta": "extended-cache-ttl-2025-04-11"} if duree_cache == "1h" else {}
            response = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=sys_blocs,
                messages=messages,
                extra_headers=extra
            )
            text = "".join(
                block.text for block in response.content
                if hasattr(block, "text")
            )
            u = response.usage
            return {
                "content": text,
                "usage": {
                    "input_tokens": u.input_tokens,
                    "output_tokens": u.output_tokens,
                    "cache_creation_input_tokens": getattr(u, "cache_creation_input_tokens", 0) or 0,
                    "cache_read_input_tokens": getattr(u, "cache_read_input_tokens", 0) or 0
                },
                "cache_ttl": duree_cache
            }
        except anthropic.RateLimitError as e:
            if attempt < max_retries - 1:
                wait = 2 ** (attempt + 1)  # 2s, 4s, 8s
                time.sleep(wait)
                continue
            raise HTTPException(429, f"Rate limit après {max_retries} tentatives : {e}")
        except anthropic.BadRequestError as e:
            # v2.13 — repli : si le refus vise la durée du cache (syntaxe ttl ou en-tête
            # inconnus du serveur), on repart en durée courte, une seule fois, sans
            # consommer de tentative — la génération ne casse jamais pour du cache.
            msg = str(e).lower()
            if duree_cache == "1h" and ("ttl" in msg or "cache" in msg or "beta" in msg):
                duree_cache = "5m"
                continue
            raise HTTPException(400, f"Prompt trop long ou invalide : {e}")
        except Exception as e:
            raise HTTPException(500, f"Erreur API : {e}")


# ─────────────────────────────────────────────
# ENDPOINT : /export
# Structure JSON adaptée → DOCX
# ─────────────────────────────────────────────
@app.post("/export")
async def export_docx(request: Request):
    """
    Reçoit la structure JSON adaptée par Claude.
    Reconstruit un DOCX avec :
    - le texte adapté
    - les images originales à leurs positions
    - le layout préservé selon le type d'exercice
    """
    body = await request.json()
    exercises = body.get("exercises", [])
    filename = body.get("filename", "tailory_adapte.docx")

    doc = DocxDocument()

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Andika"
    style.font.size = Pt(12)

    def add_image_to_para(para, img_data: str):
        """Ajoute une image base64 à un paragraphe."""
        if not img_data or not img_data.startswith("data:"):
            return
        try:
            _, b64 = img_data.split(",", 1)
            img_bytes = base64.b64decode(b64)
            img_stream = io.BytesIO(img_bytes)
            run = para.add_run()
            run.add_picture(img_stream, width=Inches(1.5))
        except Exception:
            pass

    def add_exercise_header(title: str, num: int):
        para = doc.add_paragraph()
        run = para.add_run(f"Exercice {num} — {title.upper()}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6e, 0xb7, 0x9e)

    def add_consigne(text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(13)

    def add_response_line(label: str = ""):
        para = doc.add_paragraph()
        if label:
            para.add_run(f"{label} ").bold = True
        para.add_run("_" * 20)

    def build_relier_table(images, words):
        """Template fixe pour exercice relier : image | • | mot"""
        if not words:
            return
        table = doc.add_table(rows=len(words), cols=3)
        table.style = "Table Grid"
        for i, word in enumerate(words):
            row = table.rows[i]
            # Colonne image
            cell_img = row.cells[0]
            if i < len(images):
                para = cell_img.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_image_to_para(para, images[i]["data"])
            # Colonne point
            row.cells[1].text = "•"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Colonne mot
            row.cells[2].text = word
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    ex_num = 0
    for exercise in exercises:
        ex_type = exercise.get("exercise_type", "autre")
        consigne = exercise.get("adapted_consigne") or exercise.get("consigne", "")
        blocks = exercise.get("adapted_blocks") or exercise.get("blocks", [])
        images = exercise.get("all_images", [])

        if ex_type == "header":
            for block in blocks:
                if block.get("text"):
                    doc.add_paragraph(block["text"])
            continue

        ex_num += 1

        # En-tête exercice
        add_exercise_header(ex_type, ex_num)

        # Consigne
        if consigne:
            add_consigne(consigne)

        # Corps selon le type d'exercice
        if ex_type == "relier":
            # Extraire les mots de la colonne droite depuis les blocs
            words = [
                b["text"] for b in blocks
                if b.get("text") and not b.get("is_consigne")
                and len(b["text"].split()) <= 4
            ]
            build_relier_table(images, words)

        elif ex_type in ("compléter", "numéroter"):
            for block in blocks:
                if block.get("is_consigne"):
                    continue
                if block["type"] == "table":
                    # Reconstruire le tableau
                    rows = block.get("rows", [])
                    if rows:
                        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        t.style = "Table Grid"
                        for i, row in enumerate(rows):
                            for j, cell in enumerate(row):
                                t.rows[i].cells[j].text = cell.get("text", "")
                else:
                    text = block.get("text", "")
                    if text:
                        doc.add_paragraph(text)
                    if block.get("images"):
                        para = doc.add_paragraph()
                        for img in block["images"][:2]:
                            add_image_to_para(para, img["data"])

        elif ex_type == "classer":
            # Tableau découpe compact
            all_items = [
                b["text"] for b in blocks
                if b.get("text") and not b.get("is_consigne")
            ]
            if all_items:
                t = doc.add_table(rows=len(all_items), cols=2)
                t.style = "Table Grid"
                for i, item in enumerate(all_items):
                    t.rows[i].cells[0].text = str(i + 1)
                    t.rows[i].cells[1].text = ""  # case réponse

        else:
            # Fallback : texte + images
            for block in blocks:
                if block.get("is_consigne"):
                    continue
                if block.get("text"):
                    doc.add_paragraph(block["text"])
                for img in block.get("images", [])[:2]:
                    para = doc.add_paragraph()
                    add_image_to_para(para, img["data"])

        # Séparateur
        doc.add_paragraph("─" * 30)

    # Sauvegarder et retourner
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ─────────────────────────────────────────────
# ENDPOINT : /convert (existant — conservé)
# PDF → DOCX
# ─────────────────────────────────────────────
@app.post("/convert")
async def convert_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Format PDF requis")

    uid = str(uuid.uuid4())[:8]
    pdf_path = f"/tmp/{uid}.pdf"
    docx_path = f"/tmp/{uid}.docx"

    try:
        content = await file.read()
        with open(pdf_path, "wb") as f:
            f.write(content)
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{file.filename.replace(".pdf", ".docx")}"'}
        )
    finally:
        for p in [pdf_path, docx_path]:
            if os.path.exists(p):
                os.remove(p)


# ─────────────────────────────────────────────
# ENDPOINT : /health
# ─────────────────────────────────────────────
_arasaac_probe_cache = {"t": 0.0, "state": "unknown"}

@app.post("/figures_vues")
async def figures_vues(file: UploadFile = File(...)):
    # v2.30 — LA VOIX DU DIAGNOSTIC. Condition : un PDF reçu. Effet : la liste
    # des figures telle que parse_pdf la fabrique (MÊME chemin de code, aucune
    # découpe parallèle), sans les images elles-mêmes — ordre, page, cadre en
    # pleines décimales, sceau du contenu, et l'empreinte du fichier reçu.
    # (voir JOURNAL BACKEND v2.30)
    content = await file.read()
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(status_code=400,
                            detail="figures_vues ne lit que des PDF")
    r = parse_pdf(content, file.filename)
    return {
        "version": VERSION, "empreinte": EMPREINTE, "lecteur_pdf": LECTEUR_PDF,
        "fichier": {"nom": file.filename,
                    "sha256": _hl.sha256(content).hexdigest()[:12],
                    "octets": len(content)},
        "nombre": len(r["images"]),
        "figures": [{"ordre": i["index"], "page": i["page"],
                     "cadre": i["cadre"], "sceau": i["sceau"],
                     "w": i["w"], "h": i["h"]}
                    for i in r["images"]],
    }


@app.get("/health")
def health():
    # V2.9.2 — sonde ARASAAC (mise en cache 10 min : /health est appelé à chaque
    # chargement du frontend, on ne frappe pas l'API à chaque fois)
    now = time.time()
    if now - _arasaac_probe_cache["t"] > 600:
        try:
            data = _arasaac_http_json(f"{ARASAAC_API}/fr/bestsearch/maison")
            _arasaac_probe_cache["state"] = "ok" if isinstance(data, list) and data else "empty"
        except Exception:
            _arasaac_probe_cache["state"] = "unreachable"
        _arasaac_probe_cache["t"] = now
    # V2.10 — la version du module grilles est remontée telle qu'elle est
    # DANS le fichier déployé, jamais recopiée à la main : c'est ce qui
    # permet de savoir, en ligne, laquelle tourne vraiment.
    return {"status": "ok", "version": VERSION, "empreinte": EMPREINTE, "lecteur_pdf": LECTEUR_PDF,
            "grilles": (getattr(_grilles, "VERSION", "inconnue")
                        if GRILLES_ACTIVES else "absent"),
            "formes": (getattr(_formes, "VERSION", "inconnue")
                       if FORMES_ACTIVES else "absent"),
            "arasaac": _arasaac_probe_cache["state"]}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
